"""Structured resume → DOCX bytes (strict one-page, template-matched)."""
from __future__ import annotations

import io
import re
from typing import Optional


def _build_docx_bytes_from_structured(
    structured: dict,
    accepted_edits: Optional[dict] = None,
) -> bytes:
    accepted_edits = accepted_edits or {}
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx is not installed.")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def _set_font(run, name="Calibri", size=11, bold=False, color=(0, 0, 0)):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor(*color)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:cs'), name)

    def _add_horizontal_line(paragraph, color=(0, 0, 0), width=4):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(width))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '%02x%02x%02x' % color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing = line_spacing

    def _add_section_header(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        _set_font(run, size=10, bold=True, color=(0, 0, 0))
        _set_paragraph_spacing(p, before=6, after=1)
        _add_horizontal_line(p, color=(0, 0, 0), width=4)
        return p

    def _strip_bullet_prefix(text: str) -> str:
        text = str(text).strip()
        while text:
            if text[0] in ('•', '-', '‐', '–', '—', '*', '·'):
                text = text[1:].strip()
            else:
                break
        return text

    def _add_experience_line(role: str, company: str, dates: str, location: str = ""):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tab_stops = p.paragraph_format.tab_stops
        right_pos = Inches(7.4)
        tab_stops.add_tab_stop(right_pos, WD_TAB_ALIGNMENT.RIGHT)
        left_parts = [x for x in [role, company] if x]
        left_text = " | ".join(left_parts)
        run_left = p.add_run(left_text)
        _set_font(run_left, size=10.5, bold=True)
        p.add_run("\t")
        right_parts = [x for x in [dates, location] if x]
        right_text = " | ".join(right_parts)
        run_right = p.add_run(right_text)
        _set_font(run_right, size=9, bold=False, color=(100, 100, 100))
        _set_paragraph_spacing(p, before=4, after=0, line_spacing=1.0)
        return p

    def _add_education_line(school: str, dates: str, degree: str = "", location: str = "", gpa: str = ""):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tab_stops = p.paragraph_format.tab_stops
        right_pos = Inches(7.4)
        tab_stops.add_tab_stop(right_pos, WD_TAB_ALIGNMENT.RIGHT)
        run_left = p.add_run(school)
        _set_font(run_left, size=10.5, bold=True)
        p.add_run("\t")
        run_right = p.add_run(dates)
        _set_font(run_right, size=9, bold=False, color=(100, 100, 100))
        _set_paragraph_spacing(p, before=4, after=0, line_spacing=1.0)
        sub_parts = [x for x in [degree, location, gpa] if x]
        if sub_parts:
            p2 = doc.add_paragraph()
            run_sub = p2.add_run(" | ".join(sub_parts))
            _set_font(run_sub, size=9, bold=False, color=(100, 100, 100))
            _set_paragraph_spacing(p2, before=0, after=0, line_spacing=1.0)
        return p

    def _add_compact_plain(text: str, size=9.5, bold=False, color=(0, 0, 0)):
        """Plain text line — no bullet, no indent."""
        clean = _strip_bullet_prefix(text)
        if not clean:
            return None
        p = doc.add_paragraph()
        run = p.add_run(clean)
        _set_font(run, size=size, bold=bold, color=color)
        _set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
        return p

    # ── Header ─────────────────────────────────────────────────
    full_name = structured.get("full_name") or "Candidate"
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_name = p_name.add_run(full_name)
    _set_font(run_name, size=20, bold=True)
    _set_paragraph_spacing(p_name, before=0, after=0, line_spacing=1.0)

    contact_parts = []
    for key in ("email", "phone", "linkedin", "github", "location"):
        val = structured.get(key)
        if val:
            contact_parts.append(str(val))

    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_contact = p_contact.add_run(" | ".join(contact_parts))
        _set_font(run_contact, size=9, bold=False, color=(100, 100, 100))
        _set_paragraph_spacing(p_contact, before=0, after=2, line_spacing=1.0)

    # ── Summary ─────────────────────────────────────────────────
    if structured.get("summary"):
        _add_section_header("SUMMARY")
        p_sum = doc.add_paragraph()
        run_sum = p_sum.add_run(structured["summary"])
        _set_font(run_sum, size=9.5)
        _set_paragraph_spacing(p_sum, before=0, after=1, line_spacing=1.0)

    # ── Education (FIRST) ──────────────────────────────────────
    education = structured.get("education") or []
    if education:
        _add_section_header("EDUCATION")
        for edu in education:
            if isinstance(edu, dict):
                school = edu.get("school") or edu.get("institution") or ""
                degree = edu.get("degree") or ""
                dates = edu.get("dates") or edu.get("graduation_date") or ""
                location = edu.get("location") or ""
                gpa = edu.get("gpa") or ""
                _add_education_line(school, dates, degree, location, gpa)

    # ── Experience (SECOND) ────────────────────────────────────
    experience = structured.get("experience") or []
    if experience:
        _add_section_header("EXPERIENCE")
        for ei, exp in enumerate(experience):
            role = exp.get("role", "")
            company = exp.get("company", "")
            dates = exp.get("dates", "")
            location = exp.get("location", "")
            _add_experience_line(role, company, dates, location)

            bullets = exp.get("bullets") or []
            ei_str = str(ei)
            for bi, bullet in enumerate(bullets):
                bi_str = str(bi)
                text = accepted_edits.get(ei_str, {}).get(bi_str) or bullet
                if text and str(text).strip():
                    _add_compact_plain(str(text).strip())

    # ── Projects ───────────────────────────────────────────────
    projects = structured.get("projects") or []
    if projects:
        _add_section_header("PROJECTS")
        for proj in projects:
            if isinstance(proj, dict):
                name = proj.get("name") or proj.get("title") or ""
                tech = proj.get("tech") or proj.get("technologies") or ""
                dates = proj.get("dates") or ""
                url = proj.get("url") or proj.get("github") or ""
                left_text = name
                if tech:
                    left_text += f" | {tech}"
                if url:
                    left_text += f" | {url}"
                _add_experience_line(left_text, "", dates)
                bullets = proj.get("bullets") or proj.get("description") or []
                if isinstance(bullets, str):
                    bullets = [bullets]
                for bullet in bullets:
                    if bullet and str(bullet).strip():
                        _add_compact_plain(str(bullet).strip())

    # ── Skills ─────────────────────────────────────────────────
    skills = structured.get("skills") or []
    if skills:
        _add_section_header("SKILLS")
        has_levels = any(sk.get("levels") for sk in skills)
        if has_levels:
            all_skills = []
            for sk in skills:
                items = sk.get("items") or []
                levels = sk.get("levels") or []
                for i, item in enumerate(items):
                    level = levels[i] if i < len(levels) else 3
                    all_skills.append((str(item), level))
            if all_skills:
                table = doc.add_table(rows=0, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                tbl = table._tbl
                tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tblBorders.append(border)
                tblPr.append(tblBorders)
                mid = (len(all_skills) + 1) // 2
                for i in range(mid):
                    row = table.add_row()
                    skill1, level1 = all_skills[i]
                    cell1 = row.cells[0]
                    cell1.text = ""
                    p1 = cell1.paragraphs[0]
                    run1 = p1.add_run(f"{skill1}  {_render_dots(level1)}")
                    _set_font(run1, size=9.5)
                    if i + mid < len(all_skills):
                        skill2, level2 = all_skills[i + mid]
                        cell2 = row.cells[1]
                        cell2.text = ""
                        p2 = cell2.paragraphs[0]
                        run2 = p2.add_run(f"{skill2}  {_render_dots(level2)}")
                        _set_font(run2, size=9.5)
        else:
            for sk in skills:
                cat = sk.get("category", "")
                items = sk.get("items") or []
                if not items:
                    continue
                p = doc.add_paragraph()
                if cat:
                    run_cat = p.add_run(f"{cat}: ")
                    _set_font(run_cat, size=9.5, bold=True)
                run_items = p.add_run(", ".join(str(i) for i in items))
                _set_font(run_items, size=9.5)
                _set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    # ── Extra Sections ─────────────────────────────────────────
    for sec in (structured.get("extra_sections") or []):
        title = sec.get("title", "")
        lines = sec.get("lines") or []
        if not lines or not title:
            continue
        _add_section_header(title.upper())
        for line in lines:
            if line and str(line).strip():
                _add_compact_plain(str(line).strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_dots(level: int, max_level: int = 5) -> str:
    try:
        level = int(level)
    except (ValueError, TypeError):
        level = 3
    level = max(1, min(level, max_level))
    return "●" * level + "○" * (max_level - level)


def _docx_attachment_filename(stem: str, fallback: str = "resume") -> str:
    safe = re.sub(r"[^\w.\-]+", "_", (stem or "").strip()).strip("._")[:80] or fallback
    return safe if safe.lower().endswith(".docx") else f"{safe}.docx"
