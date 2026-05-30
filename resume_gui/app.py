"""
Resume Generator GUI — Starlette backend
Run locally:
  .venv/Scripts/python.exe resume_gui/app.py   (Windows)
  python resume_gui/app.py                      (macOS/Linux)

Deploy on Railway — LLM env (pick one primary):
  Grok-first (recommended to avoid Gemini free-tier caps):
    XAI_API_KEY=…
    LLM_PROVIDER=grok
    (optional) GROK_MODEL=grok-4-1-fast-non-reasoning
  Gemini-first:
    GOOGLE_API_KEY=…
    LLM_PROVIDER=gemini   # also use if XAI_API_KEY is set but you want Gemini primary
  Shared: LIBRARY_ROOT, ALLOWED_ORIGINS, …
  Railway auto-detects the Procfile and runs: uvicorn resume_gui.app:app --host 0.0.0.0 --port $PORT
"""

import asyncio
import base64
import io
from functools import partial
import json
import logging
import os
import re
import sys
import threading
from uuid import uuid4
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pdfplumber

# ── Logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(name)s  |  %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("resume_gui")

sys.path.insert(0, str(Path(__file__).parent.parent / "linkedin_agent"))

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / "linkedin_agent" / ".env")

from starlette.applications import Starlette
from starlette.middleware import Middleware
from starlette.middleware.cors import CORSMiddleware
from starlette.requests import Request
from starlette.responses import HTMLResponse, JSONResponse, FileResponse, Response
from starlette.routing import Route
from sse_starlette.sse import EventSourceResponse

import uvicorn

from ats_service import ats_check, structured_ratings_from_ats
from resume_library import (
    list_resumes,
    stream_latex_resume,
    extract_jd_from_url,
    get_resume_tex,
    parse_resume_tex,
    splice_bullets_into_tex,
    recompile_resume_from_tex,
    ai_rewrite_bullet,
    ai_generate_skills,
    doctor_check_resume,
    primary_gemini_flash_model,
    primary_llm_model_for_resume_workloads,
    grok_preferred_for_throughput,
    run_tailor_research_job_context,
    coach_suggestions_llm,
    coach_suggestions_llm_stream,
    _sse_friendly_error,
    _get_resume_tex_for_user,
    _rate_resume,
    _optional_gemini_client,
    _compute_diff,
    _explain_changes,
    _extract_body,
    _sanitize_change_rationales,
)
try:
    from resume_gui.renderers.latex_renderer import (
        JinjaLatexRenderer,
        ResumeDocModel,
        ExperienceItem,
        EducationItem,
        ProjectItem,
        normalize_skill_items,
    )
except ImportError:
    from renderers.latex_renderer import (  # type: ignore
        JinjaLatexRenderer,
        ResumeDocModel,
        ExperienceItem,
        EducationItem,
        ProjectItem,
        normalize_skill_items,
    )
try:
    from resume_gui.profile_parser import parse_profile_text
except ImportError:
    from profile_parser import parse_profile_text  # type: ignore
try:
    from resume_gui.resume_extraction import (
        DEFAULT_SECTION_ORDER,
        ExtractionManifest,
        faithful_extract_prompt,
        infer_section_order_from_profile,
        inject_section_line_breaks,
        inventory_to_dict,
        log_extraction_debug,
        manifest_to_dict,
        pop_manifest_from_llm_raw,
        profile_section_inventory,
        ProfileSectionInventory,
        tailor_doc_prompt,
        validate_extraction_against_inventory,
        validate_manifest_against_doc,
        filter_education_grounded_in_source,
        dedupe_education_rows,
    )
except ImportError:
    from resume_extraction import (  # type: ignore
        DEFAULT_SECTION_ORDER,
        ExtractionManifest,
        faithful_extract_prompt,
        infer_section_order_from_profile,
        inject_section_line_breaks,
        inventory_to_dict,
        log_extraction_debug,
        manifest_to_dict,
        pop_manifest_from_llm_raw,
        profile_section_inventory,
        ProfileSectionInventory,
        tailor_doc_prompt,
        validate_extraction_against_inventory,
        validate_manifest_against_doc,
        filter_education_grounded_in_source,
        dedupe_education_rows,
    )

# Storage helper — works whether run as `uvicorn resume_gui.app:app` (Railway) or
# `python resume_gui/app.py` (local dev).
try:
    from resume_gui.storage import upload_pdf, upload_tex, download_tex, download_pdf, save_version, list_versions, load_version, download_json, storage_status
except ImportError:
    from storage import upload_pdf, upload_tex, download_tex, download_pdf, save_version, list_versions, load_version, download_json, storage_status  # type: ignore

# ── Config (env-var driven for Railway) ──────────────────────────────────────
LIBRARY_ROOT    = os.environ.get("LIBRARY_ROOT", str(Path(__file__).parent.parent / "resumes"))
HTML_FILE       = Path(__file__).parent / "index.html"
PORT            = int(os.environ.get("PORT", 8765))
USE_JINJA_LATEX_RENDERER = os.environ.get("USE_JINJA_LATEX_RENDERER", "true").strip().lower() in {"1", "true", "yes", "on"}
ENABLE_JD_URL_EXTRACT = os.environ.get("ENABLE_JD_URL_EXTRACT", "false").strip().lower() in {"1", "true", "yes", "on"}
USE_SUPABASE_TEMPLATE_BODY = os.environ.get("USE_SUPABASE_TEMPLATE_BODY", "false").strip().lower() in {"1", "true", "yes", "on"}


def _parse_entry_header(header: str) -> Tuple[str, str, str, str]:
    parts = [p.strip() for p in (header or "").split("|") if p.strip()]
    if not parts:
        return "", "", "", ""
    if len(parts) == 1:
        return "", parts[0], "", ""
    if len(parts) == 2:
        return parts[0], parts[1], "", ""
    if len(parts) == 3:
        return parts[0], parts[1], "", parts[2]
    return parts[0], parts[1], parts[2], " | ".join(parts[3:])


_DEGREE_LINE_RE = re.compile(
    r"\b(?:"
    r"B\.?\s*S\.?|B\.?\s*A\.?|Bachelor|M\.?\s*S\.?|M\.?\s*A\.?|MBA|Master'?s?|Ph\.?\s*D\.?|Doctorate|"
    r"Associate|Diploma|PG\s+Diploma|Undergraduate|Postgraduate)\b",
    re.I,
)
_SCHOOL_HINT_RE = re.compile(r"\b(?:university|college|institute|school|academy)\b", re.I)
_EDU_DATE_HINT_RE = re.compile(
    r"\b(?:19|20)\d{2}\b|(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\b",
    re.I,
)
_EXPERIENCE_SECTION_LABELS = frozenset({
    "experience",
    "work experience",
    "professional experience",
    "employment",
    "work history",
})
_SKILL_SECTION_LABELS = frozenset({"skills", "technical skills", "core competencies"})
_LOCATION_BULLET_RE = re.compile(
    r"^[A-Za-z .'-]+,\s*[A-Za-z .'-]+$",
)
_EDU_SECTION_NOISE = frozenset({"projects", "project", "education", "coursework", "relevant coursework"})
_ACCOMPLISHMENT_STARTERS = frozenset({
    "developed", "supported", "led", "built", "implemented", "automated", "collaborated",
    "assisted", "performed", "ensured", "redesigned", "constructed", "consolidated",
    "managed", "conducted", "applied", "designed", "streamlined", "owned", "contributed",
    "created", "analyzed", "analysed", "optimized", "optimised", "reduced", "increased",
})


def _is_experience_section_label(text: str) -> bool:
    low = re.sub(r"[^a-z ]", "", (text or "").lower()).strip()
    return low in _EXPERIENCE_SECTION_LABELS


def _looks_like_accomplishment_bullet(text: str) -> bool:
    t = _clean_model_text(text)
    if len(t.split()) < 8:
        return False
    first = t.split()[0].lower().rstrip(",")
    return first in _ACCOMPLISHMENT_STARTERS


def _normalize_phone_value(phone: str) -> str:
    p = _clean_model_text(phone)
    if p and p[0].isdigit() and ")" in p and "(" not in p:
        p = "(" + p
    return p


def _normalize_structured_experience(doc: ResumeDocModel) -> None:
    """Fix table-style extracts that put the word 'Experience' in the company slot."""
    for exp in doc.experience or []:
        company = _clean_model_text(exp.company or "")
        role = _clean_model_text(exp.role or "")
        loc = _clean_model_text(exp.location or "")
        if _is_experience_section_label(company):
            if "|" in role:
                parts = [p.strip() for p in role.split("|") if p.strip()]
                if len(parts) >= 2:
                    exp.role = _clean_model_text(parts[0])
                    exp.company = _clean_model_text(parts[1])
                    if len(parts) >= 3:
                        exp.location = _clean_model_text(parts[2].strip(" |"))
                else:
                    exp.company = ""
            else:
                exp.company = ""
        if _is_experience_section_label(exp.role):
            exp.role = ""
        if role and not exp.role and not _is_experience_section_label(company):
            if "|" in role:
                parts = [p.strip() for p in role.split("|") if p.strip()]
                if len(parts) >= 2 and not _is_experience_section_label(parts[0]):
                    exp.role = _clean_model_text(parts[0])
                    exp.company = _clean_model_text(parts[1]) or exp.company
                    if len(parts) >= 3:
                        exp.location = _clean_model_text(parts[2].strip(" |")) or loc
            elif not exp.role:
                exp.role = role
        exp.role = _clean_model_text(exp.role or "").strip(" |")
        exp.company = _clean_model_text(exp.company or "")
        if _is_experience_section_label(exp.company):
            exp.company = ""


# Broader degree detection than _DEGREE_LINE_RE — catches Indian / GCSE / IB
# certifications and any line with CGPA/GPA/% which is almost always a degree
# line in education context. Used by _split_collapsed_education_entries.
_DEGREE_BROAD_RE = re.compile(
    r"\b(?:"
    r"B\.?\s*Tech|M\.?\s*Tech|B\.?\s*E\b|M\.?\s*E\b|B\.?\s*Sc|M\.?\s*Sc|"
    r"B\.?\s*S\b|M\.?\s*S\b|B\.?\s*A\b|M\.?\s*A\b|MBA|Ph\.?\s*D|Doctorate|"
    r"Higher\s+Secondary|HSC|SSC|ICSE|CBSE|IGCSE|GED|A[\s-]?Level|"
    r"Certificate|Bachelor|Master|Associate|Undergraduate|Postgraduate|"
    r"Diploma"
    r")\b",
    re.I,
)


def _looks_like_education_degree_line(text: str) -> bool:
    """Heuristic: does this line describe a degree / qualification rather than an institution?"""
    t = _clean_model_text(text or "")
    if not t:
        return False
    if _DEGREE_BROAD_RE.search(t):
        return True
    if re.search(r"\b(?:CGPA|GPA)\b", t, re.I):
        return True
    if re.search(r"\d+(?:\.\d+)?\s*%", t):
        return True
    return False


def _looks_like_education_institution_line(text: str) -> bool:
    """Heuristic: short line naming a school / college / university (no degree tokens)."""
    t = _clean_model_text(text or "")
    if not t:
        return False
    if not _SCHOOL_HINT_RE.search(t):
        return False
    # If it's actually a degree line that happens to mention "school" / "college"
    # (rare, e.g. "Doctor of College Education"), prefer degree.
    if _looks_like_education_degree_line(t):
        return False
    # Real institution names are short — 1-8 words. Bullets that talk ABOUT a
    # school in prose are usually much longer.
    if len(t.split()) > 9:
        return False
    return True


def _split_collapsed_education_entries(doc: ResumeDocModel) -> None:
    """LLM extractor sometimes collapses multiple education entries into ONE entry
    by stuffing the other institutions + their degrees into `bullets[]`. Detect
    that pattern and split them back into proper EducationItem entries.

    Walking the first entry's `bullets`, every institution-like line starts a new
    EducationItem; degree-like lines attach to the current entry's `degree`.
    Finally, if the first entry's `dates` field actually holds a degree-line
    (another common LLM mis-fill), we move it down to the next entry that needs
    a degree.
    """
    if not doc.education:
        return
    EduCtor = type(doc.education[0])
    new_education: List[Any] = []
    for edu in doc.education:
        bullets = [_clean_model_text(b) for b in (edu.bullets or []) if _clean_model_text(b)]
        edu.bullets = []
        kept = [edu]
        current = edu
        for b in bullets:
            if _looks_like_education_institution_line(b):
                current = EduCtor(institution=b, degree="", dates="", location="", bullets=[])
                kept.append(current)
                continue
            if _looks_like_education_degree_line(b) and not current.degree:
                current.degree = b
                continue
            current.bullets.append(b)

        # If the first entry's `dates` actually holds a degree-line AND we
        # successfully split into multiple entries, the LLM almost certainly
        # mis-placed a later entry's degree into the first's dates. Move it.
        # Degree-line markers (Higher Secondary, CGPA, %, B.Tech, ICSE, etc.)
        # take priority over an incidental year in parentheses — real degree
        # lines often include "(2023)" as the graduation year.
        if (
            len(kept) > 1
            and edu.dates
            and _looks_like_education_degree_line(edu.dates)
        ):
            misplaced = edu.dates
            edu.dates = ""
            for k in kept[1:]:
                if not k.degree:
                    k.degree = misplaced
                    break
            else:
                # No taker — restore so we don't silently drop content
                edu.dates = misplaced

        new_education.extend(kept)
    doc.education = new_education


def _normalize_structured_education(doc: ResumeDocModel) -> None:
    """Fix degree/dates/location swapped when PDF tables split one school across lines."""
    for edu in doc.education or []:
        dates = _clean_model_text(edu.dates or "")
        degree = _clean_model_text(edu.degree or "")
        if dates and _DEGREE_LINE_RE.search(dates) and not _EDU_DATE_HINT_RE.search(dates):
            if not degree:
                edu.degree = dates
                edu.dates = ""
            elif _EDU_DATE_HINT_RE.search(degree) and not _EDU_DATE_HINT_RE.search(dates):
                edu.dates, edu.degree = degree, dates
        loc_field = _clean_model_text(edu.location or "")
        if loc_field and _EDU_DATE_HINT_RE.search(loc_field) and not _EDU_DATE_HINT_RE.search(edu.dates or ""):
            edu.dates = loc_field
            edu.location = ""
        bullets = list(edu.bullets or [])
        if bullets and not (edu.location or "").strip():
            kept: list[str] = []
            for b in bullets:
                bt = _clean_model_text(b)
                low_b = re.sub(r"[^a-z ]", "", bt.lower()).strip()
                if low_b in _EDU_SECTION_NOISE:
                    continue
                if (
                    _EDU_DATE_HINT_RE.search(bt)
                    and not _DEGREE_LINE_RE.search(bt)
                    and not edu.dates
                ):
                    edu.dates = bt
                    continue
                if (
                    not edu.location
                    and _LOCATION_BULLET_RE.match(bt)
                    and not _DEGREE_LINE_RE.search(bt)
                    and not _SCHOOL_HINT_RE.search(bt)
                ):
                    edu.location = bt
                else:
                    kept.append(bt)
            edu.bullets = kept


def _normalize_structured_skills(doc: ResumeDocModel) -> None:
    cleaned: list[tuple[str, list[str]]] = []
    for cat, items in doc.skills or []:
        label = _clean_model_text(cat or "")
        low = re.sub(r"[^a-z ]", "", label.lower()).strip()
        if low in _SKILL_SECTION_LABELS:
            label = "Skills"
        clean_items = [
            _clean_model_text(it)
            for it in (items or [])
            if _clean_model_text(it)
            and re.sub(r"[^a-z ]", "", _clean_model_text(it).lower()).strip() not in _SKILL_SECTION_LABELS
            and not _looks_like_accomplishment_bullet(it)
        ]
        if label or clean_items:
            cleaned.append((label or "Skills", clean_items))
    doc.skills = cleaned


def _normalize_structured_contact(doc: ResumeDocModel) -> None:
    doc.phone = _normalize_phone_value(doc.phone or "")


def _parse_school_date_line(line: str) -> Tuple[str, str, str]:
    """Parse ``School | May 2026`` or ``School | City | May 2026`` into institution, dates, location."""
    raw = _clean_model_text(line)
    if not raw or "|" not in raw:
        return "", "", ""
    parts = [p.strip() for p in raw.split("|") if p.strip()]
    if len(parts) < 2 or not _EDU_DATE_HINT_RE.search(parts[-1]):
        return "", "", ""
    if len(parts) >= 3:
        return " | ".join(parts[:-2]), parts[-2], parts[-1]
    return parts[0], parts[1], ""


def _looks_like_school_date_line(line: str) -> bool:
    inst, dates, _ = _parse_school_date_line(line)
    return bool(inst and dates)


_COMBINED_EDU_DATE_RE = re.compile(
    r"((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}"
    r"\s*[–—\-]\s*(?:Present|Now|Current|(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|"
    r"May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}|\d{4}))",
    re.I,
)


def _education_item_from_combined_line(line: str) -> Optional[EducationItem]:
    """One-line education rows: ``Degree … dates | location institution`` (common in PDF extracts)."""
    raw = _clean_model_text(line)
    if "|" not in raw:
        return None
    left, right = [p.strip() for p in raw.split("|", 1)]
    if not left or not right:
        return None
    if not (_DEGREE_LINE_RE.search(left) or _SCHOOL_HINT_RE.search(right)):
        return None
    dm = _COMBINED_EDU_DATE_RE.search(left)
    dates = dm.group(1).strip() if dm else ""
    degree = _clean_model_text(left.replace(dm.group(1), "")) if dm else left
    institution = right
    location = ""
    um = _SCHOOL_HINT_RE.search(right)
    if um and um.start() > 0:
        location = right[: um.start()].strip().rstrip(",")
        institution = right[um.start() :].strip()
    if not (degree or institution):
        return None
    return EducationItem(institution=institution, degree=degree, dates=dates, location=location)


def _looks_like_degree_line(line: str) -> bool:
    t = _clean_model_text(line)
    if not t or _looks_like_school_date_line(t):
        return False
    if _SCHOOL_HINT_RE.search(t) and not _DEGREE_LINE_RE.search(t):
        return False
    if _DEGREE_LINE_RE.search(t):
        return True
    # "MS Clinical Mental Health Counseling" — degree keywords without word boundaries on "MS"
    if re.match(r"^(?:M\.?\s*S\.?|M\.?\s*A\.?|B\.?\s*S\.?|B\.?\s*A\.?|PG)\b", t, re.I):
        return True
    return False


def _education_item_from_dict(edu: dict) -> Optional[EducationItem]:
    inst = _clean_model_text(edu.get("institution") or "")
    deg = _clean_model_text(edu.get("degree") or "")
    dat = _clean_model_text(edu.get("dates") or "")
    loc = _clean_model_text(edu.get("location") or "")
    bullets = [
        _clean_model_text(str(d))
        for d in (edu.get("details") or [])
        if _clean_model_text(str(d))
    ]
    if not inst and deg:
        combined = _education_item_from_combined_line(deg)
        if combined:
            return EducationItem(
                institution=combined.institution or inst,
                degree=combined.degree or deg,
                dates=combined.dates or dat,
                location=combined.location or loc,
                bullets=bullets or combined.bullets,
            )
    if not (inst or deg or dat or loc or bullets):
        return None
    return EducationItem(
        institution=inst,
        degree=deg,
        dates=dat,
        location=loc,
        bullets=bullets,
    )


def _education_items_from_flat_lines(lines: list[str]) -> list[EducationItem]:
    """Group flat education lines (degree, school|date, detail bullets) into structured entries."""
    items: list[EducationItem] = []
    degree = ""
    institution = ""
    dates = ""
    location = ""
    bullets: list[str] = []

    def flush() -> None:
        nonlocal degree, institution, dates, location, bullets
        if not (degree or institution or bullets):
            return
        if institution and not (degree or dates or location or bullets):
            return
        items.append(
            EducationItem(
                institution=institution,
                degree=degree,
                dates=dates,
                location=location,
                bullets=list(bullets),
            )
        )
        degree = ""
        institution = ""
        dates = ""
        location = ""
        bullets = []

    for raw in lines:
        line = _clean_model_text(raw)
        if not line or _is_structural_noise_line(line):
            continue

        combined = _education_item_from_combined_line(line)
        if combined:
            flush()
            items.append(combined)
            continue

        if _looks_like_degree_line(line):
            flush()
            degree = line
            continue

        if (
            _EDU_DATE_HINT_RE.search(line)
            and not _DEGREE_LINE_RE.search(line)
            and not _SCHOOL_HINT_RE.search(line)
            and (degree or institution)
            and not dates
        ):
            dates = line
            continue

        inst, dat, loc = _parse_school_date_line(line)
        if inst:
            if degree or not items:
                if institution:
                    flush()
                    degree = ""
                institution = inst
                dates = dat
                location = loc
                continue
            institution = inst
            dates = dat
            location = loc
            continue

        if not degree and not institution:
            row = _education_item_from_csv_line(line)
            if row.institution and not row.degree and not row.dates:
                institution = row.institution
                continue
            if row.institution or row.degree:
                flush()
                items.append(row)
            continue

        bullets.append(line)

    flush()
    return items


def _collect_education_flat_lines(
    education_lines: list[str],
    extra_sections: list[tuple[str, list[str]]],
) -> list[str]:
    out: list[str] = []
    for ln in education_lines or []:
        ce = _clean_model_text(ln)
        if ce:
            out.append(ce)
    for name, vals in extra_sections or []:
        if (name or "").strip().lower() != "education":
            continue
        for v in vals:
            ce = _clean_model_text(str(v))
            if ce:
                out.append(ce)
    return out


def _education_item_from_csv_line(line: str) -> EducationItem:
    """Best-effort split of a single-line education row (comma-separated) into hierarchy fields."""
    raw = _clean_model_text(line)
    if not raw:
        return EducationItem(institution="")
    parts = [p.strip() for p in raw.split(",") if p.strip()]
    if len(parts) >= 4:
        institution = parts[0]
        location = parts[-1]
        dates = parts[-2]
        degree = ", ".join(parts[1:-2])
        return EducationItem(
            institution=institution,
            degree=degree,
            dates=dates,
            location=location,
        )
    if len(parts) == 3:
        return EducationItem(institution=parts[0], degree=parts[1], dates=parts[2])
    if len(parts) == 2:
        if _SCHOOL_HINT_RE.search(parts[0]) and not _DEGREE_LINE_RE.search(parts[1]):
            return EducationItem(institution=raw)
        return EducationItem(institution=parts[0], degree=parts[1])
    return EducationItem(institution=raw)


def _project_items_from_llm_projects(raw_projects: Any) -> list[ProjectItem]:
    out: list[ProjectItem] = []
    for proj in raw_projects or []:
        if not isinstance(proj, dict):
            continue
        name_p = _clean_model_text(proj.get("name") or "")
        p_buls = [_clean_model_text(b) for b in (proj.get("bullets") or []) if _clean_model_text(b)]
        if name_p and p_buls:
            out.append(ProjectItem(name=name_p, bullets=p_buls))
        elif name_p:
            out.append(ProjectItem(name=name_p, bullets=[]))
        elif p_buls:
            out.append(ProjectItem(name="", bullets=p_buls))
    return out


def _project_items_from_prefixed_bullets(lines: list[str]) -> list[ProjectItem]:
    """Group ``ProjectName: detail`` lines so the PDF shows one title and multiple bullets."""
    grouped: list[ProjectItem] = []
    i = 0
    n = len(lines)
    while i < n:
        line = _clean_model_text(lines[i])
        if not line:
            i += 1
            continue
        if ":" in line:
            head, rest = line.split(":", 1)
            h, r = head.strip(), rest.strip()
            if h and len(r) > 8:
                bullets = [r]
                i += 1
                while i < n:
                    nxt = _clean_model_text(lines[i])
                    if not nxt:
                        i += 1
                        continue
                    if ":" in nxt:
                        nh, nr = nxt.split(":", 1)
                        nh, nr = nh.strip(), nr.strip()
                        if nh.lower() == h.lower() and len(nr) > 8:
                            bullets.append(nr)
                            i += 1
                            continue
                        break
                    bullets.append(nxt)
                    i += 1
                grouped.append(ProjectItem(name=h, bullets=bullets))
                continue
        grouped.append(ProjectItem(name="", bullets=[line]))
        i += 1
    return grouped


def _clean_model_text(value: str) -> str:
    t = (value or "").strip()
    if not t:
        return ""
    # Remove Markdown emphasis markers that leak from parser output.
    t = t.replace("**", "")
    # Remove common LaTeX control sequences while preserving plain words.
    t = re.sub(r"\\(?:begin|end)\{[^}]*\}", " ", t)
    t = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", " ", t)
    t = t.replace("{", " ").replace("}", " ")
    t = re.sub(r"\s+", " ", t).strip()
    return t


_STRUCTURAL_NOISE_PATTERNS = (
    r"^leftmargin\s*=",
    r"^label\s*=",
    r"^textbackslash$",
    r"^begin\s+itemize$",
    r"^end\s+itemize$",
    r"^item$",
    r"^\\item$",
    r"^\\textbackslash$",
)

_SECTION_HEADING_LINES = {
    "candidate",
    "summary",
    "technical skills",
    "skills",
    "experience",
    "work experience",
    "professional experience",
    "education",
    "github",
    "linkedin",
}


def _is_structural_noise_line(value: str) -> bool:
    t = _clean_model_text(value)
    if not t:
        return True
    low = t.lower().strip(" :-")
    if low in _SECTION_HEADING_LINES:
        return True
    for pat in _STRUCTURAL_NOISE_PATTERNS:
        if re.match(pat, low):
            return True
    return False


def _resume_doc_from_parsed(parsed: dict) -> ResumeDocModel:
    contact = parsed.get("contact") or {}
    sections = parsed.get("sections") or []

    doc = ResumeDocModel(
        full_name=str(contact.get("name") or "Candidate").strip() or "Candidate",
        headline="",
        location=str(contact.get("location") or "").strip(),
        email=str(contact.get("email") or "").strip(),
        phone=str(contact.get("phone") or "").strip(),
        linkedin=str(contact.get("linkedinUrl") or contact.get("linkedin") or "").strip(),
        github=str(contact.get("githubUrl") or contact.get("github") or "").strip(),
        summary="",
        skills=[],
        experience=[],
        extra_sections=[],
    )

    for sec in sections:
        sec_name = str(sec.get("name") or "").strip().lower()
        entries = sec.get("entries") or []

        if "summary" in sec_name or "profile" in sec_name:
            for ent in entries:
                bullets = ent.get("bullets") or []
                if bullets:
                    cleaned_lines = []
                    for b in bullets:
                        line = _clean_model_text(str(b.get("text") or ""))
                        if not line or _is_structural_noise_line(line):
                            continue
                        cleaned_lines.append(line)
                    doc.summary = " ".join(cleaned_lines)
                    break
            continue

        if "skill" in sec_name:
            for ent in entries:
                for b in ent.get("bullets") or []:
                    line = _clean_model_text(str(b.get("text") or ""))
                    if not line or _is_structural_noise_line(line):
                        continue
                    if ":" in line:
                        label, rest = line.split(":", 1)
                        if _is_structural_noise_line(label):
                            label = "Skills"
                        items = [x.strip() for x in rest.split(",") if x.strip()]
                        items = [x for x in items if not _is_structural_noise_line(x)]
                        clean_label = _clean_model_text(label)
                        if not clean_label:
                            clean_label = "Skills"
                        normalized = normalize_skill_items(items)
                        if normalized:
                            doc.skills.append((clean_label, normalized))
                    else:
                        normalized = normalize_skill_items([line])
                        if normalized:
                            doc.skills.append(("Skills", normalized))
            continue

        if "experience" in sec_name or "work" in sec_name:
            for ent in entries:
                role, company, location, dates = _parse_entry_header(str(ent.get("header") or ""))
                bullets = [
                    _clean_model_text(str(b.get("text") or ""))
                    for b in (ent.get("bullets") or [])
                    if _clean_model_text(str(b.get("text") or ""))
                ]
                bullets = [b for b in bullets if not _is_structural_noise_line(b)]
                if not company and not role and not bullets:
                    continue
                doc.experience.append(
                    ExperienceItem(
                        company=company or "Experience",
                        role=role,
                        location=location,
                        dates=dates,
                        bullets=bullets,
                    )
                )
            continue

        if "education" in sec_name:
            for ent in entries:
                flat: list[str] = []
                hdr = _clean_model_text(str(ent.get("header") or "")).strip()
                if hdr and not _is_structural_noise_line(hdr):
                    for part in re.split(r"\s*·\s*", hdr):
                        p = part.strip()
                        if p:
                            flat.append(p)
                for b in ent.get("bullets") or []:
                    txt = _clean_model_text(str(b.get("text") or ""))
                    if txt and not _is_structural_noise_line(txt):
                        flat.append(txt)
                doc.education.extend(_education_items_from_flat_lines(flat))
            continue

        if "project" in sec_name:
            for ent in entries:
                hdr = _clean_model_text(str(ent.get("header") or "")).strip()
                bullets = [
                    _clean_model_text(str(b.get("text") or ""))
                    for b in (ent.get("bullets") or [])
                    if _clean_model_text(str(b.get("text") or ""))
                ]
                bullets = [b for b in bullets if not _is_structural_noise_line(b)]
                name = ""
                if hdr and not _is_structural_noise_line(hdr):
                    name = re.split(r"\s*·\s*", hdr)[0].strip()
                if name or bullets:
                    doc.projects.append(ProjectItem(name=name, bullets=bullets))
            continue

        # Dynamic catch-all sections for anything beyond summary/skills/experience/education/projects.
        sec_label = _clean_model_text(str(sec.get("name") or "")).strip()
        if sec_label:
            extra_lines: list[str] = []
            for ent in entries:
                hdr = _clean_model_text(str(ent.get("header") or "")).strip()
                if hdr and not _is_structural_noise_line(hdr):
                    extra_lines.append(hdr)
                for b in ent.get("bullets") or []:
                    txt = _clean_model_text(str(b.get("text") or ""))
                    if txt and not _is_structural_noise_line(txt):
                        extra_lines.append(txt)
            if extra_lines:
                doc.extra_sections.append((sec_label, extra_lines[:30]))

    # Consolidate duplicated skill labels/items from noisy model output.
    merged_skills: list[tuple[str, list[str]]] = []
    skill_index: dict[str, int] = {}
    for label, items in doc.skills:
        clean_label = _clean_model_text(label) or "Skills"
        key = clean_label.lower()
        if key not in skill_index:
            skill_index[key] = len(merged_skills)
            merged_skills.append((clean_label, []))
        idx = skill_index[key]
        existing = merged_skills[idx][1]
        seen = {x.lower() for x in existing}
        for item in items:
            item_clean = _clean_model_text(item)
            if not item_clean or _is_structural_noise_line(item_clean):
                continue
            lk = item_clean.lower()
            if lk not in seen:
                existing.append(item_clean)
                seen.add(lk)
    doc.skills = [(label, items) for label, items in merged_skills if items]

    cleaned_exp: list[ExperienceItem] = []
    for exp in doc.experience:
        company = _clean_model_text(exp.company)
        role = _clean_model_text(exp.role)
        bullets = [b for b in exp.bullets if b and not _is_structural_noise_line(b)]
        if not bullets and not role and company.lower() in {"", "experience"}:
            continue
        cleaned_exp.append(
            ExperienceItem(
                company=company or "Experience",
                role=role,
                location=_clean_model_text(exp.location),
                dates=_clean_model_text(exp.dates),
                bullets=bullets,
            )
        )
    doc.experience = cleaned_exp

    if doc.projects or doc.education:
        filtered_extra: list[tuple[str, list[str]]] = []
        for title, lines in doc.extra_sections:
            lname = (title or "").strip().lower()
            if doc.projects and lname in ("projects", "project"):
                continue
            if doc.education and lname == "education":
                continue
            filtered_extra.append((title, lines))
        doc.extra_sections = filtered_extra

    return doc


def _resume_doc_with_updates(doc: "ResumeDocModel", **overrides: Any) -> "ResumeDocModel":
    """Rebuild ``ResumeDocModel`` preserving sections not mentioned in ``overrides``."""
    return ResumeDocModel(
        full_name=overrides.get("full_name", doc.full_name),
        headline=overrides.get("headline", doc.headline),
        location=overrides.get("location", doc.location),
        email=overrides.get("email", doc.email),
        phone=overrides.get("phone", doc.phone),
        linkedin=overrides.get("linkedin", doc.linkedin),
        github=overrides.get("github", doc.github),
        summary=overrides.get("summary", doc.summary),
        skills=overrides.get("skills", doc.skills),
        experience=overrides.get("experience", doc.experience),
        education=overrides.get("education", doc.education),
        projects=overrides.get("projects", doc.projects),
        extra_sections=overrides.get("extra_sections", doc.extra_sections),
    )


def _accepted_suggestion_section_bucket(section: Optional[str]) -> str:
    """Where to apply a suggestion when verbatim ``original`` is not found in the structured doc.

    Coach quotes ``original`` from raw profile text, while ``_resume_doc_from_profile_text`` may
    normalize via LLM — substring matches often fail. A blanket ``doc.summary = suggested`` fallback
    then pasted *skills* text into the summary (user-visible bug).
    """
    s = (section or "").strip().lower()
    if not s:
        return "other"
    if "summary" in s or "profile" in s or "objective" in s:
        return "summary"
    if "skill" in s:
        return "skills"
    if "experience" in s or "employment" in s or ("work" in s and "project" not in s):
        return "experience"
    if "project" in s:
        return "projects"
    if "education" in s or "academic" in s:
        return "education"
    return "other"


def _append_extra_section_line(doc: ResumeDocModel, section_title: str, line: str) -> None:
    lt = (line or "").strip()
    if not lt:
        return
    key = section_title.strip().lower()
    if key in ("projects", "project"):
        if doc.projects:
            doc.projects[-1].bullets.append(lt)
        else:
            doc.projects.append(ProjectItem(name="", bullets=[lt]))
        return
    if key == "education":
        if doc.education:
            doc.education[-1].bullets.append(lt)
        else:
            doc.education.extend(_education_items_from_flat_lines([lt]))
        return
    for i, (name, lines) in enumerate(doc.extra_sections):
        nl = name.strip().lower()
        if nl == key or key in nl or nl in key:
            seq = list(lines)
            if lt not in seq:
                seq.append(lt)
            doc.extra_sections[i] = (name, seq)
            return
    doc.extra_sections.append((section_title, [lt]))


def _skills_fallback_replace_or_append(doc: ResumeDocModel, original: str, suggested: str) -> None:
    """When ``original`` is not inside any skill line, still apply a skills-section approval."""
    sug = (suggested or "").strip()
    if not sug:
        return
    orig = (original or "").strip()
    # Prefer replacing a line that shares a long prefix with what the coach quoted.
    if orig and len(orig) >= 24:
        prefix = orig[:48].lower()
        for idx, (label, items) in enumerate(doc.skills):
            next_items: list[str] = []
            changed = False
            for it in items:
                if prefix in (it or "").lower():
                    next_items.append(sug)
                    changed = True
                else:
                    next_items.append(it)
            if changed:
                doc.skills[idx] = (label, [x for x in next_items if x])
                return
    # Parse "Category: a, b, c" like the profile parser.
    if ":" in sug:
        label, rest = sug.split(":", 1)
        label = _clean_model_text(label)
        rest = rest.strip()
        if not label:
            label = "Skills"
        items = [x.strip() for x in rest.replace("·", ",").split(",") if x.strip()]
        items = [x for x in items if not _is_structural_noise_line(x)]
        normalized = normalize_skill_items(items)
        if normalized:
            doc.skills.append((label, normalized))
            return
    doc.skills.append(("Skills", [sug]))


def _normalize_resume_line_for_suggestion(s: str) -> str:
    """Align with ``web/lib/suggestionResumeMatch.ts`` ``normalizeResumeLineForSuggestion``."""
    t = (s or "").strip().lower()
    for a, b in (
        ("\u2018", "'"),
        ("\u2019", "'"),
        ("\u201c", "'"),
        ("\u201d", "'"),
        ("\u2032", "'"),
        ("\u2033", "'"),
    ):
        t = t.replace(a, b)
    t = re.sub(
        r"^[\s\u2022\u00b7\u2023\u2024\u2043\u2219\-\u2013\u2014*‧·.]+",
        "",
        t,
    )
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _resume_line_matches_suggestion_original(line: str, original: str) -> bool:
    """Same contract as ``resumeLineMatchesSuggestionOriginal`` in ``suggestionResumeMatch.ts``."""
    no = _normalize_resume_line_for_suggestion(original)
    nl = _normalize_resume_line_for_suggestion(line)
    if not no or not nl:
        return False
    if nl == no:
        return True

    raw_line = line.strip().lower().replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", "'").replace("\u201d", "'")
    raw_orig = original.strip().lower().replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", "'").replace("\u201d", "'")
    if raw_line == raw_orig:
        return True

    min_contains = 8
    if len(no) >= min_contains and len(nl) >= min_contains:
        if no in nl or nl in no:
            return True

    pref_len = 55
    pref_a = nl[:pref_len]
    pref_b = no[:pref_len]
    if len(pref_a) >= 12 and len(pref_b) >= 12:
        if pref_a.startswith(pref_b) or pref_b.startswith(pref_a):
            return True

    if len(no) >= 14 and len(nl) >= 14:
        shorter = nl if len(nl) <= len(no) else no
        longer = no if len(nl) <= len(no) else nl
        window = longer[: len(shorter) + 8]
        if shorter in window:
            return True

    return False


def _apply_line_suggestion(line: str, original: str, suggested: str) -> Optional[str]:
    """If ``original`` matches this line (verbatim or fuzzy), return the new line text; else None."""
    if not original:
        return None
    if original in line:
        if suggested:
            return line.replace(original, suggested).strip()
        return ""
    if _resume_line_matches_suggestion_original(line, original):
        if suggested:
            return suggested.strip()
        return ""
    return None


def _experience_fallback_replace_by_prefix(doc: ResumeDocModel, original: str, suggested: str) -> bool:
    """Last resort when fuzzy match failed: replace the bullet whose text contains a long prefix of ``original``."""
    orig = (original or "").strip()
    sug = (suggested or "").strip()
    if not orig or not sug or len(orig) < 24:
        return False
    prefix = orig[:48].lower()
    for exp in doc.experience:
        next_bullets: list[str] = []
        changed = False
        for b in exp.bullets:
            if prefix in (b or "").lower():
                next_bullets.append(sug)
                changed = True
            else:
                next_bullets.append(b)
        if changed:
            exp.bullets = [x for x in next_bullets if x]
            return True
    return False


def _apply_accepted_edits_to_doc(doc: ResumeDocModel, accepted_suggestions: Optional[list]) -> None:
    if not isinstance(accepted_suggestions, list):
        return

    for item in accepted_suggestions:
        if not isinstance(item, dict):
            continue
        original = str(item.get("original") or "").strip()
        suggested = str(item.get("suggested") or "").strip()
        if not original:
            continue

        replaced = False

        if doc.summary and original in doc.summary:
            doc.summary = doc.summary.replace(original, suggested).strip()
            replaced = True

        for idx, (label, items) in enumerate(doc.skills):
            next_items = []
            changed = False
            for it in items:
                updated = _apply_line_suggestion(it, original, suggested)
                if updated is None:
                    next_items.append(it)
                else:
                    changed = True
                    replaced = True
                    if updated:
                        next_items.append(updated)
            if changed:
                doc.skills[idx] = (label, [x for x in next_items if x])

        for exp in doc.experience:
            next_bullets = []
            changed = False
            for b in exp.bullets:
                updated = _apply_line_suggestion(b, original, suggested)
                if updated is None:
                    next_bullets.append(b)
                else:
                    changed = True
                    replaced = True
                    if updated:
                        next_bullets.append(updated)
            if changed:
                exp.bullets = [x for x in next_bullets if x]

        for proj in doc.projects:
            next_bullets = []
            changed = False
            for b in proj.bullets:
                updated = _apply_line_suggestion(b, original, suggested)
                if updated is None:
                    next_bullets.append(b)
                else:
                    changed = True
                    replaced = True
                    if updated:
                        next_bullets.append(updated)
            if changed:
                proj.bullets = [x for x in next_bullets if x]

        for edu in doc.education:
            for attr in ("institution", "degree", "dates", "location"):
                cur = getattr(edu, attr, "") or ""
                if cur and original in cur:
                    setattr(edu, attr, cur.replace(original, suggested).strip())
                    replaced = True
            next_bullets = []
            changed = False
            for b in edu.bullets:
                updated = _apply_line_suggestion(b, original, suggested)
                if updated is None:
                    next_bullets.append(b)
                else:
                    changed = True
                    replaced = True
                    if updated:
                        next_bullets.append(updated)
            if changed:
                edu.bullets = [x for x in next_bullets if x]

        # Coach quotes raw profile lines; structured doc may normalize wording — do not
        # route unrelated sections into ``summary``.
        if not replaced and suggested:
            bucket = _accepted_suggestion_section_bucket(str(item.get("section") or ""))
            if bucket == "summary":
                doc.summary = suggested.strip()
            elif bucket == "skills":
                _skills_fallback_replace_or_append(doc, original, suggested)
            elif bucket == "projects":
                _append_extra_section_line(doc, "Projects", suggested)
            elif bucket == "education":
                _append_extra_section_line(doc, "Education", suggested)
            elif bucket == "experience":
                if not _experience_fallback_replace_by_prefix(doc, original, suggested):
                    logger.warning(
                        "accepted_suggestion could not be matched to structured experience "
                        "(no verbatim or fuzzy line match, prefix fallback failed); skipped id=%s section=%r",
                        str(item.get("id") or "")[:32],
                        (item.get("section") or "")[:80],
                    )
            else:
                logger.warning(
                    "accepted_suggestion could not be matched in structured doc; skipped id=%s section=%r",
                    str(item.get("id") or "")[:32],
                    (item.get("section") or "")[:80],
                )


def _template_name_for_reference(reference_folder: Optional[str]) -> str:
    rf = (reference_folder or "").strip().lower()
    if rf == "harshibar_template1":
        return "harshibar_resume.tex.j2"
    if rf == "adobe_fullstack":
        return "classic_resume.tex.j2"
    if rf == "maltacv_modern":
        return "classic_resume.tex.j2"
    return "classic_resume.tex.j2"


def _supabase_template_is_jinja(tex_body: Optional[str]) -> bool:
    t = (tex_body or "").strip()
    if not t:
        return False
    return "{{ doc." in t or "{%" in t


def _create_structured_output_folder(base_folder: Optional[str], reference_folder: Optional[str], role: str, company: str) -> Tuple[str, str]:
    ref_name = (reference_folder or "").strip()
    base_name = (base_folder or "").strip()

    # Never chain from prior generated artifacts like *_structured_<id>.
    # Prefer canonical template key from reference_folder.
    source_name = ref_name
    if not source_name and base_name and "_structured_" not in base_name:
        source_name = base_name
    if not source_name:
        rc = f"{role}_{company}".strip("_") or "structured"
        source_name = re.sub(r"[^A-Za-z0-9_.-]+", "", rc) or "structured"

    new_folder = f"{source_name}_structured_{uuid4().hex[:8]}"
    dst = Path(LIBRARY_ROOT) / new_folder
    dst.mkdir(parents=True, exist_ok=True)

    tex_files = [p for p in dst.iterdir() if p.suffix == ".tex"]
    if not tex_files:
        fallback = dst / "resume.tex"
        fallback.write_text("", encoding="utf-8")
        tex_files = [fallback]
    return new_folder, str(tex_files[0])


def _load_tex_from_candidate(folder: str, user_id: Optional[str]) -> Optional[str]:
    name = (folder or "").strip()
    if not name:
        return None
    tex = _get_resume_tex_for_user(name, user_id)
    if tex:
        return tex
    return get_resume_tex(name)


def _resolve_structured_source_folder(base_folder: Optional[str], reference_folder: Optional[str], user_id: Optional[str]) -> Tuple[str, str]:
    _ = user_id  # Structured template resolution is Supabase-driven.
    candidates: List[str] = []
    ref_name = (reference_folder or "").strip()
    base_name = (base_folder or "").strip()
    # Builder sends `reference_folder` as the explicit LaTeX style choice (e.g. Harshibar).
    # `base_folder` is whichever library row backed the last compile (often Adobe_FullStack) — it must
    # not override the user's selected reference style when loading `resume_templates` rows.
    if ref_name:
        candidates.append(ref_name)
    if base_name and "_structured_" not in base_name and base_name not in candidates:
        candidates.append(base_name)
    for fallback in ("Harshibar_Template1", "Adobe_FullStack", "MaltaCV_Modern"):
        if fallback not in candidates:
            candidates.append(fallback)

    for c in candidates:
        # Canonical source: Supabase `resume_templates`.
        tex_from_template = _load_template_tex_from_supabase(c)
        if tex_from_template:
            return c, tex_from_template

    raise RuntimeError(
        "Could not load template TeX from Supabase table `resume_templates` for folders: "
        + ", ".join(candidates)
        + ". Ensure rows exist with {reference_folder, tex_body, active=true}."
    )


def _build_resume_doc_from_llm_raw(
    raw: dict,
    *,
    role: str = "",
    company: str = "",
) -> ResumeDocModel:
    """Map LLM JSON to ResumeDocModel with consistent field cleaning."""
    full_name = str(raw.get("full_name") or "Candidate").strip() or "Candidate"

    skills: list[tuple[str, list[str]]] = []
    for sk in (raw.get("skills") or []):
        if not isinstance(sk, dict):
            continue
        cat = _clean_model_text(str(sk.get("category") or "Skills"))
        items = normalize_skill_items(str(i) for i in (sk.get("items") or []))
        if items:
            skills.append((cat or "Skills", items))

    experience: list[ExperienceItem] = []
    for exp in (raw.get("experience") or []):
        if not isinstance(exp, dict):
            continue
        bullets = [
            _clean_model_text(str(b))
            for b in (exp.get("bullets") or [])
            if _clean_model_text(str(b))
        ]
        company_name = _clean_model_text(str(exp.get("company") or "")) or "Company"
        experience.append(
            ExperienceItem(
                company=company_name,
                role=_clean_model_text(str(exp.get("role") or "")),
                dates=_clean_model_text(str(exp.get("dates") or "")),
                location=_clean_model_text(str(exp.get("location") or "")),
                bullets=bullets,
            )
        )

    education: list[EducationItem] = []
    for edu in (raw.get("education") or []):
        if isinstance(edu, dict):
            row = _education_item_from_dict(edu)
            if row:
                education.append(row)

    projects = _project_items_from_llm_projects(raw.get("projects"))

    return ResumeDocModel(
        full_name=full_name,
        headline=_clean_model_text(str(raw.get("headline") or role or "")),
        location=_clean_model_text(str(raw.get("location") or "")),
        email=_clean_model_text(str(raw.get("email") or "")),
        phone=_clean_model_text(str(raw.get("phone") or "")),
        linkedin=_clean_model_text(str(raw.get("linkedin") or "")),
        github=_clean_model_text(str(raw.get("github") or "")),
        summary=_clean_model_text(str(raw.get("summary") or "")),
        skills=skills,
        experience=experience,
        education=education,
        projects=projects,
        extra_sections=[],
    )


def _count_approved_suggestions(accepted_suggestions: Optional[List]) -> int:
    if not isinstance(accepted_suggestions, list):
        return 0
    n = 0
    for item in accepted_suggestions:
        if not isinstance(item, dict):
            continue
        orig = str(item.get("original") or "").strip()
        sugg = str(item.get("suggested") or "").strip()
        if orig or sugg:
            n += 1
    return n


def _rationales_from_accepted_suggestions(accepted_suggestions: Optional[List]) -> List[Dict]:
    """Human-readable change cards from user-approved suggestion rows only."""
    out: List[Dict] = []
    if not isinstance(accepted_suggestions, list):
        return out
    for item in accepted_suggestions:
        if not isinstance(item, dict):
            continue
        orig = str(item.get("original") or "").strip()
        sugg = str(item.get("suggested") or "").strip()
        why = str(item.get("reason") or "").strip() or "Approved from your suggestion cards."
        if orig and sugg and orig != sugg:
            out.append({"type": "rewrote", "previous": orig, "text": sugg, "why": why})
        elif sugg and not orig:
            out.append({"type": "added", "text": sugg, "why": why})
        elif orig and not sugg:
            out.append({"type": "removed", "text": orig, "why": why or "Removed per approved suggestion."})
    return _sanitize_change_rationales(out)


def _structured_tailor_diff_and_rationales(
    *,
    baseline_tex: Optional[str],
    new_tex: str,
    jd: str,
    model: str,
    gemini_client,
    accepted_suggestions: Optional[List],
) -> Tuple[List[Dict], int, int, List[Dict]]:
    """Line diff + change rationales for the Jinja structured PDF path (never JD gap chips)."""
    diff_lines: List[Dict] = []
    adds = removes = 0
    rationales: List[Dict] = []

    old_body = _extract_body(baseline_tex or "") if baseline_tex else ""
    new_body = _extract_body(new_tex or "") if new_tex else ""
    if old_body.strip() and new_body.strip() and old_body.strip() != new_body.strip():
        diff_lines, adds, removes = _compute_diff(old_body, new_body)
        try:
            explained = _explain_changes(gemini_client, model, old_body, new_body, (jd or "")[:1500])
            if explained:
                rationales = _sanitize_change_rationales(explained)
        except Exception as exc:
            logger.warning("structured path: change explanations failed: %s", exc)

    if not rationales:
        rationales = _rationales_from_accepted_suggestions(accepted_suggestions)

    return diff_lines, adds, removes, rationales


def _doc_extraction_counts(doc: "ResumeDocModel") -> dict:
    """Compact section counts for log lines (grep ``EXTRACT_DEBUG``)."""
    return {
        "experience_entries": len(doc.experience or []),
        "experience_bullets": sum(len(e.bullets or []) for e in (doc.experience or [])),
        "education_entries": len(doc.education or []),
        "skill_groups": len(doc.skills or []),
        "skill_items": sum(len(items) for _, items in (doc.skills or [])),
        "projects": len(doc.projects or []),
        "extra_sections": len(doc.extra_sections or []),
        "summary_chars": len((doc.summary or "").strip()),
        "section_order": list(getattr(doc, "section_order", None) or DEFAULT_SECTION_ORDER),
    }


def _log_structured_doc(stage: str, doc: "ResumeDocModel", **extra: Any) -> None:
    payload: dict = {"counts": _doc_extraction_counts(doc), "structured_doc": _resume_doc_to_dict(doc)}
    if extra:
        payload["meta"] = extra
    log_extraction_debug(stage, payload)


def _llm_tailor_doc_for_jd(
    doc: ResumeDocModel,
    jd: str,
    role: str,
    company: str,
) -> ResumeDocModel:
    """Second pass: JD-tailor summary, experience bullets, and project bullets (structure unchanged)."""
    jd_snippet = (jd or "")[:3000].strip()
    if not jd_snippet:
        return doc
    if not doc.experience and not doc.projects and not (doc.summary or "").strip():
        return doc
    before_counts = _doc_extraction_counts(doc)
    prompt = tailor_doc_prompt(doc, jd_snippet, role, company)
    try:
        raw = _llm_json_call(prompt)
        if not raw or not isinstance(raw, dict):
            log_extraction_debug(
                "jd_tailor_skipped",
                {"reason": "empty_llm_response", "before": before_counts},
            )
            return doc
        log_extraction_debug(
            "jd_tailor_llm_response",
            {
                "keys": list(raw.keys()),
                "summary_preview": (_clean_model_text(str(raw.get("summary") or "")) or "")[:400],
                "experience_patches": len(raw.get("experience") or [])
                if isinstance(raw.get("experience"), list)
                else 0,
            },
        )
        summary = _clean_model_text(str(raw.get("summary") or ""))
        if summary:
            doc.summary = summary
        exp_in = raw.get("experience")
        if isinstance(exp_in, list) and len(exp_in) == len(doc.experience):
            for i, patch in enumerate(exp_in):
                if not isinstance(patch, dict):
                    continue
                bullets = [
                    _clean_model_text(str(b))
                    for b in (patch.get("bullets") or [])
                    if _clean_model_text(str(b))
                ]
                if bullets:
                    doc.experience[i].bullets = bullets
        proj_in = raw.get("projects")
        if isinstance(proj_in, list) and doc.projects and len(proj_in) == len(doc.projects):
            for i, patch in enumerate(proj_in):
                if not isinstance(patch, dict):
                    continue
                bullets = [
                    _clean_model_text(str(b))
                    for b in (patch.get("bullets") or [])
                    if _clean_model_text(str(b))
                ]
                if bullets:
                    doc.projects[i].bullets = bullets
        _log_structured_doc(
            "jd_tailor_after",
            doc,
            before=before_counts,
            after=_doc_extraction_counts(doc),
        )
        return doc
    except Exception as exc:
        logger.warning("LLM tailor-doc second pass failed: %s", exc)
        return doc


def _structured_doc_for_generate(
    candidate_profile: Optional[str],
    jd: str,
    role: str,
    company: str,
    *,
    use_conservative_tailor: bool,
    base_tex: Optional[str] = None,
    pre_parsed: Optional[dict] = None,
) -> ResumeDocModel:
    """Faithful extract first, optional JD tailor second, then profile backfill + section order."""
    profile_norm = inject_section_line_breaks((candidate_profile or "")[:8000])
    inv = profile_section_inventory(profile_norm)
    section_order = infer_section_order_from_profile(profile_norm)

    log_extraction_debug(
        "generate_pipeline_start",
        {
            "inventory": inventory_to_dict(inv),
            "section_order_inferred": section_order,
            "profile_chars": len(profile_norm),
            "profile_preview": profile_norm[:2000],
            "use_conservative_tailor": use_conservative_tailor,
            "has_base_tex": bool(base_tex),
            "jd_chars": len((jd or "").strip()),
            "role": role,
            "company": company,
            "has_pre_parsed": bool(pre_parsed),
        },
    )

    doc: Optional[ResumeDocModel] = None
    manifest: Optional[ExtractionManifest] = None
    extract_path = "unknown"

    # Use the structured JSON produced at upload time to skip redundant LLM re-extraction.
    # Only applies on the non-conservative path (conservative paths rely on base_tex/profile text
    # for intentional minimal-diff behaviour).
    if pre_parsed and isinstance(pre_parsed, dict) and not use_conservative_tailor and not base_tex:
        try:
            doc = _build_resume_doc_from_llm_raw(pre_parsed)
            if doc and (doc.experience or doc.education or doc.summary):
                extract_path = "pre_parsed_upload_json"
                if jd.strip():
                    extract_path = "pre_parsed_upload_json_then_jd_tailor"
                    doc = _llm_tailor_doc_for_jd(doc, jd, role, company)
            else:
                logger.warning("pre_parsed upload JSON yielded empty doc — falling through to LLM extract")
                doc = None
        except Exception as exc:
            logger.warning("pre_parsed upload JSON parse failed (%s) — falling through to LLM extract", exc)
            doc = None

    if doc is not None:
        pass  # pre_parsed path succeeded — skip remaining extraction branches
    elif use_conservative_tailor and base_tex:
        extract_path = "conservative_tex"
        parsed = parse_resume_tex(base_tex)
        doc = _resume_doc_from_parsed(parsed)
    elif use_conservative_tailor and profile_norm:
        extract_path = "conservative_profile_regex"
        doc = _resume_doc_from_profile_text(profile_norm, role, company)
    elif profile_norm:
        extract_path = "faithful_llm"
        doc, manifest = _llm_extract_with_manifest(profile_norm)
        if doc is None:
            extract_path = "faithful_llm_failed_regex_fallback"
            logger.warning("Faithful LLM extract failed — falling back to regex profile parse")
            doc = _resume_doc_from_profile_text(profile_norm, role, company)
        elif jd.strip() and not use_conservative_tailor:
            extract_path = "faithful_llm_then_jd_tailor"
            doc = _llm_tailor_doc_for_jd(doc, jd, role, company)

    if doc is None and profile_norm:
        extract_path = "regex_fallback"
        doc = _resume_doc_from_profile_text(profile_norm, role, company)
    if doc is None and base_tex:
        extract_path = "tex_fallback"
        parsed = parse_resume_tex(base_tex)
        doc = _resume_doc_from_parsed(parsed)
    if doc is None:
        extract_path = "empty_profile_default"
        doc = _resume_doc_from_profile_text(profile_norm or "", role, company)

    _finalize_structured_doc(doc, profile_norm, inv, manifest, role, company)
    doc.section_order = section_order or list(DEFAULT_SECTION_ORDER)
    _log_structured_doc(
        "generate_pipeline_final_json",
        doc,
        extract_path=extract_path,
        manifest=manifest_to_dict(manifest),
    )
    return doc


def _llm_extract_with_manifest(text: str) -> tuple[Optional[ResumeDocModel], Optional[ExtractionManifest]]:
    """Faithful structured extract; returns (doc, manifest) for validation/backfill."""
    profile_norm = inject_section_line_breaks((text or "")[:8000])
    profile_snippet = profile_norm[:6000].strip()
    if not profile_snippet:
        return None, None

    inv = profile_section_inventory(profile_norm)
    log_extraction_debug(
        "faithful_extract_input",
        {
            "inventory": inventory_to_dict(inv),
            "profile_snippet_chars": len(profile_snippet),
            "profile_snippet": profile_snippet,
        },
    )
    prompt = faithful_extract_prompt(profile_snippet, inv)

    try:
        # Structured extraction is where accuracy matters most — the parsed
        # ResumeDocModel feeds both the analysis and the rendered preview. A
        # bad split here (3 schools collapsed into 1) cascades through both.
        # Route this single call to the reasoning tier (grok-4-fast-reasoning
        # / gemini-2.5-pro). Real cost: ~8s of extra latency on Analyze, ~1.6×
        # the cost of that single call. Empirically improves date extraction
        # (0/3 → 2/3 runs) and reduces miss rates on the entry-split task.
        # Other calls (analyze prompt, gap-fix, suggestions) stay on the fast
        # non-reasoning model.
        if grok_preferred_for_throughput():
            extract_model = _grok_reasoning_model()
        else:
            extract_model = _gemini_reasoning_model()
        raw = _llm_json_call(prompt, model_override=extract_model)
        if not raw or not isinstance(raw, dict):
            log_extraction_debug("faithful_extract_failed", {"reason": "empty_or_non_dict_llm_response"})
            return None, None
        body, manifest = pop_manifest_from_llm_raw(raw)
        log_extraction_debug(
            "faithful_extract_llm_raw",
            {
                "manifest": manifest_to_dict(manifest),
                "body_keys": list(body.keys()) if isinstance(body, dict) else [],
                "llm_top_level_keys": list(raw.keys()),
            },
        )
        doc = _build_resume_doc_from_llm_raw(body)
        if doc is not None:
            _log_structured_doc("faithful_extract_parsed_doc", doc, manifest=manifest_to_dict(manifest))
        return doc, manifest
    except Exception as exc:
        logger.warning("LLM faithful extract failed: %s", exc)
        log_extraction_debug("faithful_extract_failed", {"reason": str(exc)})
        return None, None


def _finalize_structured_doc(
    doc: ResumeDocModel,
    profile_norm: str,
    inv: ProfileSectionInventory,
    manifest: Optional[ExtractionManifest],
    role: str,
    company: str,
) -> None:
    """Manifest + inventory checks, then regex profile backfill for any dropped sections."""
    before = _doc_extraction_counts(doc)
    warnings: list[str] = []
    for w in validate_extraction_against_inventory(doc, inv):
        logger.warning("Structured extract completeness: %s", w)
        warnings.append(w)
    for w in validate_manifest_against_doc(doc, manifest):
        logger.warning("Structured extract manifest mismatch: %s", w)
        warnings.append(w)
    for w in filter_education_grounded_in_source(doc, profile_norm):
        logger.warning("Structured extract education grounding: %s", w)
        warnings.append(w)
    _preserve_structured_sections_from_profile(doc, profile_norm, inv, role, company)
    # Split collapsed-into-one education entries BEFORE the per-entry normalizer
    # so the date/degree swap heuristics in _normalize_structured_education see
    # each institution as its own entry.
    _split_collapsed_education_entries(doc)
    _normalize_structured_experience(doc)
    _normalize_structured_education(doc)
    _normalize_structured_skills(doc)
    _normalize_structured_contact(doc)
    after = _doc_extraction_counts(doc)
    log_extraction_debug(
        "finalize_structured_doc",
        {
            "before_counts": before,
            "after_counts": after,
            "validation_warnings": warnings,
            "manifest": manifest_to_dict(manifest),
        },
    )


# ── Vision-based PDF extraction ────────────────────────────────────────────
# Multi-column / two-frame PDF layouts (institution+location on the same row,
# role+date on the same row) break text-extractors because pdfplumber's
# y-bucketing reads columns separately. The result: locations leak from
# Education into Experience, dates float to the wrong place, and the LLM that
# parses the text then either hallucinates or collapses entries.
#
# A vision-capable model reading the PDF as an image sidesteps all of this —
# it sees the same visual layout a recruiter does. Empirically (3 résumés,
# multiple runs) Grok-4 vision returns a perfectly-structured JSON in ~10–17s
# with every numeral, every location, and every wrapped bullet intact.
#
# We try this path first for PDF uploads and fall back to the text-based
# reasoning extract on any failure or for non-PDF uploads.

_VISION_EXTRACT_PROMPT = """You are reading a one-page (or rarely multi-page) résumé image. Produce a COMPLETE, FAITHFUL JSON of every field on the page — do not skip anything, do not paraphrase. Schema:

{
  "full_name": "",
  "headline": "",
  "phone": "",
  "email": "",
  "linkedin": "",
  "github": "",
  "portfolio": "",
  "location": "",
  "summary": "",
  "education": [
    {"institution":"", "degree":"", "location":"", "dates":"", "gpa_or_grade":"", "notes":""}
  ],
  "experience": [
    {"company":"", "role":"", "location":"", "dates":"", "bullets":["..."]}
  ],
  "projects": [
    {"name":"", "tech":"", "dates":"", "bullets":["..."]}
  ],
  "skills_groups": [
    {"category":"", "items":["..."]}
  ],
  "certifications": [{"name":"", "issuer":"", "date":""}],
  "awards": ["..."],
  "publications": ["..."],
  "languages": ["..."],
  "activities_or_leadership": [{"title":"", "organization":"", "dates":"", "bullets":["..."]}],
  "extra_sections": [{"section_name":"", "items":["..."]}]
}

CRITICAL RULES:
- Every institution = its own education entry. Never merge two schools into one.
- Every distinct role/job = its own experience entry.
- Bullets must be FULL multi-line text exactly as visible; do not truncate, do not paraphrase.
- Preserve every numeral, percent, year, CGPA, count (e.g. "5 refinement cycles", "up to 3 retries", "9,000+ records", "50,000+ HEIs", "~70%", "CGPA: 9.266").
- If a field is absent in the résumé, set it to "" or [].
- Skills MUST be grouped exactly as the résumé groups them (e.g. "Languages and Frameworks", "ML / AI"). Do not invent groups.
- Put anything that doesn't fit a named schema field into extra_sections.
Return ONLY the JSON object, no markdown fences, no commentary."""


def _render_pdf_pages_to_b64_pngs(
    pdf_bytes: bytes, *, dpi: int = 200, max_pages: int = 2
) -> List[str]:
    """Render the first ``max_pages`` of a PDF to base64-encoded PNGs."""
    try:
        import fitz  # type: ignore  # PyMuPDF
    except Exception as exc:
        logger.warning("vision-extract: PyMuPDF unavailable (%s)", exc)
        return []
    out: List[str] = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        logger.warning("vision-extract: PDF open failed (%s)", exc)
        return []
    try:
        for i in range(min(len(doc), max_pages)):
            pix = doc[i].get_pixmap(dpi=dpi)
            png_bytes = pix.tobytes("png")
            out.append(base64.standard_b64encode(png_bytes).decode())
    finally:
        doc.close()
    return out


def _vision_raw_to_resume_doc(raw: dict) -> ResumeDocModel:
    """Map the vision-extract JSON schema onto ResumeDocModel.

    The vision schema separates ``degree`` from ``gpa_or_grade`` and emits a
    standalone ``activities_or_leadership`` array; we merge the two grade
    fields into ``EducationItem.degree`` and append activities (plus any
    certifications / awards / publications / languages) into the doc's
    ``extra_sections`` so nothing the candidate wrote is dropped.
    """
    EduCtor = EducationItem
    ExpCtor = ExperienceItem
    ProjCtor = ProjectItem

    # ── Education ─────────────────────────────────────────────────────────
    # The vision LLM is run-to-run inconsistent about whether the grade lives
    # in `degree` (as "B.Tech ... CGPA: 9.266 (Ongoing)"), in `gpa_or_grade`
    # ("CGPA: 9.266 (Ongoing)"), or both. Combine only when the degree text
    # doesn't already include the grade — otherwise we end up with
    # "B.Tech ... CGPA: 9.266 (Ongoing) — CGPA: 9.266 (Ongoing)".
    education: List[EducationItem] = []
    for e in (raw.get("education") or []):
        if not isinstance(e, dict):
            continue
        degree = str(e.get("degree") or "").strip()
        grade = str(e.get("gpa_or_grade") or "").strip()
        # Strip a possible trailing date / status from the grade for the
        # containment check ("CGPA: 9.266 (Ongoing)" vs "CGPA: 9.266").
        grade_core = re.sub(r"\s*\(\s*[^)]+\s*\)\s*$", "", grade).strip()
        if degree and grade:
            if grade_core and grade_core.lower() in degree.lower():
                combined = degree
            elif grade.lower() in degree.lower():
                combined = degree
            else:
                combined = f"{degree} — {grade}"
        else:
            combined = degree or grade
        notes = str(e.get("notes") or "").strip()
        education.append(EduCtor(
            institution=str(e.get("institution") or "").strip(),
            degree=combined,
            dates=str(e.get("dates") or "").strip(),
            location=str(e.get("location") or "").strip(),
            bullets=[notes] if notes else [],
        ))

    # ── Experience ────────────────────────────────────────────────────────
    # Vision is run-to-run inconsistent about whether the location lives in
    # `location` or gets embedded into `company` as "Company · Location" /
    # "Company, Location". Detect the embedded form and strip the location
    # out so we don't end up with "Co · Tardeo, Mumbai | Tardeo, Mumbai".
    experience: List[ExperienceItem] = []
    for w in (raw.get("experience") or []):
        if not isinstance(w, dict):
            continue
        company = str(w.get("company") or "").strip()
        location = str(w.get("location") or "").strip()
        if company and location:
            # Try common embedded patterns: " · ", " - ", " | ", ", "
            for sep in (" · ", " - ", " | "):
                if sep in company:
                    left, right = company.split(sep, 1)
                    if location.lower() == right.strip().lower() or location.lower() in right.strip().lower():
                        company = left.strip()
                        break
            # Trailing ", Location" pattern — only strip when the comma-tail
            # is a clear location match (avoid eating legit "Company, Inc.").
            if company.lower().endswith(f", {location.lower()}"):
                company = company[: -(len(location) + 2)].strip()
        bullets = [str(b).strip() for b in (w.get("bullets") or []) if str(b).strip()]
        experience.append(ExpCtor(
            company=company,
            role=str(w.get("role") or "").strip(),
            dates=str(w.get("dates") or "").strip(),
            location=location,
            bullets=bullets,
        ))

    # ── Projects ──────────────────────────────────────────────────────────
    projects: List[ProjectItem] = []
    for p in (raw.get("projects") or []):
        if not isinstance(p, dict):
            continue
        name = str(p.get("name") or "").strip()
        tech = str(p.get("tech") or "").strip()
        bullets = [str(b).strip() for b in (p.get("bullets") or []) if str(b).strip()]
        # Prepend tech stack as the first bullet so the stack list is preserved
        # in display (Harshibar's project template doesn't have a dedicated
        # tech field yet). Skip when the tech string would duplicate the name.
        if tech and tech not in bullets:
            bullets = [tech] + bullets
        projects.append(ProjCtor(name=name, bullets=bullets))

    # ── Skills (preserve the candidate's own grouping) ────────────────────
    # The LLM is run-to-run inconsistent on the items[] shape: sometimes a
    # proper array (["Python", "Java", ...]), sometimes a single CSV string
    # (["Python, Java, C, ..."]). Split any comma-joined entries so the
    # rendered preview shows them as individual chips.
    skills: List[tuple[str, List[str]]] = []
    for g in (raw.get("skills_groups") or []):
        if not isinstance(g, dict):
            continue
        cat = str(g.get("category") or "Skills").strip() or "Skills"
        raw_items = g.get("items") or []
        items: List[str] = []
        for it in raw_items:
            s = str(it).strip()
            if not s:
                continue
            if "," in s and len(s) > 18:
                # Split on commas while preserving compound names like "Power BI"
                # and parenthesized clarifications.
                parts = [p.strip().strip(",") for p in s.split(",")]
                items.extend(p for p in parts if p)
            else:
                items.append(s)
        if items:
            skills.append((cat, items))

    # ── extra_sections: activities, certs, awards, etc. ──────────────────
    extra_sections: List[tuple[str, List[str]]] = []

    activities = raw.get("activities_or_leadership") or []
    if activities:
        lines: List[str] = []
        for a in activities:
            if not isinstance(a, dict):
                continue
            title = str(a.get("title") or "").strip()
            org = str(a.get("organization") or "").strip()
            dates = str(a.get("dates") or "").strip()
            header = " | ".join(p for p in (title, org, dates) if p)
            if header:
                lines.append(header)
            for b in (a.get("bullets") or []):
                bs = str(b).strip()
                if bs:
                    lines.append(f"• {bs}")
        if lines:
            extra_sections.append(("Activities & Leadership", lines))

    for key, label in (
        ("certifications", "Certifications"),
        ("awards", "Awards"),
        ("publications", "Publications"),
        ("languages", "Languages"),
    ):
        items = raw.get(key) or []
        if not items:
            continue
        lines = []
        for it in items:
            if isinstance(it, dict):
                pieces = [str(it.get(k) or "").strip() for k in ("name", "issuer", "date")]
                line = " | ".join(p for p in pieces if p)
                if line:
                    lines.append(line)
            else:
                s = str(it).strip()
                if s:
                    lines.append(s)
        if lines:
            extra_sections.append((label, lines))

    for s in (raw.get("extra_sections") or []):
        if not isinstance(s, dict):
            continue
        name = str(s.get("section_name") or "").strip() or "Other"
        items = [str(it).strip() for it in (s.get("items") or []) if str(it).strip()]
        if items:
            extra_sections.append((name, items))

    return ResumeDocModel(
        full_name=str(raw.get("full_name") or "Candidate").strip() or "Candidate",
        headline=str(raw.get("headline") or "").strip(),
        location=str(raw.get("location") or "").strip(),
        email=str(raw.get("email") or "").strip(),
        phone=str(raw.get("phone") or "").strip(),
        linkedin=str(raw.get("linkedin") or "").strip(),
        github=str(raw.get("github") or "").strip(),
        summary=str(raw.get("summary") or "").strip(),
        skills=skills,
        experience=experience,
        education=education,
        projects=projects,
        extra_sections=extra_sections,
    )


def _llm_extract_pdf_vision(pdf_bytes: bytes) -> Optional[ResumeDocModel]:
    """Try the vision-PDF structured extract. Returns None on any failure
    (no PDF bytes, no XAI key, PyMuPDF missing, API error, empty result)
    so callers can fall back to the text-based reasoning extract."""
    if not pdf_bytes or len(pdf_bytes) < 100:
        return None
    if os.environ.get("DISABLE_VISION_EXTRACT", "").lower() in ("1", "true", "yes"):
        return None
    if not os.environ.get("XAI_API_KEY"):
        # Vision via Gemini isn't wired here yet; require xAI key for the fast path.
        return None
    images_b64 = _render_pdf_pages_to_b64_pngs(pdf_bytes)
    if not images_b64:
        return None
    try:
        from openai import OpenAI  # type: ignore
        model = (os.environ.get("VISION_EXTRACT_MODEL") or "grok-4").strip()
        xai = OpenAI(api_key=os.environ["XAI_API_KEY"], base_url="https://api.x.ai/v1")
        content: List[dict] = [{"type": "text", "text": _VISION_EXTRACT_PROMPT}]
        for b64 in images_b64:
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{b64}"},
            })
        r = xai.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": content}],
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        text = (r.choices[0].message.content or "").strip()
        text = re.sub(r"^```[a-z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
        raw = json.loads(text)
    except Exception as exc:
        logger.warning("vision-extract LLM call failed (%s) — falling back", exc)
        return None
    if not isinstance(raw, dict):
        logger.warning("vision-extract returned non-dict — falling back")
        return None
    doc = _vision_raw_to_resume_doc(raw)
    # Sanity check — if nothing meaningful was extracted, refuse and let the
    # text path try.
    if not (doc.experience or doc.education or doc.projects or doc.skills):
        logger.warning("vision-extract returned empty doc — falling back")
        return None
    logger.info(
        "vision-extract OK | edu=%d exp=%d proj=%d skills=%d extra=%d",
        len(doc.education), len(doc.experience), len(doc.projects),
        len(doc.skills), len(doc.extra_sections),
    )
    return doc


def _synthesize_text_from_resume_doc(doc: ResumeDocModel) -> str:
    """Produce a clean text representation of a ResumeDocModel for the preview
    + analysis prompt.

    When the vision-PDF extract succeeds, the raw MarkItDown / pdfplumber text
    is no longer the source of truth — the structured doc is. The Analyze
    right-panel preview reads ``extractedText`` line-by-line, and so does the
    analysis prompt. Synthesizing both from the (clean) doc means the user
    sees the layout the parser actually understood, and the LLM grades the
    real content instead of column-extraction artifacts ("fragmented
    locations under Experience", "disorganized header").

    Format mirrors what the frontend renderer expects:
      - Header: NAME on line 1, contact pipe-joined on line 2
      - Section labels in ALL-CAPS (EDUCATION / EXPERIENCE / ...)
      - Entry headers use ``" | "`` separators (Company | Role | Location | Date)
        so ``looksLikeEntryHeader`` picks them up for bold styling
      - Bullets prefixed with ``"• "``
      - Blank line between entries / sections
    """
    out: List[str] = []

    name = (doc.full_name or "").strip()
    if name:
        out.append(name)

    contact_parts: List[str] = []
    for v in (doc.phone, doc.email, doc.linkedin, doc.github):
        s = (v or "").strip()
        if s:
            contact_parts.append(s)
    if contact_parts:
        out.append(" | ".join(contact_parts))

    if (doc.location or "").strip():
        out.append(doc.location.strip())

    summary = (doc.summary or "").strip()
    if summary:
        out.extend(["", "SUMMARY", summary])

    if doc.education:
        out.extend(["", "EDUCATION"])
        # Year-or-status hint embedded as "(YYYY)" / "(Ongoing|Present|Current|Now)"
        # at the end of the degree string. Lift it onto the entry header line
        # so it renders right-aligned (Harshibar style) instead of leaking out
        # as a parenthesized fragment that the frontend treats as a bullet.
        _trailing_date_re = re.compile(
            r"\s*\(\s*((?:19|20)\d{2}|Ongoing|Present|Current|Now|"
            r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+(?:19|20)\d{2})\s*\)\s*$",
            re.IGNORECASE,
        )
        for e in doc.education:
            inst = (e.institution or "").strip()
            loc = (e.location or "").strip()
            explicit_dates = (e.dates or "").strip()
            degree = (e.degree or "").strip()
            cleaned_degree = degree
            year_from_degree = ""
            m = _trailing_date_re.search(degree) if degree else None
            if m:
                year_from_degree = m.group(1).strip()
                cleaned_degree = degree[: m.start()].rstrip()
            dates = explicit_dates or year_from_degree
            header_pieces = [p for p in (inst, loc, dates) if p]
            if header_pieces:
                out.append(" | ".join(header_pieces))
            if cleaned_degree:
                out.append(cleaned_degree)
            for b in (e.bullets or []):
                bs = (b or "").strip()
                if bs:
                    out.append(f"• {bs}")
            out.append("")
        if out and out[-1] == "":
            out.pop()

    if doc.experience:
        out.extend(["", "EXPERIENCE"])
        for w in doc.experience:
            role = (w.role or "").strip()
            company = (w.company or "").strip()
            loc = (w.location or "").strip()
            dates = (w.dates or "").strip()
            header_pieces = [p for p in (role, company, loc, dates) if p]
            if header_pieces:
                out.append(" | ".join(header_pieces))
            for b in (w.bullets or []):
                bs = (b or "").strip()
                if bs:
                    out.append(f"• {bs}")
            out.append("")
        if out and out[-1] == "":
            out.pop()

    if doc.projects:
        out.extend(["", "PROJECTS"])
        for p in doc.projects:
            name_p = (p.name or "").strip()
            bullets = [b.strip() for b in (p.bullets or []) if b and b.strip()]
            # First bullet may be the tech stack (synthesized that way by
            # _vision_raw_to_resume_doc). Promote it onto the project header
            # line for nicer rendering.
            tech = ""
            if bullets and " · " in bullets[0] and not bullets[0].startswith(("Built ", "Designed ", "Engineered ", "Architected ", "Developed ", "Delivered ", "Implemented ")):
                tech = bullets[0]
                bullets = bullets[1:]
            if name_p and tech:
                out.append(f"{name_p} | {tech}")
            elif name_p:
                out.append(name_p)
            for b in bullets:
                out.append(f"• {b}")
            out.append("")
        if out and out[-1] == "":
            out.pop()

    if doc.skills:
        out.extend(["", "SKILLS"])
        for cat, items in doc.skills:
            items_clean = [i.strip() for i in (items or []) if i and i.strip()]
            if not items_clean:
                continue
            label = (cat or "Skills").strip() or "Skills"
            out.append(f"{label}: {', '.join(items_clean)}")

    for title, lines in (doc.extra_sections or []):
        title_clean = (title or "").strip()
        if not title_clean or not lines:
            continue
        out.extend(["", title_clean.upper()])
        for ln in lines:
            s = (ln or "").strip()
            if s:
                out.append(s)

    text = "\n".join(out).strip()
    # Collapse runs of >2 newlines so we don't accidentally create giant gaps.
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _llm_extract(text: str, pdf_bytes: Optional[bytes] = None) -> "Optional[ResumeDocModel]":
    """Extract structured resume data faithfully — no tailoring or bullet rewriting.

    Used by the Analyze path so the structured model mirrors what the user actually wrote.
    Returns None on failure.

    When ``pdf_bytes`` is provided, tries the vision-based PDF extract first
    (cleaner on multi-column layouts) and falls back to the text-based
    reasoning extract on any failure.
    """
    profile_norm = inject_section_line_breaks((text or "")[:8000])
    profile_snippet = profile_norm[:6000].strip()
    if not profile_snippet and not pdf_bytes:
        return None

    # ── Vision-first for PDF uploads ──────────────────────────────────────
    vision_doc = _llm_extract_pdf_vision(pdf_bytes) if pdf_bytes else None
    if vision_doc is not None:
        # Vision output is already clean — skip the manifest/inventory backfill
        # and the entry-splitter post-processors (which were built for the
        # text-extract failure modes). Keep the lightweight normalizers as a
        # safety net.
        _normalize_structured_experience(vision_doc)
        _normalize_structured_education(vision_doc)
        _normalize_structured_skills(vision_doc)
        _normalize_structured_contact(vision_doc)
        vision_doc.section_order = infer_section_order_from_profile(profile_norm or "")
        _log_structured_doc("analyze_extract_final_json_vision", vision_doc)
        return vision_doc

    # ── Fallback: text-based reasoning extract ────────────────────────────
    doc, manifest = _llm_extract_with_manifest(text)
    if doc is None:
        return None
    inv = profile_section_inventory(profile_norm)
    log_extraction_debug(
        "analyze_extract_input",
        {"inventory": inventory_to_dict(inv), "profile_chars": len(profile_norm)},
    )
    _finalize_structured_doc(doc, profile_norm, inv, manifest, "", "")
    doc.section_order = infer_section_order_from_profile(profile_norm)
    _log_structured_doc("analyze_extract_final_json", doc)
    return doc


def _resume_doc_to_dict(doc: "ResumeDocModel") -> dict:
    """Serialize a ResumeDocModel to a plain dict for JSON transport."""
    return {
        "full_name": doc.full_name,
        "headline": doc.headline,
        "location": doc.location,
        "email": doc.email,
        "phone": doc.phone,
        "linkedin": doc.linkedin,
        "github": doc.github,
        "summary": doc.summary,
        "skills": [{"category": cat, "items": items} for cat, items in (doc.skills or [])],
        "experience": [
            {
                "company": e.company,
                "role": e.role,
                "dates": e.dates,
                "location": e.location,
                "bullets": list(e.bullets),
            }
            for e in (doc.experience or [])
        ],
        "education": [
            {
                "institution": e.institution,
                "degree": e.degree,
                "dates": e.dates,
                "location": e.location,
                "bullets": list(e.bullets),
            }
            for e in (doc.education or [])
        ],
        "projects": [
            {"name": p.name, "bullets": list(p.bullets)}
            for p in (doc.projects or [])
        ],
        "extra_sections": [
            {"title": title, "lines": lines}
            for title, lines in (doc.extra_sections or [])
        ],
        "section_order": list(getattr(doc, "section_order", None) or DEFAULT_SECTION_ORDER),
    }


def _resume_doc_from_profile_text(candidate_profile: Optional[str], role: str, company: str) -> ResumeDocModel:
    parsed = parse_profile_text(candidate_profile)
    extra_sec = list(parsed.extra_sections or [])

    full_name = _clean_model_text(parsed.full_name or "Candidate") or "Candidate"
    summary = _clean_model_text(parsed.summary or "")
    skills: list[tuple[str, list[str]]] = []
    for ln in parsed.skills_lines or []:
        clean = _clean_model_text(ln)
        if not clean or _is_structural_noise_line(clean) or _looks_like_accomplishment_bullet(clean):
            continue
        if ":" in clean:
            label, rest = clean.split(":", 1)
            items = [x.strip() for x in rest.split(",") if x.strip()]
            if items:
                skills.append((_clean_model_text(label) or "Skills", normalize_skill_items(items)))
        else:
            skills.append(("Skills", normalize_skill_items([clean])))

    if parsed.experience_entries:
        experience_list = [
            ExperienceItem(
                company=_clean_model_text(ec),
                role=_clean_model_text(er or ""),
                dates=_clean_model_text(ed or ""),
                location=_clean_model_text(el or ""),
                bullets=[_clean_model_text(b) for b in eb if _clean_model_text(b)],
            )
            for ec, er, ed, el, eb in parsed.experience_entries
        ]
    else:
        reg_bullets = [_clean_model_text(b) for b in (parsed.experience_bullets or []) if _clean_model_text(b)]
        experience_list = [
            ExperienceItem(
                company=company or "",
                role=role or "",
                dates="",
                location="",
                bullets=reg_bullets,
            )
        ]

    structured_projects: list[ProjectItem] = []
    if parsed.projects_bullets:
        clean_proj = [_clean_model_text(b) for b in parsed.projects_bullets if _clean_model_text(b)]
        if clean_proj:
            structured_projects = _project_items_from_prefixed_bullets(clean_proj)
            if not structured_projects:
                structured_projects = [ProjectItem(name="", bullets=clean_proj)]
    if not structured_projects:
        for name, vals in extra_sec:
            if (name or "").strip().lower() in ("projects", "project") and vals:
                structured_projects = _project_items_from_prefixed_bullets(
                    [_clean_model_text(v) for v in vals if _clean_model_text(v)]
                )
                break

    edu_flat = _collect_education_flat_lines(parsed.education_lines, extra_sec)
    structured_education = _education_items_from_flat_lines(edu_flat)

    extra_sec_filtered: list[tuple[str, list[str]]] = []
    for name, vals in extra_sec:
        lname = (name or "").strip().lower()
        if structured_projects and lname in ("projects", "project"):
            continue
        if structured_education and lname == "education":
            continue
        cleaned = (name, [_clean_model_text(v) for v in vals if _clean_model_text(v)])
        if cleaned[0]:
            extra_sec_filtered.append(cleaned)

    return ResumeDocModel(
        full_name=full_name,
        headline=_clean_model_text(parsed.headline or role or ""),
        location=_clean_model_text(parsed.location or ""),
        email=_clean_model_text(parsed.email or ""),
        phone=_normalize_phone_value(_clean_model_text(parsed.phone or "")),
        linkedin=_clean_model_text(parsed.linkedin or ""),
        github=_clean_model_text(parsed.github or ""),
        summary=summary,
        skills=skills,
        experience=experience_list,
        education=structured_education,
        projects=structured_projects,
        extra_sections=extra_sec_filtered,
    )


def _preserve_structured_sections_from_profile(
    doc: ResumeDocModel,
    candidate_profile: Optional[str],
    inv: ProfileSectionInventory,
    role: str,
    company: str,
) -> None:
    """When LLM tailoring drops sections, backfill from regex profile parse (upload/PDF text)."""
    profile_norm = inject_section_line_breaks((candidate_profile or "")[:8000])
    if not profile_norm.strip():
        return
    try:
        source = _resume_doc_from_profile_text(profile_norm, role, company)
    except Exception as exc:
        logger.warning("preserve_structured_sections_from_profile failed: %s", exc)
        return

    src_edu = list(source.education or [])
    doc_edu = list(doc.education or [])
    expected_rows = max(1, inv.estimated_education_lines) if inv.expects_education() else 0

    if not doc_edu and src_edu:
        doc.education = src_edu
        logger.info(
            "Structured renderer: restored %s education row(s) from profile parse",
            len(doc.education),
        )
    elif inv.expects_education() and src_edu:
        if len(doc_edu) > len(src_edu) and len(src_edu) <= max(2, expected_rows):
            logger.warning(
                "Structured renderer: replacing %s LLM education row(s) with %s profile-parsed row(s) "
                "(LLM over-extracted vs source EDUCATION section)",
                len(doc_edu),
                len(src_edu),
            )
            doc.education = src_edu
        elif len(doc_edu) < len(src_edu):
            seen = {
                (e.institution.strip().lower(), e.degree.strip().lower())
                for e in doc_edu
            }
            for row in src_edu:
                key = (row.institution.strip().lower(), row.degree.strip().lower())
                if key not in seen and (row.institution or row.degree):
                    doc.education.append(row)
                    seen.add(key)
    doc.education = dedupe_education_rows(list(doc.education or []))

    if not doc.projects and source.projects:
        doc.projects = list(source.projects)

    if not doc.skills and source.skills:
        doc.skills = list(source.skills)
    elif inv.has_skills_header and source.skills:
        src_count = sum(len(items) for _, items in source.skills)
        doc_count = sum(len(items) for _, items in doc.skills)
        if doc_count < max(3, src_count // 2):
            doc.skills = list(source.skills)
            logger.info("Structured renderer: restored skills from profile parse (LLM under-extracted)")

    if inv.expects_experience() and source.experience:
        if not doc.experience:
            doc.experience = list(source.experience)
            logger.info(
                "Structured renderer: restored %s experience entries from profile parse",
                len(doc.experience),
            )
        elif len(doc.experience) < max(1, inv.estimated_job_blocks - 1) and len(source.experience) > len(doc.experience):
            doc.experience = list(source.experience)
            logger.warning(
                "Structured renderer: replaced experience with profile parse "
                "(%s jobs vs LLM %s)",
                len(source.experience),
                len(doc.experience),
            )

    for field_name, src_val, doc_val in (
        ("email", source.email, doc.email),
        ("phone", source.phone, doc.phone),
        ("linkedin", source.linkedin, doc.linkedin),
        ("github", source.github, doc.github),
    ):
        if src_val and not doc_val:
            setattr(doc, field_name, src_val)

    for w in validate_extraction_against_inventory(doc, inv):
        logger.warning("After profile backfill still incomplete: %s", w)

# CORS: localhost dev + production site (merge env extras; never drop defaults)
_CORS_DEFAULT_ORIGINS = [
    "http://localhost:3000",
    "http://localhost:3001",
    "http://localhost:3002",
    "http://localhost:8765",
    "https://www.resunova.io",
    "https://resunova.io",
]
_extra = [o.strip() for o in os.environ.get("ALLOWED_ORIGINS", "").split(",") if o.strip()]
ALLOWED_ORIGINS = list(dict.fromkeys(_CORS_DEFAULT_ORIGINS + _extra))


async def homepage(request: Request):
    return HTMLResponse(HTML_FILE.read_text(encoding="utf-8"))


async def api_health(_request: Request):
    """GET /api/health — liveness probe for Railway / uptime checks."""
    return JSONResponse({"ok": True, "service": "resume_gui"})


async def api_resumes(request: Request):
    user_id = (request.query_params.get("user_id") or "").strip()

    # 1. Try the Supabase `resumes` *table* first — this is the authoritative source
    #    and is properly scoped to the authenticated user.
    if user_id:
        supabase = _supabase_client()
        if supabase:
            try:
                rows = (
                    supabase.table("resumes")
                    .select("folder, company, role, score, pdf_url, created_at, job_description")
                    .eq("user_id", user_id)
                    .order("created_at", desc=True)
                    .execute()
                    .data or []
                )
                if rows:
                    logger.info(f"GET /api/resumes  |  {len(rows)} DB rows  user={user_id}")
                    return JSONResponse(rows)
            except Exception as exc:
                logger.warning(f"api_resumes: DB query failed: {exc}")

    # 2. Fall back to local disk (dev mode only — never exposes other users' data
    #    because local resumes have no user_id concept).
    resumes = list_resumes()
    logger.info(f"GET /api/resumes  |  {len(resumes)} local entries  user={user_id or 'anon'}")
    return JSONResponse(resumes)


async def api_generate_stream(request: Request):
    """SSE endpoint — streams events as the resume is generated."""
    try:
        body = await request.json()
    except json.JSONDecodeError:
        async def err_gen():
            yield {
                "data": json.dumps({
                    "event": "error",
                    "msg": "Invalid JSON body. Send strict JSON with double-quoted property names.",
                })
            }
        return EventSourceResponse(err_gen())

    if not isinstance(body, dict):
        async def err_gen():
            yield {
                "data": json.dumps({
                    "event": "error",
                    "msg": "Invalid request body. Expected a JSON object.",
                })
            }
        return EventSourceResponse(err_gen())

    company           = (body.get("company") or "").strip()
    role              = (body.get("role") or "").strip()
    jd                = (body.get("job_description") or "").strip()
    model = primary_llm_model_for_resume_workloads((body.get("model") or "").strip() or None)
    base_folder       = (body.get("base_folder") or "").strip() or None
    reference_folder  = (body.get("reference_folder") or "").strip() or None
    candidate_profile = (body.get("candidate_profile") or "").strip() or None
    user_id           = (body.get("user_id") or "").strip() or "local"
    layout_compile    = bool(body.get("layout_compile"))
    accepted_suggestions = body.get("accepted_suggestions")
    suggest_research_digest = (body.get("suggest_research_digest") or "").strip() or None
    post_suggestion_coach_run = bool(body.get("post_suggestion_coach_run"))
    _tb = body.get("tailor_body_with_ai")
    tailor_body_with_ai = True if _tb is None else bool(_tb)
    use_jinja_renderer = USE_JINJA_LATEX_RENDERER
    if body.get("use_jinja_renderer") is not None:
        use_jinja_renderer = bool(body.get("use_jinja_renderer"))
    _sr = body.get("structured_resume")
    pre_parsed_upload: Optional[dict] = _sr if isinstance(_sr, dict) else None

    logger.info(
        f"STREAM  |  {role} @ {company}  |  model={model}  |  base={base_folder}  "
        f"|  reference_folder={reference_folder}  "
        f"|  custom_profile={bool(candidate_profile)}  |  layout_compile={layout_compile}  "
        f"|  user={user_id or 'anon'}  |  accepted_suggestions={len(accepted_suggestions) if isinstance(accepted_suggestions, list) else 0}"
        f"  |  reuse_suggest_digest={bool(suggest_research_digest)}"
        f"  |  post_suggestion_coach_run={post_suggestion_coach_run}"
        f"  |  tailor_body_with_ai={tailor_body_with_ai}"
        f"  |  use_jinja_renderer={use_jinja_renderer}"
        f"  |  has_pre_parsed_upload={bool(pre_parsed_upload)}"
    )

    if not company or not role or not jd:
        async def err_gen():
            yield {"data": json.dumps({"event": "error", "msg": "company, role, and job_description required"})}
        return EventSourceResponse(err_gen())

    loop  = asyncio.get_event_loop()
    queue: asyncio.Queue = asyncio.Queue()

    def run_sync():
        # Track local file paths from the "saved" event so we can upload to
        # Supabase Storage when the matching "pdf" event fires.
        saved_folder: Optional[str] = None
        saved_tex_path: Optional[str] = None
        # "local" = anonymous / no Supabase user — do not write under that prefix in Storage.
        storage_user = user_id if user_id and user_id != "local" else ""

        if use_jinja_renderer:
            try:
                source_folder = (reference_folder or "").strip() or (base_folder or "").strip() or "structured"
                base_tex: Optional[str] = None
                if USE_SUPABASE_TEMPLATE_BODY:
                    try:
                        _sf, _tex = _resolve_structured_source_folder(base_folder, reference_folder, user_id)
                        source_folder, base_tex = _sf, _tex
                    except Exception:
                        base_tex = None
                else:
                    base_tex = _load_tex_from_candidate(source_folder, user_id)

                asyncio.run_coroutine_threadsafe(queue.put({
                    "event": "status",
                    "msg": f"Structured renderer: loading source ({source_folder})…",
                }), loop).result()

                # Content pipeline: LLM extraction+tailoring → regex fallback → base .tex fallback.
                # The selected template controls layout; the content pipeline controls substance.
                n_approved = _count_approved_suggestions(
                    accepted_suggestions if isinstance(accepted_suggestions, list) else None
                )
                use_conservative_tailor = (
                    post_suggestion_coach_run
                    and not tailor_body_with_ai
                    and n_approved == 0
                )
                if use_conservative_tailor:
                    logger.info(
                        "Structured renderer: conservative content — skipping full LLM JD rewrite "
                        "(post_suggestion_coach_run, no approved edits)"
                    )

                if candidate_profile and jd and not use_conservative_tailor:
                    asyncio.run_coroutine_threadsafe(queue.put({
                        "event": "status",
                        "msg": "Extracting résumé sections faithfully…",
                    }), loop).result()
                doc = _structured_doc_for_generate(
                    candidate_profile,
                    jd,
                    role,
                    company,
                    use_conservative_tailor=use_conservative_tailor,
                    base_tex=base_tex,
                    pre_parsed=pre_parsed_upload,
                )
                logger.info(
                    "Structured generate | extract_path logged above | counts=%s",
                    _doc_extraction_counts(doc),
                )
                if candidate_profile and jd and not use_conservative_tailor:
                    asyncio.run_coroutine_threadsafe(queue.put({
                        "event": "status",
                        "msg": "Tailoring summary, experience, and project bullets to the job…",
                    }), loop).result()
                _apply_accepted_edits_to_doc(doc, accepted_suggestions if isinstance(accepted_suggestions, list) else None)

                asyncio.run_coroutine_threadsafe(queue.put({
                    "event": "status",
                    "msg": "Structured renderer: generating deterministic LaTeX…",
                }), loop).result()

                renderer = JinjaLatexRenderer()
                template_name = _template_name_for_reference(source_folder or reference_folder)
                new_tex = renderer.render(doc, template_name=template_name)
                asyncio.run_coroutine_threadsafe(queue.put({
                    "event": "status",
                    "msg": f"Structured renderer: using file template ({template_name})…",
                }), loop).result()

                out_folder, _ = _create_structured_output_folder(base_folder, reference_folder, role, company)
                compiled = recompile_resume_from_tex(out_folder, new_tex)
                tex_path = compiled.get("tex_path")
                pdf_path = compiled.get("pdf_path")
                filename = Path(tex_path).name if tex_path else "resume.tex"

                saved_event = {
                    "event": "saved",
                    "folder": out_folder,
                    "tex_path": tex_path,
                }

                # mirror storage upload handling for .tex
                saved_folder = saved_event.get("folder")
                saved_tex_path = saved_event.get("tex_path")
                if storage_user and saved_folder and saved_tex_path:
                    try:
                        tex_url = upload_tex(storage_user, saved_folder, saved_tex_path)
                        if tex_url:
                            asyncio.run_coroutine_threadsafe(queue.put({
                                "event": "storage",
                                "artifact": "tex",
                                "stored": True,
                                "url": tex_url,
                            }), loop).result()
                        else:
                            asyncio.run_coroutine_threadsafe(queue.put({
                                "event": "storage",
                                "artifact": "tex",
                                "stored": False,
                                "reason": storage_status().get("reason") or "Supabase upload returned no public URL",
                            }), loop).result()
                    except Exception as exc:
                        logger.warning(f"upload_tex failed: {exc}")

                asyncio.run_coroutine_threadsafe(queue.put(saved_event), loop).result()

                if not compiled.get("compiled"):
                    raise RuntimeError(compiled.get("compile_error") or "Structured renderer compile failed")

                rel_pdf = f"/pdf/{out_folder}/{Path(pdf_path).name}" if pdf_path else None
                pdf_event = {"event": "pdf", "url": rel_pdf}

                if storage_user and saved_folder and pdf_path:
                    try:
                        public = upload_pdf(storage_user, saved_folder, pdf_path)
                        if public and public.startswith(("http://", "https://")):
                            pdf_event = {"event": "pdf", "url": public}
                            asyncio.run_coroutine_threadsafe(queue.put({
                                "event": "storage",
                                "artifact": "pdf",
                                "stored": True,
                                "url": public,
                            }), loop).result()
                        else:
                            asyncio.run_coroutine_threadsafe(queue.put({
                                "event": "storage",
                                "artifact": "pdf",
                                "stored": False,
                                "reason": storage_status().get("reason") or "Supabase upload returned no public URL",
                            }), loop).result()
                    except Exception as exc:
                        logger.warning(f"upload_pdf failed: {exc}")

                # Frontend/library contract + LLM-based ratings.
                gemini_client = _optional_gemini_client()
                try:
                    llm_ratings = _rate_resume(gemini_client, model, new_tex, jd[:1500])
                except Exception:
                    llm_ratings = None
                if llm_ratings and isinstance(llm_ratings, dict):
                    if "qualifications" in llm_ratings or "responsibilities" in llm_ratings:
                        # New detailed schema
                        kw = llm_ratings.get("keywords") or {}
                        overall = int(llm_ratings.get("overall_score") or llm_ratings.get("match_score") or 0)
                        # Support both new categorized schema and legacy flat arrays
                        if isinstance(kw, dict) and ("direct_skills" in kw or "contextual" in kw):
                            # New categorized keyword schema
                            ds = kw.get("direct_skills") or {}
                            ctx = kw.get("contextual") or {}
                            ds_found = ds.get("found") or [] if isinstance(ds, dict) else []
                            ds_missing = ds.get("missing") or [] if isinstance(ds, dict) else []
                            ctx_found = ctx.get("found") or [] if isinstance(ctx, dict) else []
                            ctx_missing = ctx.get("missing") or [] if isinstance(ctx, dict) else []
                            # Normalise ctx_found to list of {keyword, count} dicts
                            ctx_found_norm = []
                            for item in ctx_found:
                                if isinstance(item, dict):
                                    ctx_found_norm.append({"keyword": str(item.get("keyword", "")), "count": int(item.get("count", 1))})
                                else:
                                    ctx_found_norm.append({"keyword": str(item), "count": 1})
                            kw_payload = {
                                "direct_skills": {"found": ds_found, "missing": ds_missing},
                                "contextual": {"found": ctx_found_norm, "missing": ctx_missing},
                                "found_count": len(ds_found) + len(ctx_found_norm),
                                "total_count": len(ds_found) + len(ds_missing) + len(ctx_found_norm) + len(ctx_missing),
                            }
                        else:
                            # Legacy flat arrays — wrap into new shape
                            found_kw = kw.get("found") or [] if isinstance(kw, dict) else []
                            missing_kw = kw.get("missing") or [] if isinstance(kw, dict) else []
                            kw_payload = {
                                "direct_skills": {"found": found_kw, "missing": missing_kw},
                                "contextual": {"found": [], "missing": []},
                                "found_count": len(found_kw),
                                "total_count": len(found_kw) + len(missing_kw),
                            }
                        ratings_payload = {
                            # New schema fields
                            "overall_score": overall,
                            "job_title": llm_ratings.get("job_title") or {},
                            "qualifications": llm_ratings.get("qualifications") or {"score": 0, "covered": [], "missing": []},
                            "responsibilities": llm_ratings.get("responsibilities") or {"score": 0, "covered": [], "missing": []},
                            "keywords": kw_payload,
                            "whats_working": llm_ratings.get("whats_working") or [],
                            "gaps": llm_ratings.get("gaps") or [],
                            "verdict": llm_ratings.get("verdict", ""),
                            # Backwards compat: keep match_score and criteria so old consumers don't break
                            "match_score": overall,
                            "criteria": [],
                        }
                    else:
                        # Old schema (fallback if model returns old format)
                        ratings_payload = {
                            "match_score": llm_ratings.get("match_score", 0),
                            "criteria": (llm_ratings.get("criteria") or [])[:12],
                            "whats_working": llm_ratings.get("whats_working") or [],
                            "gaps": llm_ratings.get("gaps") or [],
                            "verdict": llm_ratings.get("verdict", ""),
                        }
                else:
                    # Fallback to ATS-based scoring.
                    ats = ats_check(
                        out_folder, jd, storage_user or user_id,
                        target_role=role, parsed=parse_resume_tex(new_tex),
                    )
                    ratings_payload = structured_ratings_from_ats(ats)

                asyncio.run_coroutine_threadsafe(queue.put({
                    "event": "base",
                    "folder": source_folder,
                    "loaded": True,
                }), loop).result()

                diff_data, diff_adds, diff_removes, rationales_data = _structured_tailor_diff_and_rationales(
                    baseline_tex=base_tex,
                    new_tex=new_tex,
                    jd=jd,
                    model=model,
                    gemini_client=gemini_client,
                    accepted_suggestions=accepted_suggestions if isinstance(accepted_suggestions, list) else None,
                )
                asyncio.run_coroutine_threadsafe(queue.put({
                    "event": "diff",
                    "data": diff_data,
                    "adds": diff_adds,
                    "removes": diff_removes,
                }), loop).result()
                if rationales_data:
                    asyncio.run_coroutine_threadsafe(queue.put({
                        "event": "rationales",
                        "data": rationales_data,
                    }), loop).result()
                asyncio.run_coroutine_threadsafe(queue.put({
                    "event": "ratings",
                    "data": ratings_payload,
                }), loop).result()

                asyncio.run_coroutine_threadsafe(queue.put(pdf_event), loop).result()
                asyncio.run_coroutine_threadsafe(queue.put({"event": "done"}), loop).result()
                asyncio.run_coroutine_threadsafe(queue.put(None), loop).result()
                return
            except Exception as exc:
                logger.exception("structured renderer failed")
                asyncio.run_coroutine_threadsafe(queue.put({
                    "event": "error",
                    "msg": f"Structured renderer failed: {exc}",
                }), loop).result()
                asyncio.run_coroutine_threadsafe(queue.put(None), loop).result()
                return

        for event in stream_latex_resume(
            company, role, jd,
            reference_folder=reference_folder,
            model=model, base_folder=base_folder, candidate_profile=candidate_profile, user_id=user_id,
            layout_compile=layout_compile,
            accepted_suggestions=accepted_suggestions if isinstance(accepted_suggestions, list) else None,
            pre_research_digest=suggest_research_digest,
            post_suggestion_coach_run=post_suggestion_coach_run,
            tailor_body_with_ai=tailor_body_with_ai,
        ):
            ev_name = event.get("event")

            if ev_name == "saved":
                saved_folder   = event.get("folder")
                saved_tex_path = event.get("tex_path")
                # Upload the .tex source straight away — even if pdflatex fails
                # later, we still want the source preserved for diff/use-as-base.
                if storage_user and saved_folder and saved_tex_path:
                    try:
                        tex_url = upload_tex(storage_user, saved_folder, saved_tex_path)
                        if tex_url:
                            asyncio.run_coroutine_threadsafe(queue.put({
                                "event": "storage",
                                "artifact": "tex",
                                "stored": True,
                                "url": tex_url,
                            }), loop).result()
                        else:
                            asyncio.run_coroutine_threadsafe(queue.put({
                                "event": "storage",
                                "artifact": "tex",
                                "stored": False,
                                "reason": storage_status().get("reason") or "Supabase upload returned no public URL",
                            }), loop).result()
                    except Exception as exc:
                        logger.warning(f"upload_tex failed: {exc}")
                        asyncio.run_coroutine_threadsafe(queue.put({
                            "event": "storage",
                            "artifact": "tex",
                            "stored": False,
                            "reason": str(exc),
                        }), loop).result()

            elif ev_name == "pdf" and storage_user and saved_folder:
                # The library emits a relative URL like "/pdf/<folder>/<file>.pdf".
                # Resolve the local file path, push to Supabase Storage, and
                # rewrite the event so the frontend gets a durable absolute URL.
                rel_url  = event.get("url") or ""
                filename = rel_url.rsplit("/", 1)[-1] if rel_url else None
                if filename:
                    pdf_path = os.path.join(LIBRARY_ROOT, saved_folder, filename)
                    try:
                        public = upload_pdf(storage_user, saved_folder, pdf_path)
                        if public and public.startswith(("http://", "https://")):
                            event = {**event, "url": public}
                            asyncio.run_coroutine_threadsafe(queue.put({
                                "event": "storage",
                                "artifact": "pdf",
                                "stored": True,
                                "url": public,
                            }), loop).result()
                        else:
                            asyncio.run_coroutine_threadsafe(queue.put({
                                "event": "storage",
                                "artifact": "pdf",
                                "stored": False,
                                "reason": storage_status().get("reason") or "Supabase upload returned no public URL",
                            }), loop).result()
                    except Exception as exc:
                        logger.warning(f"upload_pdf failed: {exc}")
                        asyncio.run_coroutine_threadsafe(queue.put({
                            "event": "storage",
                            "artifact": "pdf",
                            "stored": False,
                            "reason": str(exc),
                        }), loop).result()

            asyncio.run_coroutine_threadsafe(queue.put(event), loop).result()
        asyncio.run_coroutine_threadsafe(queue.put(None), loop).result()

    threading.Thread(target=run_sync, daemon=True).start()

    async def event_gen():
        while True:
            item = await queue.get()
            if item is None:
                break
            yield {"data": json.dumps(item)}

    return EventSourceResponse(event_gen())


async def api_upload_resume(request: Request):
    """Extract résumé text from PDF or DOCX (MarkItDown + LLM structured parse).

    Returns ``text`` (plain, builder-friendly), ``markdown`` (raw extract),
    optional ``structured`` (JSON), ``parse_status`` (``ready`` | ``llm_failed``),
    and optional ``hints`` when structured parsing did not complete.

    Guards mirror common patterns from Resume Matcher (type, size, empty extract, clear errors).
    """
    try:
        form = await request.form()
        file = form.get("file")
        if file is None:
            return JSONResponse({"error": "No file uploaded", "code": "no_file"}, status_code=400)
        content = await file.read()
        if not content:
            return JSONResponse({"error": "Empty file", "code": "empty_file"}, status_code=400)

        filename = getattr(file, "filename", None) or "resume.pdf"
        content_type = getattr(file, "content_type", None)

        from resume_upload_parse import (
            RESUME_UPLOAD_MAX_BYTES,
            extract_upload_markdown,
            message_for_empty_resume_extract,
            parse_upload_bytes,
            validate_resume_upload_file,
        )

        try:
            validate_resume_upload_file(content_type, filename)
        except ValueError as ve:
            return JSONResponse({"error": str(ve), "code": "invalid_file_type"}, status_code=400)

        if len(content) > RESUME_UPLOAD_MAX_BYTES:
            mb = RESUME_UPLOAD_MAX_BYTES // (1024 * 1024)
            return JSONResponse(
                {
                    "error": f"File too large (maximum {mb} MB). Try compressing images or a shorter document.",
                    "code": "file_too_large",
                },
                status_code=413,
            )

        loop = asyncio.get_event_loop()

        def _extract_sync():
            return extract_upload_markdown(content, filename, pdf_plain_fallback=None)

        outcome = await loop.run_in_executor(None, _extract_sync)
        if not (outcome.markdown or "").strip():
            return JSONResponse(
                {
                    "error": message_for_empty_resume_extract(outcome.empty_reason),
                    "code": "no_extractable_text",
                    "detail": outcome.empty_reason,
                },
                status_code=422,
            )

        markdown_content = inject_section_line_breaks(outcome.markdown)

        def _pipeline_sync():
            return parse_upload_bytes(content, filename, markdown_content)

        structured, plain_text, parse_status, hints = await loop.run_in_executor(None, _pipeline_sync)

        logger.info(
            "Resume upload  |  %s  |  md_chars=%s  plain_chars=%s  parse_status=%s",
            filename,
            len(markdown_content),
            len(plain_text or ""),
            parse_status,
        )

        payload: Dict[str, Any] = {
            "text": plain_text,
            "markdown": markdown_content,
            "parse_status": parse_status,
        }
        if parse_status in ("ready", "ready_deterministic") and structured:
            payload["structured"] = structured
            log_extraction_debug(
                "upload_pipeline_final_structured",
                {
                    "education_rows": len(structured.get("education") or []),
                    "experience_rows": len(structured.get("experience") or []),
                    "structured": structured,
                },
            )
        if hints:
            payload["hints"] = hints

        return JSONResponse(payload)
    except Exception:
        logger.exception("Resume upload failed")
        return JSONResponse(
            {
                "error": "Something went wrong while processing your résumé. Please try again in a moment.",
                "code": "server_error",
            },
            status_code=500,
        )


async def api_extract_jd(request: Request):
    """Fetch a job posting URL and extract structured {company, role, location, job_description}."""
    if not ENABLE_JD_URL_EXTRACT:
        return JSONResponse(
            {
                "error": "JD URL extraction is temporarily disabled.",
                "code": "feature_disabled",
            },
            status_code=503,
        )
    try:
        body = await request.json()
        url  = (body.get("url") or "").strip()
        if not url:
            return JSONResponse({"error": "url required"}, status_code=400)
        logger.info(f"EXTRACT-JD  |  {url}")
        loop = asyncio.get_event_loop()
        data = await loop.run_in_executor(None, extract_jd_from_url, url)
        return JSONResponse(data)
    except ValueError as exc:
        return JSONResponse({"error": str(exc)}, status_code=422)
    except Exception as exc:
        logger.exception("extract-jd failed")
        return JSONResponse({"error": str(exc)}, status_code=500)


async def api_resume_parsed(request: Request):
    """GET /api/resume/{folder} — return parsed bullet tree for the editor.

    Source-of-truth resolution:
      1. Local filesystem (Railway has the freshly-generated copy in /tmp).
      2. Supabase Storage (covers re-deploys / cross-machine reads).
    """
    folder  = request.path_params["folder"]
    user_id = (request.query_params.get("user_id") or "").strip()
    if user_id == "local":
        user_id = ""
    if ".." in folder or "/" in folder:
        return JSONResponse({"error": "invalid folder"}, status_code=400)

    tex = get_resume_tex(folder)
    if tex is None and user_id:
        tex = download_tex(user_id, folder)
    if tex is None:
        return JSONResponse({"error": "resume not found"}, status_code=404)

    try:
        parsed = parse_resume_tex(tex)
    except Exception as exc:
        logger.exception("parse_resume_tex failed")
        return JSONResponse({"error": f"parse failed: {exc}"}, status_code=500)
    return JSONResponse(parsed)


async def api_resume_save(request: Request):
    """POST /api/resume/{folder} — accept edited tree, splice bullets into the
    original .tex, re-run pdflatex, push refreshed PDF to Supabase Storage,
    and return the new public URL."""
    folder = request.path_params["folder"]
    if ".." in folder or "/" in folder:
        return JSONResponse({"error": "invalid folder"}, status_code=400)

    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid json"}, status_code=400)

    user_id = (body.get("user_id") or "").strip() or "local"
    storage_user = user_id if user_id != "local" else ""
    parsed  = body.get("parsed") or {}
    if not isinstance(parsed, dict) or "sections" not in parsed:
        return JSONResponse({"error": "missing parsed.sections"}, status_code=400)

    # Source .tex — same fallback chain as the GET endpoint.
    raw_tex = parsed.get("rawTex") or get_resume_tex(folder)
    if not raw_tex and storage_user:
        raw_tex = download_tex(storage_user, folder)
    if not raw_tex:
        return JSONResponse({"error": "source .tex not found"}, status_code=404)

    new_tex = splice_bullets_into_tex(raw_tex, parsed)

    loop   = asyncio.get_event_loop()
    layout = parsed.get("pdfLayout") if isinstance(parsed, dict) else None
    result = await loop.run_in_executor(None, recompile_resume_from_tex, folder, new_tex, layout)

    if not result.get("compiled"):
        return JSONResponse({
            "error":          "recompile failed",
            "compile_error":  result.get("compile_error"),
        }, status_code=500)

    # Refresh both artifacts in Supabase so the Download button picks up the
    # new PDF and future GET /api/resume/{folder} reads see the new bullets.
    pdf_url: Optional[str] = None
    try:
        if storage_user and result.get("pdf_path"):
            pdf_url = upload_pdf(storage_user, folder, result["pdf_path"])
    except Exception as exc:
        logger.warning(f"upload_pdf (post-edit) failed: {exc}")
    try:
        if storage_user and result.get("tex_path"):
            upload_tex(storage_user, folder, result["tex_path"])
    except Exception as exc:
        logger.warning(f"upload_tex (post-edit) failed: {exc}")

    return JSONResponse({
        "folder":   folder,
        "pdf_url":  pdf_url,
        "tex_path": result.get("tex_path"),
    })


async def api_ai_edit_bullet(request: Request):
    """POST /api/ai-edit-bullet — single bullet AI rewrite for the editor."""
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid json"}, status_code=400)

    bullet_text = (body.get("bullet_text") or "").strip()
    instruction = (body.get("instruction") or "").strip()
    jd_snippet  = (body.get("jd") or "").strip()
    if not bullet_text:
        return JSONResponse({"error": "bullet_text required"}, status_code=400)

    loop = asyncio.get_event_loop()
    try:
        new_text = await loop.run_in_executor(
            None, ai_rewrite_bullet, bullet_text, instruction, jd_snippet,
        )
    except Exception as exc:
        logger.exception("ai_rewrite_bullet failed")
        return JSONResponse({"error": str(exc)}, status_code=500)
    return JSONResponse({"text": new_text})


async def api_generate_skills(request: Request):
    """POST /api/generate-skills — generate a list of skills for a job role.

    Body: {"role": "Software Engineer", "existing_skills": ["Python", ...]}
    Returns: {"skills": ["Skill 1", "Skill 2", ...]}
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid json"}, status_code=400)

    role = (body.get("role") or "").strip()
    if not role:
        return JSONResponse({"error": "role required"}, status_code=400)
    existing = body.get("existing_skills") or []
    if not isinstance(existing, list):
        existing = []

    loop = asyncio.get_event_loop()
    try:
        skills = await loop.run_in_executor(None, ai_generate_skills, role, existing)
    except Exception as exc:
        logger.exception("ai_generate_skills failed")
        return JSONResponse({"error": str(exc)}, status_code=500)
    return JSONResponse({"skills": skills})


async def api_ats_check(request: Request):
    """POST /api/ats-check/{folder} — ATS + JD alignment on the compiled PDF.

    Body JSON: ``jd`` (optional), ``user_id`` (**required** — authenticated Supabase user; rejects ``local`` or empty),
      ``target_role`` / ``role``, optional ``parsed`` sections for bullet metrics.

    Runs pdfplumber extraction and rule-based alignment in a thread-pool executor.
    """
    folder = request.path_params["folder"]
    if ".." in folder or "/" in folder:
        return JSONResponse({"error": "invalid folder"}, status_code=400)

    try:
        body = await request.json()
    except Exception:
        body = {}
    jd      = (body.get("jd") or "").strip()
    user_id = (body.get("user_id") or "").strip()
    if not user_id or user_id == "local":
        return JSONResponse({"error": "authentication required"}, status_code=401)
    target_role = (body.get("target_role") or body.get("role") or "").strip()
    parsed = body.get("parsed") if isinstance(body.get("parsed"), dict) else None

    loop = asyncio.get_event_loop()
    try:
        result = await loop.run_in_executor(
            None,
            partial(ats_check, folder, jd, user_id, None, target_role=target_role, parsed=parsed),
        )
    except FileNotFoundError as exc:
        return JSONResponse({"error": str(exc)}, status_code=404)
    except Exception as exc:
        logger.exception("ats_check failed")
        return JSONResponse({"error": str(exc)}, status_code=500)
    return JSONResponse(result)


async def api_doctor_check(request: Request):
    """POST /api/doctor-check — analyze a parsed resume tree for writing-quality
    issues (passive voice, weak verbs, missing metrics, ...). Pure regex-based,
    runs synchronously, no LLM cost.

    Body: {"parsed": ParsedResume}
    Returns: {"issues": {bullet_id: [issue, ...]}, "total": int}
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid json"}, status_code=400)
    parsed = body.get("parsed")
    if not isinstance(parsed, dict):
        return JSONResponse({"error": "parsed required"}, status_code=400)

    try:
        issues = doctor_check_resume(parsed)
    except Exception as exc:
        logger.exception("doctor_check_resume failed")
        return JSONResponse({"error": str(exc)}, status_code=500)
    total = sum(len(v) for v in issues.values())
    return JSONResponse({"issues": issues, "total": total})


def _analysis_section_scores(parsed: dict, issues: dict) -> list:
    sections = parsed.get("sections") or []
    out = []
    for sec in sections:
        name = (sec.get("name") or "Section").strip()
        bullets = []
        for e in sec.get("entries", []):
            bullets.extend(e.get("bullets", []))
        warn = 0
        info = 0
        for b in bullets:
            bid = b.get("id")
            for it in (issues.get(bid) or []):
                if it.get("severity") == "warn":
                    warn += 1
                else:
                    info += 1
        score = max(1, min(10, round(9 - warn * 1.1 - info * 0.4)))
        summary = (
            "Strong section with clear, ATS-friendly wording."
            if warn == 0 and info <= 1 else
            "Good structure, but tighten phrasing and add concrete impact in a few bullets."
            if warn <= 2 else
            "Needs cleanup: too many weak or ambiguous bullets may hurt recruiter confidence."
        )
        out.append({"name": name, "score": score, "summary": summary, "warn": warn, "info": info})
    return out


def _analysis_tips(ats: dict, sections: list, issues: dict) -> tuple[list, dict]:
    tips = []
    checks = ats.get("checks") or []
    for c in checks:
        if c.get("pass"):
            continue
        sev = "critical"
        if c.get("id") in {"word_count", "page_count", "single_column"}:
            sev = "urgent"
        tips.append({
            "severity": sev,
            "title": c.get("name") or "Fix ATS issue",
            "detail": c.get("detail") or "",
        })

    missing = [k for k in (ats.get("keywords") or []) if k.get("status") == "missing"]
    for k in missing[:3]:
        tips.append({
            "severity": "optional",
            "title": f"Add missing keyword: {k.get('keyword')}",
            "detail": "Include this term naturally in experience bullets where factual.",
        })

    jd_match = ats.get("jdMatch") or {}
    for s in (jd_match.get("missingRequiredSkills") or [])[:3]:
        title = f"JD skill to cover if true: {s}"
        if any(t.get("title") == title for t in tips):
            continue
        tips.append({
            "severity": "optional",
            "title": title,
            "detail": "Mirror the job language only where it matches your real experience.",
        })

    warn_total = sum(1 for arr in issues.values() for it in arr if it.get("severity") == "warn")
    if warn_total >= 4:
        tips.append({
            "severity": "critical",
            "title": "Strengthen weak bullets",
            "detail": "Several bullets look vague or low-signal. Lead with action + measurable outcome.",
        })

    # Dedup by title while preserving order
    seen = set()
    deduped = []
    for t in tips:
        title = t.get("title") or ""
        if title in seen:
            continue
        seen.add(title)
        deduped.append(t)

    counts = {
        "urgent": sum(1 for t in deduped if t["severity"] == "urgent"),
        "critical": sum(1 for t in deduped if t["severity"] == "critical"),
        "optional": sum(1 for t in deduped if t["severity"] == "optional"),
    }
    return deduped[:8], counts


async def api_resume_analysis(request: Request):
    """POST /api/resume-analysis/{folder} — combined ATS + writing analysis.

    Body: {"user_id": "...", "jd": "...", "parsed": ParsedResume?}
    Returns section scores + prioritized fixes for UX-friendly report views.
    """
    folder = request.path_params["folder"]
    if ".." in folder or "/" in folder:
        return JSONResponse({"error": "invalid folder"}, status_code=400)

    try:
        body = await request.json()
    except Exception:
        body = {}

    jd = (body.get("jd") or "").strip()
    user_id = (body.get("user_id") or "").strip()
    if not user_id or user_id == "local":
        return JSONResponse({"error": "authentication required"}, status_code=401)
    target_role = (body.get("target_role") or body.get("role") or "").strip()
    parsed = body.get("parsed") if isinstance(body.get("parsed"), dict) else None

    if parsed is None:
        tex = get_resume_tex(folder)
        if tex is None and user_id:
            tex = download_tex(user_id, folder)
        if tex is None:
            return JSONResponse({"error": "resume not found"}, status_code=404)
        parsed = parse_resume_tex(tex)

    loop = asyncio.get_event_loop()
    try:
        ats = await loop.run_in_executor(
            None,
            partial(ats_check, folder, jd, user_id, None, target_role=target_role, parsed=parsed),
        )
        issues = doctor_check_resume(parsed)
    except FileNotFoundError as exc:
        return JSONResponse({"error": str(exc)}, status_code=404)
    except Exception as exc:
        logger.exception("resume analysis failed")
        return JSONResponse({"error": str(exc)}, status_code=500)

    sections = _analysis_section_scores(parsed, issues)
    tips, counts = _analysis_tips(ats, sections, issues)
    overall = round(((ats.get("score") or 0) / 10 + (sum(s["score"] for s in sections) / max(1, len(sections)))) / 2)
    summary = (
        "Overall structure is strong and ATS-friendly, with a few high-impact fixes needed before submitting."
        if overall >= 7 else
        "Resume is promising but needs structural and wording improvements before applying broadly."
    )

    return JSONResponse({
        "overall": {"score": overall, "summary": summary},
        "sections": [{"name": s["name"], "score": s["score"], "summary": s["summary"]} for s in sections],
        "tips": tips,
        "counts": counts,
        "ats": ats,
    })


# ── Share links (Phase 8b) ───────────────────────────────────────────────────
import secrets
import string
_SHORTID_ALPHABET = string.ascii_lowercase + string.digits


def _gen_shortid(n: int = 8) -> str:
    return "".join(secrets.choice(_SHORTID_ALPHABET) for _ in range(n))


def _share_table():
    """Return the supabase share_links table or None if storage isn't configured."""
    try:
        try:
            from resume_gui.storage import _get_client  # type: ignore
        except ImportError:
            from storage import _get_client  # type: ignore
        client = _get_client()
        if client is None:
            return None
        return client.table("share_links")
    except Exception as exc:
        logger.warning(f"share_table unavailable: {exc}")
        return None


def _supabase_table(table_name: str):
    """Return a Supabase table handle via service-role client, else None."""
    try:
        try:
            from resume_gui.storage import _get_client  # type: ignore
        except ImportError:
            from storage import _get_client  # type: ignore
        client = _get_client()
        if client is None:
            return None
        return client.table(table_name)
    except Exception as exc:
        logger.warning(f"supabase table unavailable [{table_name}]: {exc}")
        return None


def _load_template_tex_from_supabase(reference_folder: str) -> Optional[str]:
    """Load canonical template tex from `resume_templates` table when available.

    Expected columns: reference_folder text, tex_body text, active bool.
    """
    rf = (reference_folder or "").strip()
    if not rf:
        return None
    table = _supabase_table("resume_templates")
    if table is None:
        return None
    try:
        res = (
            table.select("tex_body")
            .eq("reference_folder", rf)
            .eq("active", True)
            .limit(1)
            .execute()
        )
        rows = res.data or []
        if rows and isinstance(rows[0], dict):
            tex = str(rows[0].get("tex_body") or "").strip()
            if tex:
                logger.info(f"Loaded template tex from Supabase  |  reference_folder={rf}")
                return tex
    except Exception as exc:
        logger.warning(f"template lookup failed  |  reference_folder={rf}  |  {exc}")
    return None


async def api_share_create(request: Request):
    """POST /api/share/{folder} — mint a shortid for `folder`. Idempotent if
    the same user already created one — returns the existing one.

    Body: {"user_id": "...", "pdf_url": "..."}
    """
    folder = request.path_params["folder"]
    if ".." in folder or "/" in folder:
        return JSONResponse({"error": "invalid folder"}, status_code=400)

    try:
        body = await request.json()
    except Exception:
        body = {}
    user_id = (body.get("user_id") or "").strip()
    pdf_url = (body.get("pdf_url") or "").strip()
    if not user_id:
        return JSONResponse({"error": "user_id required"}, status_code=400)

    table = _share_table()
    if table is None:
        return JSONResponse({"error": "share storage not configured"}, status_code=503)

    # Ensure the resume exists and belongs to the caller before minting a public link.
    # Also lets us fall back to the stored PDF URL if the client did not send one.
    try:
        try:
            from resume_gui.storage import _get_client  # type: ignore
        except ImportError:
            from storage import _get_client  # type: ignore
        client = _get_client()
        if client is None:
            return JSONResponse({"error": "share storage not configured"}, status_code=503)
        resume_res = (
            client.table("resumes")
                  .select("id, pdf_url")
                  .eq("user_id", user_id)
                  .eq("folder", folder)
                  .limit(1)
                  .execute()
        )
        if not resume_res.data:
            return JSONResponse(
                {
                    "error": "Résumé not in your library yet — wait a moment and tap Share again, "
                    "or sign in and generate so we can save the row Supabase needs for links.",
                },
                status_code=404,
            )
        pdf_url = pdf_url or (resume_res.data[0].get("pdf_url") or "")
    except Exception as exc:
        logger.exception("share resume ownership lookup failed")
        return JSONResponse({"error": f"share lookup failed: {exc}"}, status_code=500)

    # Reuse existing shortid if one already exists for this user+folder.
    try:
        existing = (
            table.select("shortid, pdf_url, views, revoked")
                 .eq("user_id", user_id).eq("folder", folder)
                 .eq("revoked", False)
                 .limit(1).execute()
        )
        if existing.data:
            row = existing.data[0]
            return JSONResponse({
                "shortid": row["shortid"], "pdf_url": row.get("pdf_url"),
                "views":   row.get("views", 0), "reused": True,
            })
    except Exception as exc:
        logger.warning(f"share lookup failed: {exc}")

    # Mint a new one — retry on the (vanishingly unlikely) collision.
    for _ in range(5):
        shortid = _gen_shortid()
        try:
            table.insert({
                "shortid": shortid, "user_id": user_id,
                "folder":  folder,  "pdf_url": pdf_url or None,
            }).execute()
            return JSONResponse({"shortid": shortid, "pdf_url": pdf_url, "reused": False})
        except Exception as exc:
            msg = str(exc)
            logger.warning(f"share insert failed: {msg}")
            # Only retry actual shortid collisions; other DB errors need to be
            # surfaced so the UI/operator sees the real Supabase problem.
            if "duplicate key" in msg.lower() or "unique" in msg.lower():
                continue
            return JSONResponse({"error": f"share insert failed: {msg}"}, status_code=500)
    return JSONResponse({"error": "could not mint unique shortid after retries"}, status_code=500)


def _share_token_slug_shape(token: str) -> bool:
    """Lowercase resume `public_slug`: 3–50 chars, letters/digits, single hyphens between segments."""
    if len(token) < 3 or len(token) > 50:
        return False
    return re.match(r"^[a-z0-9]+(?:-[a-z0-9]+)*$", token) is not None


async def api_share_resolve(request: Request):
    """GET /api/share/{shortid} — resolve a share shortid or a resume `public_slug` to folder + pdf_url.
    Share rows increment the view counter; slug-based resolves do not.

    Public endpoint — used by the recipient page (no auth).
    """
    raw = (request.path_params.get("shortid") or "").strip().lower()
    if not raw:
        return JSONResponse({"error": "invalid id"}, status_code=400)

    table = _share_table()
    if table is None:
        return JSONResponse({"error": "share storage not configured"}, status_code=503)

    # ── 1) Legacy minted shortids (6–16 lowercase alnum) in share_links ─────
    if re.match(r"^[a-z0-9]{6,16}$", raw):
        try:
            rows = table.select("shortid, folder, pdf_url, views, revoked, created_at") \
                        .eq("shortid", raw).limit(1).execute()
        except Exception as exc:
            logger.exception("share resolve query failed")
            return JSONResponse({"error": str(exc)}, status_code=500)
        if rows.data:
            row = rows.data[0]
            if row.get("revoked"):
                return JSONResponse({"error": "link revoked"}, status_code=410)
            try:
                table.update({"views": (row.get("views") or 0) + 1}).eq("shortid", raw).execute()
            except Exception as exc:
                logger.warning(f"share view-counter update failed: {exc}")
            return JSONResponse({
                "shortid":    row["shortid"],
                "folder":     row["folder"],
                "pdf_url":    row.get("pdf_url"),
                "views":      (row.get("views") or 0) + 1,
                "created_at": row.get("created_at"),
            })

    # ── 2) Custom per-resume slug (resumes.public_slug, service-role read) ───
    if not _share_token_slug_shape(raw):
        return JSONResponse({"error": "invalid id"}, status_code=400)

    try:
        try:
            from resume_gui.storage import _get_client  # type: ignore
        except ImportError:
            from storage import _get_client  # type: ignore
        client = _get_client()
        if client is None:
            return JSONResponse({"error": "share storage not configured"}, status_code=503)
        resume_res = (
            client.table("resumes")
                  .select("folder, pdf_url, public_slug")
                  .eq("public_slug", raw)
                  .limit(1)
                  .execute()
        )
        if not resume_res.data:
            return JSONResponse({"error": "not found"}, status_code=404)
        r0 = resume_res.data[0]
        pdf_url = (r0.get("pdf_url") or "").strip()
        if not pdf_url:
            return JSONResponse({"error": "not found"}, status_code=404)
        return JSONResponse({
            "shortid":    raw,
            "folder":     r0.get("folder"),
            "pdf_url":    pdf_url,
            "views":      0,
            "created_at": None,
        })
    except Exception as exc:
        logger.exception("resume slug resolve failed")
        return JSONResponse({"error": str(exc)}, status_code=500)


async def api_share_revoke(request: Request):
    """DELETE /api/share/{shortid} — owner-only revoke. Body: {"user_id": "..."}.
    We require user_id match because this is what the frontend has after login.
    Service-role on the backend would let us bypass RLS, but we still scope by
    user_id to prevent cross-user revocation by a logged-in attacker."""
    shortid = request.path_params["shortid"]
    if not re.match(r"^[a-z0-9]{6,16}$", shortid or ""):
        return JSONResponse({"error": "invalid shortid"}, status_code=400)
    try:
        body = await request.json()
    except Exception:
        body = {}
    user_id = (body.get("user_id") or "").strip()
    if not user_id:
        return JSONResponse({"error": "user_id required"}, status_code=400)

    table = _share_table()
    if table is None:
        return JSONResponse({"error": "share storage not configured"}, status_code=503)
    try:
        table.update({"revoked": True}).eq("shortid", shortid).eq("user_id", user_id).execute()
    except Exception as exc:
        logger.exception("share revoke failed")
        return JSONResponse({"error": str(exc)}, status_code=500)
    return JSONResponse({"ok": True})


# ── Version History ───────────────────────────────────────────────────

async def api_version_save(request: Request):
    """POST /api/version/{folder} — save current editor state as a version.
    Body: {"user_id": "...", "parsed": "..."}"""
    folder = request.path_params["folder"]
    if ".." in folder or "/" in folder:
        return JSONResponse({"error": "invalid folder"}, status_code=400)
    
    try:
        body = await request.json()
    except Exception:
        body = {}
    user_id = (body.get("user_id") or "").strip()
    parsed = body.get("parsed")
    
    if not user_id or not parsed:
        return JSONResponse({"error": "user_id and parsed required"}, status_code=400)
    
    result = save_version(user_id, folder, json.dumps(parsed))
    if result is None:
        return JSONResponse({"error": "failed to save version"}, status_code=500)
    
    return JSONResponse(result)


async def api_version_list(request: Request):
    """GET /api/version/{folder}?user_id=xxx — list all versions."""
    folder = request.path_params["folder"]
    user_id = request.query_params.get("user_id", "").strip()
    
    if not user_id:
        return JSONResponse({"error": "user_id required"}, status_code=400)
    
    versions = list_versions(user_id, folder)
    if versions is None:
        return JSONResponse({"error": "failed to list versions"}, status_code=500)
    
    return JSONResponse({"versions": versions})


async def api_version_load(request: Request):
    """GET /api/version/{folder}/{version}?user_id=xxx — load a specific version."""
    folder = request.path_params["folder"]
    try:
        version = int(request.path_params.get("version", 0))
    except ValueError:
        return JSONResponse({"error": "invalid version"}, status_code=400)
    user_id = request.query_params.get("user_id", "").strip()
    
    if not user_id or version < 1:
        return JSONResponse({"error": "user_id and version required"}, status_code=400)
    
    parsed = load_version(user_id, folder, version)
    if parsed is None:
        return JSONResponse({"error": "version not found"}, status_code=404)
    
    return JSONResponse({"parsed": json.loads(parsed)})


async def api_storage_status(request: Request):
    """GET /api/storage-status?user_id=<uuid> — diagnose missing .tex files.

    Compares the `resumes` table for user_id against what's in the resume-tex
    Storage bucket AND the local LIBRARY_ROOT (which only ever has data in
    local dev — Railway's filesystem is empty after each deploy).

    Returns: { rows: [{folder, company, role, has_storage_tex, has_local_tex, status}], summary }
    """
    user_id = (request.query_params.get("user_id") or "").strip()
    if not user_id:
        return JSONResponse({"error": "user_id required"}, status_code=400)

    try:
        try:
            from resume_gui.storage import _get_client  # type: ignore
        except ImportError:
            from storage import _get_client  # type: ignore
        client = _get_client()
        if client is None:
            return JSONResponse({"error": "storage not configured"}, status_code=503)
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)

    loop = asyncio.get_event_loop()

    def _scan():
        db_rows = client.table("resumes").select("folder, company, role, created_at") \
                        .eq("user_id", user_id).order("created_at", desc=True).execute().data or []
        try:
            objs = client.storage.from_("resume-tex").list(user_id) or []
        except Exception:
            objs = []
        in_storage = {o["name"][:-4] for o in objs if o.get("name", "").endswith(".tex")}

        local_with_tex: set = set()
        if os.path.isdir(LIBRARY_ROOT):
            for entry in os.listdir(LIBRARY_ROOT):
                p = os.path.join(LIBRARY_ROOT, entry)
                if os.path.isdir(p) and any(f.endswith(".tex") for f in os.listdir(p)):
                    local_with_tex.add(entry)

        rows = []
        in_storage_count = recoverable = lost = 0
        for r in db_rows:
            folder   = r["folder"]
            has_stor = folder in in_storage
            has_loc  = folder in local_with_tex
            if has_stor:
                status = "in_storage"; in_storage_count += 1
            elif has_loc:
                status = "recoverable"; recoverable += 1
            else:
                status = "lost"; lost += 1
            rows.append({
                "folder":   folder,
                "company":  r.get("company"),
                "role":     r.get("role"),
                "has_storage_tex": has_stor,
                "has_local_tex":   has_loc,
                "status":   status,
            })
        return {
            "rows": rows,
            "summary": {
                "total":       len(db_rows),
                "in_storage":  in_storage_count,
                "recoverable": recoverable,
                "lost":        lost,
            },
        }

    try:
        result = await loop.run_in_executor(None, _scan)
    except Exception as exc:
        logger.exception("storage_status failed")
        return JSONResponse({"error": str(exc)}, status_code=500)
    return JSONResponse(result)


async def api_backfill_tex(request: Request):
    """POST /api/backfill-tex — upload any local-only .tex/.pdf into Storage.

    Body: {"user_id": "<uuid>", "folders": ["folder1", ...]}.
    `folders` is optional — if omitted, attempts every recoverable folder
    found by api_storage_status. Each folder must exist in the local
    LIBRARY_ROOT, otherwise it's skipped.

    Designed for one-shot recovery from a logged-in browser session — no admin
    auth gate (single-user app today). Locks down: only the row's own user_id
    can re-upload, since the path is partitioned by user_id.
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid json"}, status_code=400)
    user_id = (body.get("user_id") or "").strip()
    folders = body.get("folders") or None
    if not user_id:
        return JSONResponse({"error": "user_id required"}, status_code=400)
    if folders is not None and (not isinstance(folders, list) or not all(isinstance(f, str) for f in folders)):
        return JSONResponse({"error": "folders must be list[str]"}, status_code=400)

    loop = asyncio.get_event_loop()

    def _run():
        # Resolve the candidate list — either explicit `folders` or auto-detect.
        if folders is None:
            try:
                try:
                    from resume_gui.storage import _get_client  # type: ignore
                except ImportError:
                    from storage import _get_client  # type: ignore
                client = _get_client()
                rows = client.table("resumes").select("folder").eq("user_id", user_id).execute().data or []
                candidates = [r["folder"] for r in rows]
                objs = client.storage.from_("resume-tex").list(user_id) or []
                in_storage = {o["name"][:-4] for o in objs if o.get("name", "").endswith(".tex")}
                target = [f for f in candidates if f not in in_storage]
            except Exception as exc:
                return {"error": f"could not list candidates: {exc}"}
        else:
            target = list(folders)

        report = {"fixed": [], "skipped": [], "errors": []}
        for folder in target:
            if ".." in folder or "/" in folder:
                report["errors"].append({"folder": folder, "msg": "invalid folder name"})
                continue
            local = os.path.join(LIBRARY_ROOT, folder)
            if not os.path.isdir(local):
                report["skipped"].append({"folder": folder, "reason": "not in local LIBRARY_ROOT"})
                continue
            tex_files = [f for f in os.listdir(local) if f.endswith(".tex")]
            if not tex_files:
                report["skipped"].append({"folder": folder, "reason": "no .tex inside local folder"})
                continue
            try:
                tex_url = upload_tex(user_id, folder, os.path.join(local, tex_files[0]))
            except Exception as exc:
                report["errors"].append({"folder": folder, "msg": f"upload_tex: {exc}"})
                continue
            pdf_url = None
            pdf_files = [f for f in os.listdir(local) if f.endswith(".pdf")]
            if pdf_files:
                try:
                    pdf_url = upload_pdf(user_id, folder, os.path.join(local, pdf_files[0]))
                except Exception as exc:
                    # Non-fatal — tex is the important one for the editor.
                    logger.warning(f"backfill upload_pdf failed for {folder}: {exc}")
            report["fixed"].append({"folder": folder, "tex_url": tex_url, "pdf_url": pdf_url})
        report["summary"] = {
            "fixed":   len(report["fixed"]),
            "skipped": len(report["skipped"]),
            "errors":  len(report["errors"]),
        }
        return report

    try:
        result = await loop.run_in_executor(None, _run)
    except Exception as exc:
        logger.exception("backfill_tex failed")
        return JSONResponse({"error": str(exc)}, status_code=500)
    if isinstance(result, dict) and result.get("error"):
        return JSONResponse(result, status_code=500)
    return JSONResponse(result)


# ── Recruiter Checks ──────────────────────────────────────────────────────────


def _merge_split_word_tokens(tokens: list[str]) -> list[str]:
    """Fuse a lone uppercase letter with the following lowercase token.

    pdfplumber occasionally splits words like ``Led`` into ``L`` + ``ed``.
    """
    skip_single = frozenset({"I", "A"})
    out: list[str] = []
    i = 0
    while i < len(tokens):
        t = tokens[i]
        if (
            len(t) == 1
            and t.isalpha()
            and t.isupper()
            and t not in skip_single
            and i + 1 < len(tokens)
        ):
            nxt = tokens[i + 1]
            if nxt and len(nxt) >= 2 and nxt[0].islower():
                out.append(t + nxt)
                i += 2
                continue
        out.append(t)
        i += 1
    return out


_SECTION_KW = re.compile(
    r"^(?:"
    r"EXPERIENCE|"
    r"WORK\s+HISTORY|"
    r"WORK\s+EXPERIENCE|"
    r"PROFESSIONAL\s+EXPERIENCE|"
    r"PROFESSIONAL\s+HISTORY|"
    r"EMPLOYMENT(?:\s+HISTORY)?|"
    r"CAREER(?:\s+HISTORY|\s+OVERVIEW|\s+SUMMARY)?|"
    r"EDUCATION|"
    r"SKILLS|"
    r"SUMMARY|"
    r"PROFILE|"
    r"PROJECTS|"
    r"CERTIFICATIONS|"
    r"AWARDS|"
    r"PUBLICATIONS|"
    r"LANGUAGES|"
    r"VOLUNTEER|"
    r"PROFESSIONAL\s+SUMMARY|"
    r"TECHNICAL\s+SKILLS|"
    r"ACHIEVEMENTS?|"
    r"REFERENCES|"
    r"OBJECTIVE|"
    r"ACTIVITIES|"
    r"HONORS|"
    r"LEADERSHIP|"
    r"INTERESTS|"
    r"EXTRACURRICULAR"
    r")\s*$",
    re.IGNORECASE,
)

_HEADER_CONTACT_ANCHOR = re.compile(
    r"@|linkedin\.com/|www\.linkedin\.com/|github\.com/|www\.github\.com/|"
    r"\bportfolio\b|\bsite\b|\bmobile\b|\bphone\b|"
    r"[\[\(]?\d{3}[\])]?[\s.\-]?\d{3}[\s.\-]?\d{4}",
    re.IGNORECASE,
)

_BULLET_START_HEADER = re.compile(
    r"^[\s\ufeff]*(?:[-*•●◦·‣⁃▪►➤○⚫—–‑]|\d{1,2}[\).]\s?)",
    re.UNICODE,
)

_HEADER_JOB_ROLE = re.compile(
    r"\b(Engineer|Developer|Architect|Scientist|Analyst|Designer|Consultant|"
    r"Specialist|Manager|Director|Lead|Intern|Associate|Executive)\b",
    re.IGNORECASE,
)


def _strip_header_candidate_lines(lines: list[str], start: int, end: int) -> list[str]:
    out: list[str] = []
    lo = max(0, start)
    hi = min(len(lines), end)
    for j in range(lo, hi):
        line = lines[j].replace("\ufeff", "").strip()
        if not line:
            if len(out) >= 2:
                break
            continue
        if _SECTION_KW.match(line):
            continue
        if _BULLET_START_HEADER.match(line):
            continue
        if len(line) > 180:
            continue
        if re.search(r"%|↑|€|\$\d", line):
            continue
        out.append(line)
        if len(out) >= 8:
            break
    return out[:8]


def _header_window(lines: list[str], center_idx: int, before: int, after: int) -> list[str]:
    return _strip_header_candidate_lines(lines, center_idx - before, center_idx + after)


def _looks_like_all_caps_person_name(line: str) -> bool:
    t = line.strip()
    words = [re.sub(r"[''\-‐‑]", "", w) for w in t.split() if w]
    if len(words) < 2 or len(words) > 5 or len(t) > 48:
        return False
    if not words[0][0].isalpha():
        return False
    caps_words = [w for w in words if len(w) > 1 and w == w.upper()]
    if len(caps_words) < 2:
        return False
    return _SECTION_KW.match(t) is None


def _looks_like_title_person_name(line: str) -> bool:
    t = line.strip()
    if len(t) < 5 or len(t) > 44:
        return False
    if _HEADER_JOB_ROLE.search(t):
        return False
    words = [w for w in t.split() if w]
    if len(words) < 2 or len(words) > 4:
        return False

    def _tok_ok(w: str) -> bool:
        w = re.sub(r"[''.,]", "", w)
        return bool(re.fullmatch(r"[A-Z][a-z]+(?:-[A-Z][a-z]+)*", w))

    if not all(_tok_ok(w) for w in words):
        return False
    return _SECTION_KW.match(t) is None


def _extract_resume_header(text: str) -> list[str]:
    """Name + contact: top-of-document pass, then contact anchors / name heuristics in-body."""
    if not text.strip():
        return []

    raw_lines = text.split("\n")
    lines = [ln.replace("\ufeff", "").strip() for ln in raw_lines]

    primary: list[str] = []
    for line in lines:
        if not line:
            if len(primary) >= 2:
                break
            continue
        if _SECTION_KW.match(line):
            break
        if _BULLET_START_HEADER.match(line):
            break
        primary.append(line)
        if len(primary) >= 6:
            break
    if primary:
        return primary[:6]

    limit = min(220, len(lines))
    best: list[str] = []
    for i in range(limit):
        line = lines[i]
        if not line or len(line) > 200:
            continue
        if _HEADER_CONTACT_ANCHOR.search(line):
            chunk = _header_window(lines, i, 10, 6)
            if len(chunk) > len(best):
                best = chunk
    if not best:
        for i in range(min(360, len(lines))):
            line = lines[i]
            if _looks_like_all_caps_person_name(line) or _looks_like_title_person_name(line):
                best = _header_window(lines, i, 2, 6)
                break
    if not best:
        email_re = re.compile(r"\S+@\S+\.\S+")
        for i in range(min(400, len(lines))):
            line = lines[i]
            if line and email_re.search(line):
                best = _header_window(lines, i, 10, 6)
                break
    return best[:6]


# Bullet-glyph prefix detector used by the continuation stitcher.
_BULLET_GLYPH_LEAD_RE = re.compile(r"^[•\-\*▪▸●◦‧·・‣⁃►➤○⚫—–‑]")

# Tokens that mark a line as a structural anchor (section header, role/company
# row, education degree, contact line, project header) — these end any
# bullet-continuation chain even when the previous bullet looks unterminated.
_STITCH_STOP_LINE_RE = re.compile(
    r"^(?:"
    r"[A-Z][A-Z\s/&]+$"                  # all-caps section heading
    r"|.{0,90}\|.{0,90}\|"                # multi-pipe entry header (Co | Role | Date)
    r"|.*\b(?:19|20)\d{2}\s*[–—\-]\s*(?:(?:19|20)\d{2}|Present|Current)\b"
    r"|.*\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+(?:19|20)\d{2}\b"
    r")",
    re.IGNORECASE,
)


def _stitch_wrapped_bullets(text: str) -> str:
    """Rejoin display-line wrapped bullet bodies back onto their bullet.

    Multi-column / tightly-set PDFs make bullets wrap to several display lines.
    pdfplumber's word-bucketing then emits each wrapped line as its own row, so
    a bullet like:

        • Designed and implemented secure PHP server-side modules for form
          handling and MySQL database operations, following input-validation
          best practices to prevent injection vulnerabilities.

    comes out as 3 separate lines, and ``bulletAnalysis`` only sees the first
    one — truncating at "following" and tagging the bullet as "incomplete
    sentence". The structured-extract LLM does stitch them, but the per-bullet
    analyzer doesn't. Fix here in the extraction so both paths see the same
    complete bullet.

    Stitching rule: a non-bullet line joins the previous bullet when ALL of:
      - previous line is/was a bullet (starts with a bullet glyph OR we're
        currently mid-stitch)
      - previous bullet did NOT end with sentence-final punctuation
        (``.``/``!``/``?``/``;``) — open-clause continuations are the ones we
        want to glue
      - current line is short prose (no bullet glyph, no pipe table, not a
        section heading, not a date line)
      - current line starts with lowercase OR a hyphen-continuation
        ("input-validation" after "form handling and MySQL operations,
        following")
    """
    lines = text.splitlines()
    out: List[str] = []
    in_bullet = False
    blank_gap = 0  # consecutive blanks since the last non-blank line
    for raw in lines:
        ln = raw.rstrip()
        stripped = ln.strip()
        if not stripped:
            out.append(ln)
            blank_gap += 1
            continue
        if _BULLET_GLYPH_LEAD_RE.match(stripped):
            out.append(ln)
            in_bullet = True
            blank_gap = 0
            continue
        # Not a bullet — decide if it's a continuation. Allow up to one blank
        # line between a wrapped bullet and its continuation, since
        # pdfplumber's y-bucketing sometimes emits an empty bucket between
        # display lines that share the same logical bullet.
        if in_bullet and blank_gap <= 1 and out:
            # Find the previous non-blank line (the bullet head we want to join).
            j = len(out) - 1
            while j >= 0 and not out[j].strip():
                j -= 1
            if j >= 0:
                prev = out[j].rstrip()
                prev_tail = prev[-1:] if prev else ""
                terminated = prev_tail in (".", "!", "?", ";", ":")
                looks_anchor = bool(_STITCH_STOP_LINE_RE.match(stripped))
                starts_lower_or_continuation = bool(
                    re.match(r"^[a-z]", stripped) or stripped.startswith("-")
                )
                if (
                    not terminated
                    and not looks_anchor
                    and starts_lower_or_continuation
                    and len(stripped) <= 200
                ):
                    out[j] = f"{prev} {stripped}"
                    # Drop any blank lines we accumulated between bullet and continuation.
                    while len(out) > j + 1 and not out[-1].strip():
                        out.pop()
                    blank_gap = 0
                    # Still in bullet — next continuation can keep stitching.
                    continue
        out.append(ln)
        in_bullet = False
        blank_gap = 0
    return "\n".join(out)


def _post_clean_resume_text(text: str) -> str:
    """Normalize PDF placeholder glyphs; keep line breaks for section structure."""
    text = re.sub(r"\(\s*cid\s*:\s*\d+\)", " • ", text, flags=re.IGNORECASE)
    text = re.sub(r"\bUsed\s*:\s*to\b", "Used to", text, flags=re.IGNORECASE)
    lines = [" ".join(ln.split()) for ln in text.splitlines()]
    text = "\n".join(lines).strip()
    # Stitch BEFORE inject_section_line_breaks so the section breaker doesn't
    # see continuation fragments as standalone lines.
    text = _stitch_wrapped_bullets(text)
    try:
        return inject_section_line_breaks(text)
    except Exception:
        return text


def _sanitize_bullet_display(s: str) -> str:
    """Strip CID placeholders and collapse spaces in LLM bullet fields."""
    if not isinstance(s, str):
        return ""
    s = re.sub(r"\(\s*cid\s*:\s*\d+\)", "", s, flags=re.IGNORECASE)
    s = " ".join(s.split())
    return s.strip()


def _extract_pdf_text(pdf) -> str:
    """Extract text from a pdfplumber PDF object, reconstructing proper word spacing.

    pdfplumber's default extract_text() sometimes collapses spaces between words
    (especially in multi-column or tightly-set PDFs), producing concatenated blobs
    like 'IamaSoftwareDeveloper'.  extract_words() uses glyph bounding boxes to
    identify individual words, which we then reconstruct line-by-line.
    """
    pages_text = []
    for page in pdf.pages:
        words = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False)
        if not words:
            # Fallback to plain extract_text for image-heavy pages
            pages_text.append(page.extract_text() or "")
            continue
        # Group words by their approximate y-position (line)
        line_map: dict[int, list[str]] = {}
        for w in words:
            y_key = round(float(w["top"]) / 4) * 4  # bucket every 4pt
            line_map.setdefault(y_key, []).append(w["text"])
        page_lines = [
            " ".join(_merge_split_word_tokens(tokens))
            for _, tokens in sorted(line_map.items())
        ]
        pages_text.append("\n".join(page_lines))
    return _post_clean_resume_text("\n".join(pages_text))

_WEAK_VERBS = re.compile(
    r"\b(helped|assisted|worked on|worked with|was responsible for|participated in|"
    r"involved in|contributed to|duties included|tasked with|"
    r"did|made|got|went|had to|tried to)\b",
    re.IGNORECASE,
)
_ACTION_VERB_START_RE = re.compile(
    r"^[•\-–*▪▸]\s*(?:"
    # Management
    r"administered|analyzed|assigned|attained|chaired|consolidated|contracted|coordinated|delegated|developed|directed|evaluated|executed|improved|increased|organized|oversaw|planned|prioritized|produced|recommended|reviewed|scheduled|strengthened|supervised|"
    # Communication
    r"addressed|arbitrated|arranged|authored|collaborated|convinced|corresponded|drafted|edited|enlisted|formulated|influenced|interpreted|lectured|mediated|moderated|negotiated|persuaded|promoted|publicized|reconciled|recruited|spoke|translated|wrote|"
    # Research
    r"clarified|collected|critiqued|diagnosed|examined|extracted|identified|inspected|interviewed|investigated|summarized|surveyed|systematized|"
    # Technical
    r"assembled|built|calculated|computed|designed|devised|engineered|fabricated|maintained|operated|overhauled|programmed|remodeled|repaired|solved|upgraded|"
    # Teaching
    r"adapted|advised|coached|communicated|demystified|enabled|encouraged|explained|facilitated|guided|informed|instructed|set\s+goals|stimulated|trained|"
    # Financial / creative accomplishments
    r"acted|conceptualized|created|customized|established|fashioned|founded|illustrated|initiated|instituted|integrated|introduced|invented|originated|performed|revitalized|shaped|"
    # Helping
    r"assessed|assisted|counseled|demonstrated|educated|expedited|familiarized|motivated|referred|rehabilitated|represented|"
    # Clerical / detail
    r"approved|cataloged|classified|compiled|dispatched|generated|implemented|monitored|prepared|processed|purchased|recorded|retrieved|screened|specified|tabulated|validated|"
    # More accomplishment verbs
    r"achieved|expanded|pioneered|reduced|resolved|restored|spearheaded|transformed|"
    # Existing in-product high-signal verbs
    r"architected|automated|launched|drove|delivered|optimized"
    r")\b",
    re.IGNORECASE,
)
_PRONOUN_RE  = re.compile(r"\b(I|[Mm]e|[Mm]y|[Ww]e|[Oo]ur|[Uu]s)\b")
# Note: no IGNORECASE — we need to distinguish "us" from "US" (country code)
_DATE_RE     = re.compile(
    r"\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|"
    r"Dec(?:ember)?)[,. ]+\d{4}\b"
    r"|\b\d{4}\s*[-–]\s*(?:\d{4}|Present|Current)\b",
    re.IGNORECASE,
)
_NUMBER_RE   = re.compile(r"\b\d[\d,]*%?|\$[\d,]+[KkMmBb]?")
_EMAIL_RE    = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
_PHONE_RE    = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")
_LINKEDIN_RE = re.compile(r"linkedin\.com/in/", re.IGNORECASE)
_UNNECESSARY = re.compile(
    r"\b(references available|references furnished|references upon|"
    r"list of references|responsible for|"
    r"duties included|objective:|career objective|to obtain a|"
    r"seeking a position)\b",
    re.IGNORECASE,
)
# Passive / weak copula patterns in bullets (career-center “active not passive” guidance).
_PASSIVE_BULLET_RE = re.compile(
    r"\b(?:was|were|is|are|been|being)\s+[a-z]{2,22}(?:ed|en)\b",
    re.IGNORECASE,
)
# Experience bullets should not *start* with a date range — dates belong on role headers.
_BULLET_DATE_LEAD_RE = re.compile(
    r"^[•\-–*▪▸]\s*(?:"
    r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*[,. ]+\d{4}\b"
    r"|(?:19|20)\d{2}\s*[-–/]\s*(?:(?:19|20)\d{2}|[Pp]resent|[Cc]urrent)"
    r"|(?:19|20)\d{2}\b\s*[,–-]\s*(?:19|20)\d{2}\b"
    r")",
    re.IGNORECASE,
)


def _recruiter_checks(text: str) -> dict:
    """Run 10 recruiter checks on plain-text resume content."""
    raw_lines = [l.strip() for l in text.splitlines() if l.strip()]

    # Pass 1: merge orphan bullet glyphs (some PDFs emit "•" on its own line)
    merged: list[str] = []
    i = 0
    while i < len(raw_lines):
        ln = raw_lines[i]
        if re.match(r"^[•\-–*▪▸]\s*$", ln) and i + 1 < len(raw_lines):
            merged.append(ln.rstrip() + " " + raw_lines[i + 1])
            i += 2
        else:
            merged.append(ln)
            i += 1

    # Pass 2: merge wrapped continuation lines back into their parent bullet.
    # A continuation line is one that starts with a bullet glyph followed by a
    # lowercase letter (or a conjunction/preposition) — this happens when
    # pdfplumber assigns the bullet glyph to the second visual line of a
    # long bullet that wraps.  e.g.:
    #   "• Collaborated with the sales team …"   ← real bullet start
    #   "• the sales team to identify routes …"  ← wrapped continuation
    _CONTINUATION_RE = re.compile(r"^[•\-–*▪▸]\s+[a-z]")
    lines: list[str] = []
    for ln in merged:
        if _CONTINUATION_RE.match(ln) and lines:
            # Strip the leading glyph and space, append to previous line
            tail = re.sub(r"^[•\-–*▪▸]\s+", "", ln)
            lines[-1] = lines[-1].rstrip() + " " + tail
        else:
            lines.append(ln)

    # Lines that are clearly contact / header noise — skip from bullet lists
    _NOISE_RE = re.compile(
        r"@|linkedin|github|\.com|phone|email|mobile|location|"
        r"^\s*(education|experience|projects|skills|summary|objective|certifications)\s*$",
        re.IGNORECASE,
    )

    def _is_content_bullet(line: str) -> bool:
        if _NOISE_RE.search(line):
            return False
        words = line.split()
        # Need at least 5 whitespace-separated tokens
        if len(words) < 5:
            return False
        # Reject merged-word blobs from bad PDF extraction:
        # e.g. "IamaSoftwareDeveloperwithover5years..." → max word > 20 chars
        if max(len(w) for w in words) > 20:
            return False
        # Starts with an explicit bullet glyph — highest confidence
        if re.match(r"^[•\-–*▪▸]", line):
            return True
        # Substantive sentence starting with a capital letter (≥ 60 chars)
        if re.match(r"^[A-Z][a-z]", line) and len(line) > 60:
            return True
        return False

    bullets = [l for l in lines if _is_content_bullet(l)]
    # Explicit-bullet lines only (start with •/-/–/*): used for density/action-verb checks
    explicit_bullets = [l for l in bullets if re.match(r"^[•\-–*▪▸]", l)]

    checks = []

    # 1. Quantify impact
    unquant = [b for b in bullets if not _NUMBER_RE.search(b)]
    q_score = max(0, round(10 - len(unquant) * 1.2))
    checks.append({
        "id": "quantify", "name": "Quantified Impact",
        "score": q_score, "passed": q_score >= 7,
        "detail": (
            "Recruiters scan for numbers — percentages, dollar amounts, team sizes, "
            "timeframes. Bullets without metrics feel vague. Aim for at least 50% of "
            "your bullets to contain a quantified result."
        ),
        "items": unquant[:8],
    })

    # 2. Weak verbs
    weak_hits = [b for b in bullets if _WEAK_VERBS.search(b)]
    wv_score  = max(0, round(10 - len(weak_hits) * 2))
    checks.append({
        "id": "weak_verbs", "name": "Strong Action Verbs",
        "score": wv_score, "passed": wv_score >= 7,
        "detail": (
            "Passive or generic verbs ('helped', 'assisted', 'was responsible for') "
            "dilute impact. Replace them with strong, specific verbs from action-verb families "
            "(e.g., Managed: coordinated/oversaw; Communication: negotiated/presented; "
            "Technical: designed/engineered/programmed; Results: achieved/reduced/transformed)."
        ),
        "items": weak_hits[:8],
    })

    # 3. Action verb at start — only check explicit bullet lines
    no_action = [b for b in explicit_bullets if not _ACTION_VERB_START_RE.match(b)]
    av_score  = max(0, round(10 - len(no_action) * 1.5))
    checks.append({
        "id": "action", "name": "Action Verb Start",
        "score": av_score, "passed": av_score >= 7,
        "detail": (
            "Every bullet should start with a strong action verb (UMBC-style action list), "
            "not a noun phrase or weak helper verb. This signals initiative and makes bullets skimmable."
        ),
        "items": no_action[:6],
    })

    # 4. Pronouns
    pronoun_hits = [l for l in lines if _PRONOUN_RE.search(l)]
    pron_score   = 10 if not pronoun_hits else max(0, 10 - len(pronoun_hits) * 3)
    checks.append({
        "id": "pronouns", "name": "No Personal Pronouns",
        "score": pron_score, "passed": pron_score >= 8,
        "detail": (
            "Resumes should be written in the implied first person without using "
            "'I', 'me', 'my', 'we', etc. Remove all personal pronouns."
        ),
        "items": pronoun_hits[:6],
    })

    # 5. Repetition
    verb_counts: dict[str, int] = {}
    for b in bullets:
        m = re.match(r"^([A-Z][a-z]+)", b)
        if m:
            verb_counts[m.group(1)] = verb_counts.get(m.group(1), 0) + 1
    repeated = [f"'{v}' used {n} times" for v, n in verb_counts.items() if n >= 3]
    rep_score = 10 if not repeated else max(0, 10 - len(repeated) * 2)
    checks.append({
        "id": "repetition", "name": "Verb Variety",
        "score": rep_score, "passed": rep_score >= 7,
        "detail": (
            "Using the same action verb repeatedly makes your resume monotonous. "
            "Vary your verbs across bullets to showcase a broader skill set."
        ),
        "items": repeated,
    })

    # 6. Dates present
    has_dates = bool(_DATE_RE.search(text))
    date_score = 10 if has_dates else 0
    checks.append({
        "id": "dates", "name": "Dates Present",
        "score": date_score, "passed": has_dates,
        "detail": (
            "Recruiters need dates to understand your career timeline. "
            "Every job and education entry should include start and end dates "
            "(or 'Present' for your current role)."
        ),
        "items": [] if has_dates else ["No dates detected in the resume"],
    })

    # 7. Contact info
    has_email    = bool(_EMAIL_RE.search(text))
    has_phone    = bool(_PHONE_RE.search(text))
    has_linkedin = bool(_LINKEDIN_RE.search(text))
    contact_issues = []
    if not has_email:    contact_issues.append("Email address not found")
    if not has_phone:    contact_issues.append("Phone number not found")
    if not has_linkedin: contact_issues.append("LinkedIn URL not found")
    contact_score = 10 - len(contact_issues) * 3
    checks.append({
        "id": "contact", "name": "Contact Information",
        "score": contact_score, "passed": contact_score >= 7,
        "detail": (
            "Your contact section should include an email, phone number, and LinkedIn "
            "profile URL. Missing any of these reduces your chances of being contacted."
        ),
        "items": contact_issues,
    })

    # 8. Resume length
    word_count = len(text.split())
    if word_count < 300:
        len_score, len_note = 4, f"Too short ({word_count} words) — aim for 400–700"
    elif word_count > 900:
        len_score, len_note = 5, f"Too long ({word_count} words) — aim for 400–700"
    else:
        len_score, len_note = 10, f"Good length ({word_count} words)"
    checks.append({
        "id": "length", "name": "Resume Length",
        "score": len_score, "passed": len_score >= 7,
        "detail": (
            "A one-page resume (400–700 words) is ideal for most candidates. "
            "Two pages are acceptable for 10+ years of experience. "
            "Anything shorter looks thin; anything longer loses the reader."
        ),
        "items": [] if len_score >= 7 else [len_note],
    })

    # 9. Unnecessary phrases
    unnec_hits = list({m.group(0).lower() for m in _UNNECESSARY.finditer(text)})
    un_score   = 10 if not unnec_hits else max(0, 10 - len(unnec_hits) * 3)
    checks.append({
        "id": "unnecessary", "name": "Unnecessary Phrases",
        "score": un_score, "passed": un_score >= 8,
        "detail": (
            "Phrases like 'References available upon request' or 'Objective: To obtain a '  "
            "waste space and signal an outdated template. Remove them entirely."
        ),
        "items": [f'Remove: "{p}"' for p in unnec_hits],
    })

    # 10. Bullet density (short explicit bullets only)
    short_bullets = [b for b in explicit_bullets if len(b.split()) < 8]
    dens_score    = max(0, round(10 - len(short_bullets) * 2))
    checks.append({
        "id": "density", "name": "Bullet Depth",
        "score": dens_score, "passed": dens_score >= 7,
        "detail": (
            "Bullets under 6 words are too thin to convey impact. "
            "Each bullet should tell a mini-story: Action + Context + Result."
        ),
        "items": short_bullets[:6],
    })

    # 11. Passive / copula-heavy wording in bullets (active voice & ownership)
    passive_hits = [b for b in bullets if _PASSIVE_BULLET_RE.search(b)]
    pv_score = max(0, round(10 - len(passive_hits) * 2))
    checks.append({
        "id": "passive_voice", "name": "Active Voice & Ownership",
        "score": pv_score, "passed": pv_score >= 7,
        "detail": (
            "Use strong action verbs and active phrasing recruiters can skim in seconds. "
            "Reword passive 'was/were … done' lines and vague 'responsible for' duty dumps "
            "into owned outcomes."
        ),
        "items": passive_hits[:6],
    })

    # 12. Dates leading bullet lines (keep dates on headers; open bullets with verbs)
    date_led = [b for b in explicit_bullets if _BULLET_DATE_LEAD_RE.match(b)]
    dl_score = max(0, round(10 - len(date_led) * 2.5))
    checks.append({
        "id": "date_led_bullet", "name": "Skimmable Bullet Openings",
        "score": dl_score, "passed": dl_score >= 7,
        "detail": (
            "Start bullets with an accomplishment or action, not a calendar. "
            "Put date ranges on the role or education header line."
        ),
        "items": date_led[:6],
    })

    # 13. Role depth — detect under-described work experience entries
    # Match realistic years (1960-2029) or month+year — avoids phone-number false positives
    _ROLE_HEADER_RE = re.compile(
        r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* (?:19|20)\d{2}"
        r"|(?:19|20)\d{2}\s*[-–]\s*(?:(?:19|20)\d{2}|[Pp]resent|[Cc]urrent)",
        re.IGNORECASE,
    )
    # Lines that look like job/edu titles (contain a separator + no bullet glyph)
    _TITLE_LINE_RE = re.compile(r"[|\-–•@]|\bat\b|\bfor\b", re.IGNORECASE)

    def _parse_role_blocks(all_lines):
        roles, cur_header, cur_bullets = [], None, []
        prev_non_bullet = ""
        for ln in all_lines:
            is_bullet = bool(re.match(r"^[•\-–*▪▸]", ln))
            has_date  = bool(_ROLE_HEADER_RE.search(ln))
            if has_date and not is_bullet:
                if cur_header is not None:
                    roles.append((cur_header, cur_bullets))
                # If this line is ONLY a date range (≤ 4 tokens), prepend the previous
                # title line so we capture "Job Title | Company Name" + date
                if len(ln.split()) <= 4 and prev_non_bullet and not _ROLE_HEADER_RE.search(prev_non_bullet):
                    cur_header = prev_non_bullet + "  " + ln
                else:
                    cur_header = ln
                cur_bullets = []
            elif cur_header and is_bullet:
                cur_bullets.append(ln)
            if not is_bullet:
                prev_non_bullet = ln
        if cur_header:
            roles.append((cur_header, cur_bullets))
        return roles

    role_blocks = _parse_role_blocks(lines)
    weak_roles = []
    for header, role_bullets in role_blocks:
        has_numbers  = any(_NUMBER_RE.search(b) for b in role_bullets)
        bullet_count = len(role_bullets)
        if bullet_count < 3 or not has_numbers:
            reason = []
            if bullet_count < 3:
                reason.append(f"only {bullet_count} bullet{'s' if bullet_count != 1 else ''}")
            if not has_numbers:
                reason.append("no quantified results")
            weak_roles.append(f"{header}  [{', '.join(reason)}]")

    if role_blocks:
        rd_score = max(0, round(10 - len(weak_roles) * (10 / max(len(role_blocks), 1))))
    else:
        rd_score = 10
    checks.append({
        "id": "role_depth", "name": "Role Descriptions",
        "score": rd_score, "passed": rd_score >= 7,
        "detail": (
            "Each role should have at least 3 bullets and at least one quantified result "
            "(a percentage, dollar amount, team size, or time saved). "
            "Thin roles signal low impact to recruiters — flesh them out with specific achievements."
        ),
        "items": weak_roles[:6],
    })

    overall = round(sum(c["score"] for c in checks) / len(checks) * 10)
    passed_names = [c["name"] for c in checks if c["passed"]]
    failed_names = [c["name"] for c in checks if not c["passed"]]
    summary_ok  = ("Scored well in " + ", ".join(passed_names[:3])) if passed_names else ""
    summary_bad = ("Needs work on "  + ", ".join(failed_names[:3])) if failed_names else ""

    return {
        "overall":     overall,
        "summary_ok":  summary_ok,
        "summary_bad": summary_bad,
        "checks":      checks,
        "bullets":     bullets[:20],
    }


def _latex_to_plain(tex: str) -> str:
    """Strip common LaTeX markup to produce readable plain text for analysis."""
    # Remove comments
    tex = re.sub(r"%.*", "", tex)
    # Extract text from common resume macros
    tex = re.sub(r"\\resumeQuadHeading\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}", r"\1 \2 \3 \4", tex)
    tex = re.sub(r"\\resumeTrioHeading\{([^}]*)\}\{([^}]*)\}\{([^}]*)\}", r"\1 \2 \3", tex)
    tex = re.sub(r"\\resumeItem\{([^}]*)\}", r"• \1", tex)
    tex = re.sub(r"\\resumeItem\s*\{([^}]*)\}", r"• \1", tex)
    tex = re.sub(r"\\textbf\{([^}]*)\}", r"\1", tex)
    tex = re.sub(r"\\textit\{([^}]*)\}", r"\1", tex)
    tex = re.sub(r"\\emph\{([^}]*)\}", r"\1", tex)
    tex = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", tex)
    tex = re.sub(r"\\section\*?\{([^}]*)\}", r"\n\1\n", tex)
    # Remove remaining commands
    tex = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})*", "", tex)
    # Clean up
    tex = re.sub(r"[{}]", "", tex)
    tex = re.sub(r"\n{3,}", "\n\n", tex)
    return tex.strip()


# ── Comprehensive AI-powered resume analysis ──────────────────────────────────

_ANALYSIS_PROMPT = """\
You are an expert resume reviewer and career coach. Analyze the following resume \
using the principles below and return ONLY a valid JSON object — no markdown, no \
code fences, no prose outside the JSON.

Career-center rubric (score and write `issues[]` strings consistent with this — \
prioritize what blocks interviews and ATS passes). Primary format reference: UMBC Career Center \
\"Resume Guidelines\" (https://careers.umbc.edu/wp-content/uploads/sites/221/2015/06/Resume-Guidelines.pdf).
• UMBC section-by-section checklist (when each appears in RESUME TEXT): HEADER — name; address, city, state, zip, \
email, and phone as available for a quick contact block. OBJECTIVE — optional; one concise statement tying \
relevant skills and/or education and career goals to the target position. SUMMARY — optional; \
two to five bullets highlighting greatest strengths and skills consistent with the rest of the document; \
UMBC notes Objective and Summary are often optional when space is tight and it may be unnecessary to include both. \
EDUCATION — university (or main school): name, city, state; degree and major; graduation date; minor and/or \
certifications line when used; GPA only when explicitly stated and above 3.00; community college line if present \
with degree or dates attended pattern. CERTIFICATIONS/LICENSES — credential title and date received. \
RESEARCH, PUBLICATIONS AND PRESENTATIONS — each item: title; place or organization presented; type \
(poster, paper, oral presentation, etc.); date. RELEVANT PROJECTS — title (class/course project without course number), \
semester and year; one to two bullets on role, actions, and results; tools or techniques gained; learning \
outcomes when present. RELEVANT COURSEWORK — optional; bulleted; most applicable major/minor courses for the role; \
no more than about three lines total. SKILLS — subcategories should match the candidate's field (e.g. Laboratory / \
Quantitative / Interpersonal; clinical systems or charting for healthcare; legal research or languages for law; \
creative software for design; Programming / Software only when computing is the focus) with proficiency tiers \
(Advanced/Proficient/Novice) when used; LANGUAGES with level (conversational/fluent) when relevant. \
PROFESSIONAL EXPERIENCE (or role-focused Experience) — position title, organization, city, state, start–end dates on the \
header line; two to five action bullets emphasizing achievements, contributions, and tangible outcomes. ADDITIONAL \
EXPERIENCE — other paid roles: one to three similar bullets each; achievements not only duties. Activities tied to \
the target role may belong under Professional/Relevant Experience per UMBC. HONORS AND AWARDS — organization, award, \
date. ACTIVITIES/INTERESTS — role, organization/club, dates; one to three achievement-oriented bullets with action \
verbs. SERVICE EXPERIENCE/COMMUNITY ENGAGEMENT — organization, role, dates involved. \
• Undergraduate recency rule (only flag when resume text clearly indicates class year or high-school-era items): \
first-year students may still list high school work/activities; after second year, work and activities should be \
college-level only — mention as a sectionStructure issue only when there is explicit tension, not by guessing.
• Top problem patterns to catch: spelling/grammar/punctuation slips; missing or hard-to-parse \
contact (email, phone); passive or duty-only wording instead of owned achievements; walls of \
text or disorganized sections that fail a fast skim; bullets that never show scale, results, \
or proof of impact.
• Resume language should be: specific rather than generic; active rather than passive; \
written to express (clear facts) not to impress with fluff; fact-based — quantify and qualify \
when truthful; formatted so humans and parsers can scan headings and bullets quickly.
• Avoid: personal pronouns (I, we, my, our); unexplained heavy abbreviation; long narrative \
paragraphs where bullets are expected; slang or overly casual phrasing; “references available” \
or reference lists; opening bullets with a date or date range (dates belong on role headers).
• Encourage: consistent formatting and emphasis (spacing, bold/italics/caps); strong \
section headings in sensible order; reverse chronological experience where applicable; \
no unexplained timeline gaps when the résumé shows them; bullets that lead with strong \
action verbs and end with outcomes/scale when possible.
• Improved bullets should use precise action verbs from proven families. Prefer verbs like: \
Management (administered, coordinated, oversaw, prioritized), Communication (authored, \
negotiated, promoted, translated), Research (examined, investigated, summarized), \
Technical (engineered, programmed, upgraded), Teaching (coached, facilitated, trained), \
Financial/Creative (conceptualized, initiated, integrated), Helping (assessed, counseled, \
expedited), Clerical/Detail (implemented, monitored, validated), and Accomplishment verbs \
(achieved, reduced, spearheaded, transformed). Keep tense sensible (prior roles past tense; \
present role may mix present for ongoing scope and past for shipped wins).

ANALYSIS PRINCIPLES (map to categoryScores and bulletAnalysis):
1. READABILITY: Short, skimmable bullets; survives a ~30-second skim; white space balance; \
avoid paragraph-long bullet blobs and cluttered layout signals in text.
2. ATS COMPATIBILITY: Tables, columns, text boxes, headers/footers, images/icons, odd \
headings; standard sections (Experience, Education, Skills); contact lines machine-readable.
3. JOB MATCH: If a JD is provided — keywords, tools, responsibilities overlap; natural \
keyword placement (no stuffing). If no JD: null scores as specified below.
4. ACHIEVEMENT QUALITY: Outcomes and ownership vs. vague duties (“responsible for”, \
“worked on”, task lists without impact). Align with results-focused bullet craft. \
Do NOT fold this into quantification — weak verbs and duty-only wording are achievement problems even when numbers exist.
5. QUANTIFICATION: %, $, scale, time saved, users, rankings, before/after — reward \
truthful metrics. Aim for roughly 50% of experience bullets with measurable results (not 100%). \
Flag only the highest-impact opportunities to add numbers—especially bullets that match the JD or already show strong verbs/outcomes. \
Do NOT tag every unquantified line for quantification. \
This is separate from achievement quality: duty-language is achievement; missing metrics on a strong outcome bullet is quantification.
6. SECTION STRUCTURE: Sections and order aligned with the UMBC checklist above (header, optional Objective/Summary, \
education, optional certs/research/projects/coursework, skills, professional vs additional experience, honors, activities, \
service); enforce bullet-count norms where visible (Summary 2–5; Professional 2–5; Additional 1–3; Activities 1–3; \
Projects 1–2); flag redundant Objective + Summary when space is tight; coursework over ~3 lines; GPA not per UMBC \
(only if ≥3.0 and stated); research/pubs missing venue or presentation type when items are listed.
7. LANGUAGE QUALITY: Spelling/grammar; passive voice and buzzwords; tense; clarity over \
flowery phrasing; minimal unexplained jargon/acronyms.
8. FIELD SIGNALS & PROFESSIONAL DEPTH (JSON key MUST stay `technicalBranding` for compatibility): How clearly the \
résumé signals fit for the candidate's discipline — skills and tools grouped sensibly; field-appropriate evidence \
(e.g. portfolio or code samples for computing/design; writing or teaching clips for communications/education; \
licenses and certifications for regulated professions; publications or posters for research; patient volume or \
outcomes only when already stated). Score high when domain-relevant depth is obvious without hollow buzzwords. Do \
NOT penalize non-STEM résumés for lacking \"tech stack\" or GitHub; judge instead on clarity of training, tools, \
credentials, and outcomes that employers in THAT field expect.

SCORING GUIDANCE:
90-100 = Excellent, highly recruiter-friendly and ATS-safe.
75-89  = Strong but needs minor improvements.
60-74  = Decent but has several missed opportunities.
40-59  = Weak; needs major restructuring.
<40    = Poor; likely to fail ATS and recruiter screens.

CRITICAL RULES:
- The candidate may be in any discipline (STEM, healthcare, business, arts, education, trades, public service, etc.). \
Infer the field from the résumé text and score against that field's expectations — never assume a software-only audience.
- Be SPECIFIC, not generic. Tell exactly WHERE and HOW to fix each issue.
- When rewriting bullets, PRESERVE TRUTHFULNESS. Mark invented metrics \
  as "[X%]", "[$Y]", or "[~N]".
- For bulletAnalysis: analyze ONLY the 8 WEAKEST bullets (lowest-quality ones). \
  Skip good bullets.
- For each weak bullet, label issues clearly (quantification vs achievement vs language). \
  improvedBullet must match the primary weakness. Only ~half of weak bullets should carry a \
  quantification issue; pick JD-relevant, high-impact lines first. categoryRewrites.quantification and \
  categoryRewrites.achievementQuality must be DIFFERENT rewrites when both weaknesses apply — \
  never paste the same text into both fields.
- REWRITE HONESTY (server will reject rewrites that violate this): every improvedBullet and every \
  categoryRewrites.* value MUST preserve every numeral (digits, percentages, counts, durations) and \
  every concrete proper noun (Title-Case names, ALL-CAPS acronyms, CamelCase tech names like PostgreSQL, \
  CI/CD, AWS Lambda, gRPC) that appeared in the originalBullet. Removing "5 refinement cycles", \
  "3 retries", "PostgreSQL", or "AWS Lambda" while calling the change "readability" or \
  "language quality" is a lie — the rewrite must add JD vocabulary, not strip the original's content. \
  Rewrites should be the same length or longer than the original.
- QUANTIFICATION TRUTHFULNESS: only emit categoryRewrites.quantification when your rewrite actually \
  adds a number, percent, scale, or [X]/[%]/[$Y] placeholder that was NOT in the original. If you \
  cannot add a real or [placeholder] number, omit the quantification rewrite entirely. Never tag a \
  rewrite as quantification when no new numerals appear.
- EVIDENCE BEFORE CLAIM (server will drop unsupported claims): never claim "missing impact metrics", \
  "no quantification", or "lacks numbers" as a topIssue / atsWarning / bullet issue / final \
  recommendation when the résumé text contains numerals (digits, %, $, ×, k+, CGPA, GPA, dates, \
  counts like "9,000+ records" or "5 refinement cycles"). Count first, then claim. If the résumé \
  has 5+ numerals overall, do NOT add a "missing metrics" topIssue — the right move is to flag \
  SPECIFIC bullets that are weakest, not to claim the whole résumé lacks quantification.
- NEVER claim "duty-only bullets", "weak action verbs", or "lacks ownership" when the actual bullets \
  start with strong ownership verbs (Architected, Built, Delivered, Designed, Engineered, Developed, \
  Reduced, Implemented, Shipped, Launched, Led, Drove, Owned, Spearheaded, Scaled, Optimized, etc.). \
  Read the verbs first. If most bullets lead with strong verbs, the achievementQuality / \
  languageQuality scores cannot legitimately be below 55.
- NO TAUTOLOGICAL SUGGESTIONS: do not recommend a structural change ("create separate lines for each \
  institution", "add bullets for each role", "move skills to top") when the résumé already has that \
  structure. Read the layout before recommending a layout change.
- For each originalBullet field: copy the wording EXACTLY from RESUME TEXT (including • or -), \
  after normalizing; do not drop the first letters of words.
- If no JD is provided: set jobMatch in categoryScores to null, set \
  keywordScore to null, leave matchedKeywords/missingKeywords empty.
- Prioritize improvements that increase interview chances most.
{jd_section}

STRUCTURAL SIGNALS (deterministic pre-scan — verify against RESUME TEXT; do not invent problems):
{structural_signals}

RESUME TEXT:
{resume_text}

Return ONLY this JSON (no markdown fences, no explanation):
{{
  "overallScore": <integer 0-100>,
  "categoryScores": {{
    "readability": <0-100>,
    "atsCompatibility": <0-100>,
    "jobMatch": <0-100 or null>,
    "achievementQuality": <0-100>,
    "quantification": <0-100>,
    "sectionStructure": <0-100>,
    "languageQuality": <0-100>,
    "technicalBranding": <0-100>
  }},
  "summary": "<2-3 sentence specific overall assessment>",
  "topStrengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "topIssues": [
    {{
      "issue": "<short problem title>",
      "severity": "<low|medium|high>",
      "whyItMatters": "<1-2 sentences on impact>",
      "suggestion": "<concrete actionable fix>"
    }}
  ],
  "atsWarnings": [
    {{"warning": "<ATS issue>", "suggestion": "<fix>"}}
  ],
  "keywordAnalysis": {{
    "matchedKeywords": ["<keyword>"],
    "missingKeywords": ["<keyword>"],
    "keywordScore": <0-100 or null>,
    "suggestions": ["<where/how to naturally add missing keyword>"]
  }},
  "bulletAnalysis": [
    {{
      "originalBullet": "<exact bullet text, truncated to 150 chars>",
      "score": <0-100>,
      "issues": ["<issue 1>", "<issue 2>"],
      "improvedBullet": "<rewrite for the bullet's PRIMARY weakness only>",
      "categoryRewrites": {{
        "quantification": "<when metrics are missing: add [X%]/[$Y]/[~N] placeholders; do not invent facts>",
        "achievementQuality": "<when verbs/duties are weak: strong verb + owned outcome; omit invented metrics unless in original>"
      }}
    }}
  ],
  "sectionFeedback": [
    {{"section": "<name>", "score": <0-100>, "feedback": "<specific feedback>"}}
  ],
  "rewriteSuggestions": [
    {{"before": "<weak line>", "after": "<improved line>", "reason": "<why better>"}}
  ],
  "finalRecommendations": [
    "<most impactful action 1>",
    "<action 2>",
    "<action 3>",
    "<action 4>"
  ]
}}
"""


def _grok_reasoning_model() -> str:
    """Reasoning-tier Grok for tasks where structural accuracy matters more
    than latency (e.g. structured résumé extraction). Override via env."""
    return (os.environ.get("GROK_REASONING_MODEL") or "grok-4-fast-reasoning").strip()


def _gemini_reasoning_model() -> str:
    """Reasoning-tier Gemini fallback."""
    return (os.environ.get("GEMINI_REASONING_MODEL") or "gemini-2.5-pro").strip()


def _analysis_model() -> str:
    """Model for the main résumé-analysis prompt (the one that produces
    bulletAnalysis with issue tags + improvedBullet and topIssues +
    categoryScores). Defaults to the full grok-4 reasoning tier — same model
    we use for vision-extract — because the analysis is where dishonest tags
    and weak rewrites cost the user most. Latency hit (~8-10s) is worth the
    quality gain. Override via env: ANALYSIS_MODEL=grok-4-fast-reasoning for
    a cheaper/faster middle ground, or =grok-4-fast-non-reasoning to revert."""
    return (os.environ.get("ANALYSIS_MODEL") or "grok-4").strip()


def _llm_json_call(prompt: str, *, model_override: Optional[str] = None) -> Optional[dict]:
    """Call Grok (primary when configured) or Gemini for a JSON response.

    Pass ``model_override`` to force a specific model (e.g. the reasoning tier
    for structured extraction). When set, we use the override for the matching
    provider and fall back to the other provider's default if the override
    fails. Without override, behavior is unchanged from before.
    """
    import time
    selected_model = (model_override or primary_llm_model_for_resume_workloads()).strip()

    def _is_grok_model(model_name: str) -> bool:
        return (model_name or "").strip().lower().startswith("grok")

    def _grok_json(model_name: Optional[str] = None) -> Optional[dict]:
        xai_key = os.environ.get("XAI_API_KEY")
        if not xai_key:
            return None
        try:
            from openai import OpenAI  # type: ignore
            model = (model_name or os.environ.get("GROK_MODEL", "grok-4-1-fast-non-reasoning")).strip()
            xai = OpenAI(api_key=xai_key, base_url="https://api.x.ai/v1")
            r = xai.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                response_format={"type": "json_object"},
            )
            text = (r.choices[0].message.content or "").strip()
            text = re.sub(r"^```[a-z]*\n?", "", text)
            text = re.sub(r"\n?```$", "", text)
            return json.loads(text)
        except Exception as exc:
            logger.warning(f"Grok analysis failed: {exc}")
        return None

    def _gemini_json(model_name: Optional[str] = None) -> Optional[dict]:
        google_key = os.environ.get("GOOGLE_API_KEY") or os.environ.get("GEMINI_API_KEY")
        if not google_key:
            return None
        try:
            from google import genai as _genai  # type: ignore
            from google.genai import types as _gtypes  # type: ignore
            client = _genai.Client(api_key=google_key)
            cfg = _gtypes.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.2,
            )
            r = client.models.generate_content(
                model=(model_name or primary_gemini_flash_model()).strip(),
                contents=prompt,
                config=cfg,
            )
            text = (r.text or "").strip()
            text = re.sub(r"^```[a-z]*\n?", "", text)
            text = re.sub(r"\n?```$", "", text)
            return json.loads(text)
        except Exception as exc:
            logger.warning(f"Gemini analysis failed: {exc}")
        return None

    if _is_grok_model(selected_model):
        out = _grok_json(selected_model)
        if out is not None:
            return out
        # Cross-provider fallback. When we were forced onto a reasoning tier
        # and it failed, try the reasoning-tier Gemini equivalent before the
        # default flash model.
        if model_override:
            out = _gemini_json(_gemini_reasoning_model())
            if out is not None:
                return out
        return _gemini_json(primary_gemini_flash_model())
    out = _gemini_json(selected_model)
    if out is not None:
        return out
    if model_override:
        out = _grok_json(_grok_reasoning_model())
        if out is not None:
            return out
    return _grok_json(os.environ.get("GROK_MODEL", "grok-4-1-fast-non-reasoning"))


# ── Rewrite-vs-original validator ──────────────────────────────────────────
#
# The LLM sometimes emits a "rewrite" that silently drops numerals or proper
# nouns from the original (e.g. dropping "5 refinement cycles", "3 retries",
# "AWS Lambda", "PostgreSQL"). When that rewrite is then tagged as
# `readability` / `languageQuality`, the user sees a category badge that
# actively misrepresents what the rewrite did. Worse: a `quantification`
# rewrite that adds no numerals is just noise.
#
# These helpers diff original → rewrite and decide whether the rewrite is
# honest enough to surface. Callers should drop or fall-back when False.

# Tokens we treat as numerals (any digit run, including "5%", "3x", "$10k").
_NUMERAL_RE = re.compile(r"\d[\d.,]*")

# Tokens we treat as concrete proper nouns / named technologies — Title Case,
# all-caps acronyms, CamelCase, or slashed/dotted product names. These are
# the recruiter Ctrl-F targets that must survive any rewrite.
_PROPER_NOUN_RE = re.compile(
    r"\b(?:"
    r"[A-Z]{2,}(?:[/.][A-Z]{2,})*"          # ALL-CAPS / CI/CD / AWS
    r"|[A-Z][a-z]+(?:[A-Z][a-z]+)+"          # CamelCase / PostgreSQL / GraphQL
    r"|[A-Z][a-zA-Z0-9]+(?:\.[a-z]+)+"       # node.js, asp.net
    r")\b"
)


def _rewrite_numerals(text: str) -> set[str]:
    return set(_NUMERAL_RE.findall(text or ""))


def _rewrite_proper_nouns(text: str) -> set[str]:
    return {m.group(0) for m in _PROPER_NOUN_RE.finditer(text or "")}


def _validate_rewrite_against_original(
    original: str,
    rewrite: str,
    *,
    category: Optional[str] = None,
) -> Tuple[bool, str]:
    """Return (is_honest, reason). The rewrite is honest iff:
      - it preserves every numeral from the original
      - it preserves every concrete proper noun / acronym / CamelCase tech
        name from the original
      - it does not shrink dramatically in length (a clear sign of dropped
        content) — allow a small trim (≥ 80% of original word count)
      - if category == "quantification", at least one numeral that was not
        in the original must appear in the rewrite (otherwise the badge is
        a lie)
    Empty rewrites or rewrites identical to the original are considered
    honest at this layer (callers strip those separately).
    """
    o = (original or "").strip()
    r = (rewrite or "").strip()
    if not o or not r or o == r:
        return True, ""

    orig_nums = _rewrite_numerals(o)
    new_nums = _rewrite_numerals(r)
    dropped_nums = orig_nums - new_nums
    if dropped_nums:
        return False, f"dropped numerals: {sorted(dropped_nums)[:5]}"

    orig_props = _rewrite_proper_nouns(o)
    new_props = _rewrite_proper_nouns(r)
    dropped_props = orig_props - new_props
    if dropped_props:
        return False, f"dropped proper nouns: {sorted(dropped_props)[:5]}"

    o_wc = len(o.split())
    r_wc = len(r.split())
    if o_wc >= 8 and r_wc < int(o_wc * 0.8):
        return False, f"shrank from {o_wc}→{r_wc} words (>20% drop)"

    if category and category.lower() in ("quantification", "quantify_impact"):
        added_nums = new_nums - orig_nums
        if not added_nums:
            return False, "tagged quantification but no new numerals added"

    return True, ""


def _normalize_for_rewrite_diff(s: str) -> str:
    """Collapse whitespace + bullet glyphs + punctuation for no-op detection."""
    if not s:
        return ""
    t = re.sub(r"^[\s•\-\*▪▸●◦‧·・‣⁃►➤○⚫—–‑]+", "", s.strip())
    t = re.sub(r"\s+", " ", t)
    return t.lower().strip(" .,:;")


def _filter_bullet_rewrites(
    original: str,
    improved: str,
    category_rewrites: Optional[dict],
    issues: Optional[List[Any]] = None,
) -> Tuple[str, dict, List[Any]]:
    """Drop any rewrite that fails _validate_rewrite_against_original, drop
    no-op rewrites that are byte-equal (after whitespace/punctuation normalize)
    to the original, and strip "quantification" from the issues tag list when
    neither the rewrite nor the category rewrites actually add new numerals.

    Returns (kept_improved, kept_category_rewrites, kept_issues). When a
    rewrite is rejected we replace it with empty string so the UI falls back
    to the original — better silence than a misleading badge.
    """
    kept_improved = improved or ""
    orig_norm = _normalize_for_rewrite_diff(original)

    if kept_improved:
        ok, why = _validate_rewrite_against_original(original, kept_improved)
        if not ok:
            logger.info(
                "rewrite-validator dropped improvedBullet: %s  | orig=%r  | rewrite=%r",
                why, original[:80], kept_improved[:80],
            )
            kept_improved = ""
        elif _normalize_for_rewrite_diff(kept_improved) == orig_norm:
            # The "rewrite" is identical to the original after whitespace /
            # punctuation normalization. A no-op rewrite shown as "AI IMPROVED"
            # is dishonest — better to surface nothing.
            logger.info(
                "rewrite-validator dropped no-op improvedBullet (identical to original): %r",
                original[:80],
            )
            kept_improved = ""

    kept_cr: dict = {}
    if isinstance(category_rewrites, dict):
        for cat, txt in category_rewrites.items():
            if not isinstance(txt, str) or not txt.strip():
                continue
            ok, why = _validate_rewrite_against_original(original, txt, category=str(cat))
            if not ok:
                logger.info(
                    "rewrite-validator dropped categoryRewrites.%s: %s  | orig=%r  | rewrite=%r",
                    cat, why, original[:80], txt[:80],
                )
                continue
            if _normalize_for_rewrite_diff(txt) == orig_norm:
                logger.info(
                    "rewrite-validator dropped no-op categoryRewrites.%s (identical to original): %r",
                    cat, original[:80],
                )
                continue
            kept_cr[str(cat)] = txt

    # ── Strip lying "quantification" tag ──────────────────────────────────
    # The LLM regularly stamps a bullet with `issues: ["quantification", ...]`
    # then emits a rewrite that adds zero numerals. The badge tells the user
    # the bullet needs metrics — but the AI's own proposed fix is identical
    # prose. Drop the tag in that case so the user isn't misled.
    kept_issues: List[Any] = list(issues or [])
    if kept_issues:
        lowered = [str(t).strip().lower() for t in kept_issues]
        quant_keys = {"quantification", "quantify_impact", "quantify"}
        has_quant_tag = any(t in quant_keys for t in lowered)
        if has_quant_tag:
            orig_nums = _rewrite_numerals(original)
            added_anywhere = False
            for candidate in [kept_improved] + list(kept_cr.values()):
                if not candidate:
                    continue
                if _rewrite_numerals(candidate) - orig_nums:
                    added_anywhere = True
                    break
            if not added_anywhere:
                kept_issues = [
                    t for t in kept_issues
                    if str(t).strip().lower() not in quant_keys
                ]
                logger.info(
                    "rewrite-validator stripped 'quantification' tag — no new numerals in any rewrite: %r",
                    original[:80],
                )

    return kept_improved, kept_cr, kept_issues


# ── Evidence-grounded issue validator ──────────────────────────────────────
#
# The LLM also fabricates *issues* and *recommendations* that contradict the
# résumé text — e.g. "missing impact metrics" on a résumé with CGPA 9.266,
# 97.16%, "up to 3 retries", "9,000+ records"; or "duty-style bullets that
# fail to highlight ownership" on bullets that start with Architected /
# Built / Engineered / Delivered.
#
# When those bogus issues land in topIssues / bulletAnalysis.issues / final-
# Recommendations they (a) mislead the user and (b) crash the calibrated
# overall score in _normalize_analysis (weak_penalty + issue_penalty +
# floor_penalty). Cross-checking each claim against the actual résumé text
# is the same pattern as the rewrite validator above.

_STRONG_OWNERSHIP_VERBS = (
    # Build / ship — the core "I made this" cluster
    "architected", "built", "delivered", "designed", "engineered",
    "developed", "implemented", "shipped", "launched", "deployed",
    "rebuilt", "refactored", "modernized", "modernised", "composed",
    "produced", "authored", "drafted", "created", "generated",
    "configured", "provisioned", "automated", "migrated", "integrated",
    "orchestrated", "executed", "wrote", "rewrote",
    # Lead / own — leadership ownership cluster
    "led", "drove", "owned", "managed", "directed", "spearheaded",
    "headed", "supervised", "oversaw", "coordinated", "championed",
    "founded", "established", "originated", "pioneered", "initiated",
    "instituted", "introduced", "rolled",
    # Outcome — measurable result verbs
    "reduced", "achieved", "increased", "decreased", "improved",
    "transformed", "accelerated", "scaled", "optimized", "optimised",
    "streamlined", "boosted", "grew", "saved", "cut", "doubled",
    "tripled", "expanded", "secured", "negotiated",
    # Teach / mentor — common in education + early-career résumés;
    # genuinely ownership verbs, not duty-style
    "taught", "trained", "coached", "mentored", "facilitated",
    "instructed", "tutored", "educated", "guided", "handled",
    "presented", "moderated",
    # Investigate / resolve — analyst / engineer / consulting verbs
    "analyzed", "analysed", "researched", "investigated", "diagnosed",
    "evaluated", "assessed", "resolved", "troubleshot", "debugged",
    "audited", "validated", "verified", "tested", "benchmarked",
    # Acquisition / partnership
    "recruited", "hired", "onboarded", "partnered", "collaborated",
    "consulted", "advised",
)
_STRONG_VERB_RE = re.compile(
    r"^[\s•\-\*▪▸–—]*(?:" + "|".join(_STRONG_OWNERSHIP_VERBS) + r")\b",
    re.IGNORECASE,
)

_MISSING_METRICS_CLAIM_RE = re.compile(
    # No trailing \b — the right-hand tokens (metric, number, quantif, result)
    # often appear pluralized or suffixed (metrics / numbers / quantifiable /
    # quantification / results), and a closing \b would block those matches.
    r"\b(?:"
    r"missing\s+(?:impact\s+)?metric|missing\s+number|no\s+metric|no\s+number|"
    r"lack(?:s|ing)?\s+(?:of\s+)?(?:metric|number|data|quantif|impact|measurable|quantifi)|"
    r"lack(?:s|ing)?\s+impact|"
    r"need\s+(?:to\s+)?(?:add|include)\s+(?:metric|number|quantif)|"
    r"no\s+quantif|not\s+quantif|un[- ]?quantif|"
    r"add\s+(?:metric|number|quantifi)|specific\s+metric|fact[- ]based|"
    r"demonstrat\w*\s+result|"
    r"quantifiable\s+outcome|measurable\s+outcome|measurable\s+result"
    r")",
    re.IGNORECASE,
)

_WEAK_VERB_CLAIM_RE = re.compile(
    r"\b(?:"
    r"weak\s+(?:action\s+)?verb|weak\s+verb|"
    r"duty[- ]only|duty[- ]style|duty[- ]focused|duties?\s+included|task[- ]focused|"
    r"fail(?:s|ed|ing)?\s+to\s+highlight\s+ownership|"
    r"no\s+ownership|lack\s+ownership|responsible\s+for|helped\s+with|assisted\s+with"
    r")",
    re.IGNORECASE,
)


def _resume_numeral_count(text: str) -> int:
    """Number of distinct numeral tokens in the résumé text (CGPA 9.266 counts as 1)."""
    return len(_NUMERAL_RE.findall(text or ""))


def _resume_strong_verb_share(text: str) -> Tuple[int, int]:
    """Return (strong_verb_lead_lines, total_bullet_like_lines) over the résumé.

    Denominator is strict — only lines that start with an actual bullet glyph
    (• – * ▪ ▸) count as bullets. Tech-stack lines like 'Python · Django · …'
    or section / project headers must NOT inflate the denominator, otherwise
    the strong-verb majority floor never fires on real résumés.
    """
    if not text:
        return 0, 0
    # Optional whitespace after the glyph — bullets in real PDFs sometimes
    # come out as `•managed` with no space, e.g. when extract_words collapses
    # the glyph and first token. Require a real word char after the glyph so
    # standalone "•" lines aren't counted.
    BULLET_LEAD = re.compile(r"^[•\-\*▪▸]\s*(?=\w)")
    strong = 0
    total = 0
    for raw in text.splitlines():
        ln = raw.strip()
        if not ln:
            continue
        if not BULLET_LEAD.match(ln):
            continue
        # Drop the leading glyph + any whitespace before matching the verb.
        body = BULLET_LEAD.sub("", ln, count=1)
        total += 1
        if _STRONG_VERB_RE.search(body):
            strong += 1
    return strong, total


def _issue_text_blob(item: dict) -> str:
    parts = []
    for k in ("issue", "suggestion", "whyItMatters", "category", "warning"):
        v = item.get(k) if isinstance(item, dict) else None
        if isinstance(v, str):
            parts.append(v)
    return " ".join(parts).lower()


def _issue_contradicts_resume(
    blob: str,
    has_plenty_of_numerals: bool,
    has_strong_verb_majority: bool,
) -> Optional[str]:
    """Return a non-empty reason string when the issue's claim contradicts evidence."""
    if has_plenty_of_numerals and _MISSING_METRICS_CLAIM_RE.search(blob):
        return "claims missing metrics but résumé has plenty of numerals"
    if has_strong_verb_majority and _WEAK_VERB_CLAIM_RE.search(blob):
        return "claims weak/duty verbs but most bullets lead with strong ownership verbs"
    return None


# Thresholds. Tuned for typical 1-page résumés: ≥6 numerals is "plenty" and
# ≥60% strong-verb bullets is a "strong-verb majority". Both are conservative
# (only drop when the contradiction is overwhelming).
_NUMERAL_PLENTY_MIN = 6
_STRONG_VERB_MAJORITY_SHARE = 0.6


# Phrasings the LLM emits as atsWarnings that are actually GOOD facts about
# the résumé, not warnings. "No tables detected" doesn't deserve a warning
# label — the absence of tables is precisely what ATS-friendly résumés want.
_NON_ISSUE_ATS_WARNING_RE = re.compile(
    r"\b(?:"
    r"no\s+(?:table|graphic|image|column|header|footer|text\s*box|chart|icon|sidebar)|"
    r"no\s+multi[- ]column|"
    r"no\s+(?:embedded|complex)\s+(?:formatting|layout)|"
    r"standard\s+(?:heading|section|format)|"
    r"plain[- ]text\s+(?:heading|format|layout)|"
    r"ats[- ]friendly\s+(?:layout|format|structure)|"
    r"no\s+(?:non[- ]standard|unusual)\s+formatting"
    r")\b",
    re.IGNORECASE,
)


def _strip_non_issue_ats_warnings(raw: dict) -> None:
    """Drop atsWarnings whose text describes a non-issue (good fact framed as
    a warning). Always-on; doesn't depend on evidence thresholds because these
    phrasings are wrong on any résumé regardless of its content."""
    warnings = raw.get("atsWarnings")
    if not isinstance(warnings, list):
        return
    kept = []
    for w in warnings:
        if not isinstance(w, dict):
            kept.append(w)
            continue
        blob = _issue_text_blob(w)
        if _NON_ISSUE_ATS_WARNING_RE.search(blob):
            logger.info(
                "evidence-validator dropped non-issue atsWarning: %r",
                str(w.get("warning") or "")[:80],
            )
            continue
        kept.append(w)
    raw["atsWarnings"] = kept


def _validate_analysis_against_resume(raw: dict, resume_text: str) -> dict:
    """Drop topIssues, atsWarnings, bullet issue tags, and finalRecommendations
    whose claim contradicts what the résumé actually shows. Runs BEFORE
    _normalize_analysis so the calibrated score penalty in that function is
    based on the cleaned set."""
    if not isinstance(raw, dict):
        return raw

    text = resume_text or ""
    numeral_count = _resume_numeral_count(text)
    strong, total = _resume_strong_verb_share(text)
    strong_share = (strong / total) if total else 0.0
    has_plenty_numerals = numeral_count >= _NUMERAL_PLENTY_MIN
    has_strong_majority = total >= 5 and strong_share >= _STRONG_VERB_MAJORITY_SHARE

    # Always-on cleanup that doesn't depend on evidence thresholds: drop
    # atsWarnings phrased as non-issues. These can fire on any résumé.
    _strip_non_issue_ats_warnings(raw)

    if not (has_plenty_numerals or has_strong_majority):
        return raw  # nothing to contradict; let everything through

    # Track whether the validator actually made any changes — when it did,
    # _normalize_analysis should no longer cap the overall at the LLM's
    # (proven-untrustworthy) overallScore. Stored on a private key so it
    # doesn't leak into the API response.
    adjustments = 0

    # 1. topIssues
    issues = raw.get("topIssues")
    if isinstance(issues, list):
        kept_issues = []
        for it in issues:
            if not isinstance(it, dict):
                kept_issues.append(it)
                continue
            reason = _issue_contradicts_resume(
                _issue_text_blob(it), has_plenty_numerals, has_strong_majority,
            )
            if reason:
                logger.info(
                    "evidence-validator dropped topIssue: %s | %r",
                    reason, str(it.get("issue") or "")[:80],
                )
                adjustments += 1
                continue
            kept_issues.append(it)
        raw["topIssues"] = kept_issues

    # 2. atsWarnings — existing evidence-contradiction check
    warnings = raw.get("atsWarnings")
    if isinstance(warnings, list):
        kept_warn = []
        for w in warnings:
            if not isinstance(w, dict):
                kept_warn.append(w)
                continue
            reason = _issue_contradicts_resume(
                _issue_text_blob(w), has_plenty_numerals, has_strong_majority,
            )
            if reason:
                logger.info(
                    "evidence-validator dropped atsWarning: %s | %r",
                    reason, str(w.get("warning") or "")[:80],
                )
                adjustments += 1
                continue
            kept_warn.append(w)
        raw["atsWarnings"] = kept_warn

    # 3. bulletAnalysis.issues  — per-bullet check uses the bullet's OWN text
    bullets = raw.get("bulletAnalysis")
    if isinstance(bullets, list):
        for ba in bullets:
            if not isinstance(ba, dict):
                continue
            orig = str(ba.get("originalBullet") or "")
            if not orig:
                continue
            bullet_nums = _resume_numeral_count(orig)
            bullet_has_strong_verb = bool(_STRONG_VERB_RE.search(orig.lstrip(" \t•-*▪▸")))
            issue_tags = ba.get("issues")
            if not isinstance(issue_tags, list):
                continue
            kept_tags = []
            for tag in issue_tags:
                t = str(tag or "")
                if not t.strip():
                    continue
                tl = t.lower()
                if bullet_nums >= 2 and _MISSING_METRICS_CLAIM_RE.search(tl):
                    logger.info(
                        "evidence-validator dropped bullet issue tag: bullet has %d numerals | %r",
                        bullet_nums, t[:80],
                    )
                    adjustments += 1
                    continue
                if bullet_has_strong_verb and _WEAK_VERB_CLAIM_RE.search(tl):
                    logger.info(
                        "evidence-validator dropped bullet issue tag: bullet leads with strong verb | %r",
                        t[:80],
                    )
                    adjustments += 1
                    continue
                kept_tags.append(t)
            ba["issues"] = kept_tags

    # 4. finalRecommendations  — text-only list; same matcher
    recs = raw.get("finalRecommendations")
    if isinstance(recs, list):
        kept_recs = []
        for r in recs:
            t = str(r or "")
            if not t.strip():
                continue
            reason = _issue_contradicts_resume(
                t.lower(), has_plenty_numerals, has_strong_majority,
            )
            if reason:
                logger.info(
                    "evidence-validator dropped finalRecommendation: %s | %r",
                    reason, t[:80],
                )
                adjustments += 1
                continue
            kept_recs.append(t)
        raw["finalRecommendations"] = kept_recs

    # 5. Re-floor categoryScores that are unjustifiably low given the evidence.
    # If the résumé has plenty of numerals, quantification can't legitimately
    # be sub-55. If strong-verb majority, achievementQuality + languageQuality
    # shouldn't be sub-55 either. These are *floors*, not caps — we only
    # raise, never lower.
    cs = raw.get("categoryScores")
    if isinstance(cs, dict):
        if has_plenty_numerals:
            try:
                q = cs.get("quantification")
                if isinstance(q, (int, float)) and q < 55:
                    logger.info("evidence-validator raised quantification %s→55 (numerals=%d)", q, numeral_count)
                    cs["quantification"] = 55
                    adjustments += 1
            except Exception:
                pass
        if has_strong_majority:
            for key in ("achievementQuality", "languageQuality"):
                try:
                    v = cs.get(key)
                    if isinstance(v, (int, float)) and v < 55:
                        logger.info(
                            "evidence-validator raised %s %s→55 (strong-verb share=%.0f%%)",
                            key, v, strong_share * 100,
                        )
                        cs[key] = 55
                        adjustments += 1
                except Exception:
                    pass
        raw["categoryScores"] = cs

    # Private flag — when ≥2 adjustments fired, the LLM has demonstrably been
    # untrustworthy on this résumé. Tell _normalize_analysis to stop using the
    # LLM's overallScore as a ceiling and trust our calibration instead.
    if adjustments >= 2:
        raw["__validator_adjusted__"] = True
        logger.info("evidence-validator: %d adjustments made → LLM overall cap disabled", adjustments)
    return raw


def _normalize_analysis(raw: dict) -> dict:
    """Ensure the LLM result has all required keys with sane defaults."""
    cs_defaults = {
        "readability": 50, "atsCompatibility": 50, "jobMatch": None,
        "achievementQuality": 50, "quantification": 50,
        "sectionStructure": 50, "languageQuality": 50, "technicalBranding": 50,
    }
    cs = raw.get("categoryScores") or {}
    for k, v in cs_defaults.items():
        if k not in cs:
            cs[k] = v
    scores = [v for v in cs.values() if isinstance(v, (int, float))]
    base_overall = round(sum(scores) / len(scores)) if scores else 50
    raw.setdefault("summary", "Analysis complete.")
    raw.setdefault("topStrengths", [])
    raw.setdefault("topIssues", [])
    raw.setdefault("atsWarnings", [])
    raw.setdefault("keywordAnalysis", {
        "matchedKeywords": [], "missingKeywords": [], "keywordScore": None, "suggestions": [],
    })
    raw.setdefault("bulletAnalysis", [])
    raw.setdefault("sectionFeedback", [])
    raw.setdefault("rewriteSuggestions", [])
    raw.setdefault("finalRecommendations", [])
    raw["categoryScores"] = cs
    bullets = raw.get("bulletAnalysis") or []
    if isinstance(bullets, list):
        for ba in bullets:
            if isinstance(ba, dict):
                ba["originalBullet"] = _sanitize_bullet_display(ba.get("originalBullet", ""))
                ba["improvedBullet"] = _sanitize_bullet_display(ba.get("improvedBullet", ""))
                cr = ba.get("categoryRewrites") or ba.get("category_rewrites")
                if isinstance(cr, dict):
                    cleaned = {}
                    for ck, cv in cr.items():
                        if isinstance(cv, str) and cv.strip():
                            cleaned[str(ck)] = _sanitize_bullet_display(cv)
                    if cleaned:
                        ba["categoryRewrites"] = cleaned
                # Reject rewrites that silently drop numerals / proper nouns /
                # named technologies from the original, claim "quantification"
                # without actually adding a number, or are byte-identical to
                # the original ("AI IMPROVED" that produces the same prose).
                # Also strip a lying "quantification" tag from issues[] when
                # no rewrite actually adds numerals.
                kept_improved, kept_cr, kept_issues = _filter_bullet_rewrites(
                    ba.get("originalBullet", ""),
                    ba.get("improvedBullet", ""),
                    ba.get("categoryRewrites"),
                    ba.get("issues"),
                )
                ba["improvedBullet"] = kept_improved
                if kept_cr:
                    ba["categoryRewrites"] = kept_cr
                elif "categoryRewrites" in ba:
                    # All variant rewrites failed validation — drop the field.
                    ba.pop("categoryRewrites", None)
                ba["issues"] = kept_issues

    # Calibrate overall score so it reflects visible weaknesses.
    # This prevents inflated "90+" overall when there are many weak bullets/issues.
    weak_bullets = 0
    if isinstance(bullets, list):
        weak_bullets = sum(
            1 for ba in bullets
            if isinstance(ba, dict) and isinstance(ba.get("score"), (int, float)) and ba.get("score", 0) < 55
        )

    issues = raw.get("topIssues") or []
    high_issues = 0
    medium_issues = 0
    if isinstance(issues, list):
        for it in issues:
            if not isinstance(it, dict):
                continue
            sev = str(it.get("severity", "")).lower()
            if sev == "high":
                high_issues += 1
            elif sev == "medium":
                medium_issues += 1

    # ── Calibration v2 ────────────────────────────────────────────────────
    # Old formula stacked three penalties (weak_bullets × 3, high_issues × 4,
    # floor of (70 - min_cat) × 0.35) that double-counted the same weaknesses
    # the category scores already encode. Result: a legitimate cashier CV
    # with mean=34 got crushed to 0, and a strong résumé got pulled from a
    # fair 71 down to 62. New formula:
    #
    #   1. base_overall  = mean of category scores. This is our truth — it
    #      already reflects every weakness.
    #   2. soft_penalty  = small, capped, ONLY counts high-severity legit
    #      issues. No floor_penalty (redundant with the mean). No weak_bullets
    #      penalty (the bullet weaknesses are already in achievementQuality /
    #      quantification / languageQuality).
    #   3. blended       = midpoint of (calibrated_base, llm_overall) when the
    #      LLM's overall is plausible. When the validator fired (LLM proven
    #      untrustworthy on this résumé), we skip the blend and use the base.
    #   4. floor at 20   = any résumé that produced category scores at all
    #      deserves at least 20 (reserve <20 for non-résumés / parse failures).
    soft_penalty = min(6, high_issues * 2 + medium_issues)
    calibrated = max(0, base_overall - soft_penalty)

    llm_overall = raw.get("overallScore")
    validator_adjusted = bool(raw.pop("__validator_adjusted__", False))
    if validator_adjusted or not isinstance(llm_overall, (int, float)):
        final = calibrated
    else:
        # Blend toward the midpoint — but if the LLM diverges from base by
        # more than 20 points, the LLM is probably wrong; lean on calibrated.
        diff = llm_overall - base_overall
        if abs(diff) > 20:
            final = calibrated
        else:
            final = round((calibrated + llm_overall) / 2)

    raw["overallScore"] = int(max(20, min(100, round(final))))
    return raw


def _regex_to_comprehensive(struct: dict, jd: str) -> dict:
    """Convert _recruiter_checks output to comprehensive format (LLM unavailable fallback)."""
    checks = struct.get("checks", [])

    def _s(cid: str) -> int:
        c = next((x for x in checks if x["id"] == cid), None)
        return round((c["score"] / 10) * 100) if c else 50

    quant  = _s("quantify")
    weak   = _s("weak_verbs")
    action = _s("action")
    pron   = _s("pronouns")
    rep    = _s("repetition")
    dens   = _s("density")
    dates  = _s("dates")
    cont   = _s("contact")
    leng   = _s("length")
    rdepth = _s("role_depth")
    unnec  = _s("unnecessary")
    passive = _s("passive_voice")
    dlead = _s("date_led_bullet")

    overall = struct.get("overall", 60)
    issues = []
    for c in checks:
        if not c.get("passed"):
            sev = "high" if c["score"] < 5 else "medium"
            first_items = "; ".join(str(x) for x in (c.get("items") or [])[:2])
            issues.append({
                "issue": c["name"],
                "severity": sev,
                "whyItMatters": c.get("detail", ""),
                "suggestion": f"Fix these: {first_items}" if first_items else c.get("detail", "")[:80],
            })

    return {
        "overallScore": overall,
        "categoryScores": {
            "readability":       round((dens + leng) / 2),
            "atsCompatibility":  round((dates + cont) / 2),
            "jobMatch":          None,
            "achievementQuality": round((weak + action + rdepth) / 3),
            "quantification":    quant,
            "sectionStructure":  round((dates + unnec + dlead) / 3),
            "languageQuality":   round((pron + rep + passive) / 3),
            "technicalBranding": 50,
        },
        "summary": (
            (struct.get("summary_ok") or "") + " " + (struct.get("summary_bad") or "")
        ).strip() or "Resume analysis complete. Fix the highlighted issues to improve your score.",
        "topStrengths": [c["name"] for c in checks if c.get("passed")][:3],
        "topIssues":    issues[:6],
        "atsWarnings":  [],
        "keywordAnalysis": {
            "matchedKeywords": [], "missingKeywords": [],
            "keywordScore": None, "suggestions": [],
        },
        "bulletAnalysis":       [],
        "sectionFeedback":      [],
        "rewriteSuggestions":   [],
        "finalRecommendations": [
            "Add quantified results (%, $, numbers) to at least 50% of your bullets.",
            "Replace weak verbs (helped, assisted, worked on) with strong action verbs.",
            "Ensure every job entry has at least 3 achievement-focused bullets.",
            "Verify contact section includes email, phone, and LinkedIn URL.",
            "Make sure the resume is formatted correctly with the correct spacing and alignment.",
            "Use a professional font like Arial, Times New Roman, or Calibri for the resume.",
            "Fix spelling/grammar and verify email + phone (and LinkedIn) are obvious in the header.",
            "Replace passive or duty-only lines with action-verb openings and measurable outcomes where truthful.",
            "Tighten layout for a 30-second skim: short bullets, dates on role lines, verbs first.",
            "Drop references blurb, unexplained abbreviations, and narrative blobs — use crisp bullets.",
        ],
    }


def _analyze_resume_comprehensive(text: str, jd: str = "") -> dict:
    """Full resume analysis: structural regex + LLM deep-dive."""
    # Structural checks (fast, always run)
    struct = _recruiter_checks(text)

    # Summarize failed structural checks for the LLM (align narrative with deterministic scan)
    sig_lines: list[str] = []
    for c in struct.get("checks", []):
        if c.get("passed"):
            continue
        samp = "; ".join(str(x) for x in (c.get("items") or [])[:2])
        title = c.get("name", "Check")
        sig_lines.append(f"- {title}: {samp}" if samp else f"- {title}")
    struct_summary = (
        "\n".join(sig_lines[:14])
        if sig_lines
        else "(All automated structural checks passed.)"
    )

    # Build LLM prompt
    jd_section = (
        f"\nJOB DESCRIPTION (analyze keyword match against this):\n{jd[:3000]}"
        if jd.strip()
        else "\n(No job description provided. Set jobMatch and keywordScore to null.)"
    )
    prompt = _ANALYSIS_PROMPT.format(
        jd_section=jd_section,
        structural_signals=struct_summary,
        resume_text=text[:6000],
    )

    # Route the main analysis through the reasoning tier (default grok-4 —
    # same model used for vision-extract). Better quality on bullet-issue
    # tagging + rewrite generation, at ~8-10s extra latency per request.
    raw = _llm_json_call(prompt, model_override=_analysis_model())
    if raw and isinstance(raw, dict):
        # Strip bogus issues / recommendations that contradict the actual
        # résumé text BEFORE _normalize_analysis runs its score calibration —
        # otherwise the calibration penalty fires on lies and the overall
        # score gets crushed to 36 on a perfectly fine résumé.
        raw = _validate_analysis_against_resume(raw, text)
        return _normalize_analysis(raw)

    logger.warning("LLM unavailable for comprehensive analysis — using regex fallback")
    return _regex_to_comprehensive(struct, jd)


async def api_analyze_upload(request: Request):
    """POST /api/analyze-upload — upload a PDF or Word (.doc / .docx) and run comprehensive AI analysis.

    Form fields:
      file  — PDF or Word binary
      jd    — optional job description text
    """
    try:
        form   = await request.form()
        upload = form.get("file")
        jd     = (form.get("jd") or "").strip()
        if not upload:
            return JSONResponse({"error": "No file provided"}, status_code=400)
        data     = await upload.read()
        if not data:
            return JSONResponse({"error": "Empty file"}, status_code=400)
        filename = getattr(upload, "filename", None) or "resume.pdf"
        content_type = getattr(upload, "content_type", None)

        from resume_upload_parse import (  # type: ignore
            extract_upload_markdown,
            message_for_empty_resume_extract,
            validate_resume_upload_file,
        )

        try:
            validate_resume_upload_file(content_type, filename)
        except ValueError as ve:
            return JSONResponse({"error": str(ve)}, status_code=400)

        loop = asyncio.get_event_loop()

        def _extract_text() -> str:
            outcome = extract_upload_markdown(data, filename, pdf_plain_fallback=None)
            text = (outcome.markdown or "").strip()
            if text:
                return text
            # pdfplumber fallback — PDF only (never treat Word as PDF)
            if filename.lower().endswith(".pdf"):
                try:
                    with pdfplumber.open(io.BytesIO(data)) as pdf:
                        return _extract_pdf_text(pdf)
                except Exception as exc:
                    logger.warning("pdfplumber fallback failed: %s", exc)
            if outcome.empty_reason:
                raise ValueError(message_for_empty_resume_extract(outcome.empty_reason))
            return ""

        try:
            text = await loop.run_in_executor(None, _extract_text)
        except ValueError as ve:
            return JSONResponse({"error": str(ve)}, status_code=422)

        if not text.strip():
            return JSONResponse(
                {"error": "Could not extract text from the uploaded file"},
                status_code=422,
            )

        # When vision-PDF extraction is available, run it FIRST and synthesize
        # a clean text representation from the structured doc. Then use that
        # clean text for both the analysis prompt AND the right-panel preview
        # so the LLM grades what the user actually wrote (not column-extraction
        # artifacts) and the preview matches the structured understanding.
        # When vision is unavailable / fails, fall back to the legacy
        # concurrent path (analysis on raw text, structured extract via the
        # reasoning text path).
        pdf_bytes_for_vision = data if filename.lower().endswith(".pdf") else None
        structured: Optional[ResumeDocModel] = None
        analysis_input_text = text
        preview_text = text

        if pdf_bytes_for_vision:
            structured = await loop.run_in_executor(
                None, _llm_extract, text, pdf_bytes_for_vision,
            )
            if structured is not None:
                synthesized = _synthesize_text_from_resume_doc(structured)
                if synthesized.strip():
                    analysis_input_text = synthesized
                    preview_text = synthesized
                    logger.info(
                        "analyze_upload: using vision-synthesized text "
                        "(%d → %d chars) for analysis + preview",
                        len(text), len(synthesized),
                    )
            result = await loop.run_in_executor(
                None, _analyze_resume_comprehensive, analysis_input_text, jd,
            )
        else:
            analysis_future = loop.run_in_executor(
                None, _analyze_resume_comprehensive, text, jd,
            )
            extract_future = loop.run_in_executor(
                None, _llm_extract, text, None,
            )
            result, structured = await asyncio.gather(analysis_future, extract_future)

        if isinstance(result, dict):
            result["extractedText"] = preview_text[:25000]
            result["resumeHeader"]  = _extract_resume_header(preview_text)

        if structured is not None:
            # Build flat bulletMap: index in bulletAnalysis → {experienceIdx, bulletIdx}
            bullet_map: list[dict] = []
            for ei, exp in enumerate(structured.experience):
                for bi in range(len(exp.bullets)):
                    bullet_map.append({"experienceIdx": ei, "bulletIdx": bi})
            result["structuredResume"] = _resume_doc_to_dict(structured)
            result["bulletMap"]        = bullet_map
            _log_structured_doc("analyze_upload_structured_resume", structured)

        # Persist analysis result for student history + cohort analytics (best-effort)
        user_id = (form.get("user_id") or "").strip()
        user_email = (form.get("user_email") or "").strip()
        if user_id and isinstance(result, dict):
            try:
                await loop.run_in_executor(None, _persist_analysis, result, user_id, user_email, bool(jd))
            except Exception:
                pass  # never block the response for analytics writes

        return JSONResponse(result)
    except Exception as exc:
        logger.exception("analyze_upload failed")
        return JSONResponse({"error": str(exc)}, status_code=500)


def _persist_analysis(result: dict, user_id: str, user_email: str, has_jd: bool) -> None:
    """Write analysis result to Supabase `resume_analyses` table for history + cohort stats."""
    table = _supabase_table("resume_analyses")
    if table is None:
        return
    header = (result.get("resumeHeader") or "").strip()
    label  = header[:80] if header else ("With JD" if has_jd else "General")
    row = {
        "user_id":    user_id,
        "user_email": user_email or None,
        "label":      label,
        "score":      result.get("overallScore"),
        "result":     result,
    }
    try:
        table.insert(row).execute()
    except Exception as exc:
        logger.warning("resume_analyses insert failed: %s", exc)


async def api_my_analyses(request: Request):
    """GET /api/my-analyses — return current user's analysis history (latest 50)."""
    user_id = (request.query_params.get("user_id") or "").strip()
    if not user_id:
        return JSONResponse({"error": "user_id required"}, status_code=400)
    table = _supabase_table("resume_analyses")
    if table is None:
        return JSONResponse({"analyses": [], "note": "storage unavailable"})
    try:
        resp = (
            table
            .select("id,label,score,created_at,result->overallScore,result->categoryScores,result->topIssues,result->topStrengths,result->keywordAnalysis")
            .eq("user_id", user_id)
            .order("created_at", desc=True)
            .limit(50)
            .execute()
        )
        rows = []
        for r in (resp.data or []):
            rows.append({
                "id":             r.get("id"),
                "label":          r.get("label"),
                "score":          r.get("score"),
                "created_at":     r.get("created_at"),
                "categoryScores": (r.get("result") or {}).get("categoryScores"),
                "topIssues":      (r.get("result") or {}).get("topIssues"),
                "topStrengths":   (r.get("result") or {}).get("topStrengths"),
                "keywordScore":   ((r.get("result") or {}).get("keywordAnalysis") or {}).get("keywordScore"),
            })
        return JSONResponse({"analyses": rows})
    except Exception as exc:
        logger.warning("my_analyses query failed: %s", exc)
        return JSONResponse({"error": str(exc)}, status_code=500)


async def api_cohort_stats(request: Request):
    """GET /api/cohort-stats — aggregate analysis stats across all students.

    Checks X-User-Email header against the ADVISOR_EMAILS env var (comma-separated list).
    Returns cohort-level score distributions, weakest dimensions, and
    most common issue types — suitable for career center dashboards.
    """
    import os as _os
    advisor_emails = {e.strip().lower() for e in _os.environ.get("ADVISOR_EMAILS", "").split(",") if e.strip()}
    user_email     = (request.headers.get("x-user-email") or "").strip().lower()
    if not advisor_emails or user_email not in advisor_emails:
        return JSONResponse({"error": "unauthorized"}, status_code=403)

    table = _supabase_table("resume_analyses")
    if table is None:
        return JSONResponse({"error": "storage unavailable"}, status_code=503)

    try:
        resp = (
            table
            .select("user_id,user_email,score,result,created_at")
            .order("created_at", desc=True)
            .limit(2000)
            .execute()
        )
        rows = resp.data or []
    except Exception as exc:
        logger.warning("cohort_stats query failed: %s", exc)
        return JSONResponse({"error": str(exc)}, status_code=500)

    if not rows:
        return JSONResponse({"student_count": 0, "analysis_count": 0})

    from collections import Counter as _Counter

    def _avg(vals):
        clean = [v for v in vals if isinstance(v, (int, float))]
        return round(sum(clean) / len(clean), 1) if clean else None

    dim_keys = ["readability", "atsCompatibility", "jobMatch", "achievementQuality",
                "quantification", "sectionStructure", "languageQuality", "technicalBranding"]

    overall_scores = [r.get("score") for r in rows if r.get("score") is not None]
    dim_avgs = {}
    for dk in dim_keys:
        vals = [(r.get("result") or {}).get("categoryScores", {}).get(dk) for r in rows]
        dim_avgs[dk] = _avg(vals)

    # --- score tier distribution ---
    tiers = {"low": 0, "mid": 0, "good": 0, "strong": 0}
    for s in overall_scores:
        if s < 50:       tiers["low"]    += 1
        elif s < 70:     tiers["mid"]    += 1
        elif s < 85:     tiers["good"]   += 1
        else:            tiers["strong"] += 1

    # --- most common issues ---
    issue_counter: _Counter = _Counter()
    for r in rows:
        for issue in ((r.get("result") or {}).get("topIssues") or []):
            title = (issue.get("title") or issue.get("issue")) if isinstance(issue, dict) else str(issue)
            if title:
                issue_counter[title.strip()] += 1
    top_issues = [{"issue": k, "count": v} for k, v in issue_counter.most_common(10)]

    # --- per-student latest scores (advisor roster) ---
    seen: dict = {}
    for r in rows:
        uid = r.get("user_id")
        if uid and uid not in seen:
            seen[uid] = {
                "user_id":        uid,
                "user_email":     r.get("user_email"),
                "latest_score":   r.get("score"),
                "latest_at":      r.get("created_at"),
                "analysis_count": sum(1 for x in rows if x.get("user_id") == uid),
            }
    student_roster = sorted(seen.values(), key=lambda x: x.get("latest_score") or 0)

    weakest_dims = sorted(
        [(dk, dim_avgs[dk]) for dk in dim_keys if dim_avgs[dk] is not None],
        key=lambda x: x[1]
    )

    return JSONResponse({
        "student_count":   len(seen),
        "analysis_count":  len(rows),
        "avg_overall":     _avg(overall_scores),
        "score_tiers":     tiers,
        "dimension_avgs":  dim_avgs,
        "weakest_dims":    [{"dimension": d, "avg": v} for d, v in weakest_dims[:3]],
        "top_issues":      top_issues,
        "student_roster":  student_roster,
        "generated_at":    __import__("datetime").datetime.utcnow().isoformat() + "Z",
    })



async def api_student_detail(request: Request):
    """GET /api/student-detail?student_id=<uuid> — full analysis history + resume list for one student.

    Advisor-only: requires X-User-Email header matching ADVISOR_EMAILS env var.
    """
    import os as _os
    advisor_emails = {e.strip().lower() for e in _os.environ.get("ADVISOR_EMAILS", "").split(",") if e.strip()}
    user_email     = (request.headers.get("x-user-email") or "").strip().lower()
    if not advisor_emails or user_email not in advisor_emails:
        return JSONResponse({"error": "unauthorized"}, status_code=403)

    student_id = (request.query_params.get("student_id") or "").strip()
    if not student_id:
        return JSONResponse({"error": "student_id required"}, status_code=400)

    analyses_table = _supabase_table("resume_analyses")
    resumes_table  = _supabase_table("resumes")
    if analyses_table is None:
        return JSONResponse({"error": "storage unavailable"}, status_code=503)

    try:
        a_resp = (
            analyses_table
            .select("id,user_email,label,score,result,created_at")
            .eq("user_id", student_id)
            .order("created_at", desc=False)
            .limit(100)
            .execute()
        )
        analyses = a_resp.data or []
    except Exception as exc:
        logger.warning("student_detail analyses query failed: %s", exc)
        return JSONResponse({"error": str(exc)}, status_code=500)

    resumes: list[dict] = []
    if resumes_table is not None:
        try:
            r_resp = (
                resumes_table
                .select("folder,company,role,score,pdf_url,created_at")
                .eq("user_id", student_id)
                .order("created_at", desc=True)
                .limit(50)
                .execute()
            )
            resumes = r_resp.data or []
        except Exception as exc:
            logger.warning("student_detail resumes query failed: %s", exc)

    # Build score history and extract latest analysis details
    score_history = [
        {"date": r.get("created_at"), "score": r.get("score"), "label": r.get("label")}
        for r in analyses
    ]
    latest = analyses[-1] if analyses else None
    latest_result = (latest or {}).get("result") or {}

    # Compute per-student averages across all analyses
    dim_keys = ["readability", "atsCompatibility", "jobMatch", "achievementQuality",
                "quantification", "sectionStructure", "languageQuality", "technicalBranding"]

    def _avg(vals):
        clean = [v for v in vals if isinstance(v, (int, float))]
        return round(sum(clean) / len(clean), 1) if clean else None

    dim_avgs = {}
    for dk in dim_keys:
        vals = [(r.get("result") or {}).get("categoryScores", {}).get(dk) for r in analyses]
        dim_avgs[dk] = _avg(vals)

    # Issue frequency across all this student's analyses
    from collections import Counter as _Counter
    issue_counter: _Counter = _Counter()
    for r in analyses:
        for issue in ((r.get("result") or {}).get("topIssues") or []):
            title = (issue.get("title") or issue.get("issue")) if isinstance(issue, dict) else str(issue)
            if title:
                issue_counter[title.strip()] += 1
    top_issues = [{"issue": k, "count": v} for k, v in issue_counter.most_common(8)]

    first_score  = analyses[0].get("score")  if analyses else None
    latest_score = analyses[-1].get("score") if analyses else None
    delta = (latest_score - first_score) if (first_score is not None and latest_score is not None) else None

    return JSONResponse({
        "student_id":     student_id,
        "user_email":     (analyses[0].get("user_email") if analyses else None),
        "analysis_count": len(analyses),
        "first_score":    first_score,
        "latest_score":   latest_score,
        "score_delta":    delta,
        "score_history":  score_history,
        "dim_avgs":       dim_avgs,
        "top_issues":     top_issues,
        "latest_strengths":     latest_result.get("topStrengths") or [],
        "latest_category_scores": latest_result.get("categoryScores") or {},
        "resumes":        resumes,
        "latest_resume_text":  (latest_result.get("extractedText") or "")[:4000],
    })


async def api_analyze_folder(request: Request):
    """POST /api/analyze-folder/{folder} — run comprehensive analysis on a stored resume."""
    folder = request.path_params.get("folder", "").strip()
    if not folder or ".." in folder:
        return JSONResponse({"error": "invalid folder"}, status_code=400)

    body: dict = {}
    try:
        body = await request.json()
    except Exception:
        pass
    user_id = body.get("user_id", "")
    jd      = (body.get("jd") or "").strip()

    loop = asyncio.get_event_loop()

    def _run():
        # 1. Local filesystem (fresh Railway deploy or local dev)
        tex_path = os.path.join(LIBRARY_ROOT, folder, "resume.tex")
        if os.path.isfile(tex_path):
            with open(tex_path, encoding="utf-8", errors="ignore") as f:
                return _latex_to_plain(f.read())

        # 2. Supabase Storage via the shared download_tex helper
        if user_id:
            try:
                tex = download_tex(user_id, folder)
                if tex:
                    return _latex_to_plain(tex)
            except Exception as e:
                logger.warning(f"analyze_folder: download_tex failed: {e}")
        return None

    plain = await loop.run_in_executor(None, _run)
    if not plain:
        return JSONResponse({"error": "Could not load resume text"}, status_code=404)

    result = await loop.run_in_executor(None, _analyze_resume_comprehensive, plain, jd)
    if isinstance(result, dict):
        result["extractedText"] = (plain or "")[:25000]
        result["resumeHeader"] = _extract_resume_header(plain or "")
    return JSONResponse(result)


async def api_rewrite_role(request: Request):
    """POST /api/rewrite-role — rewrite a weak role using AI.
    Body: { "header": "Job Title | Company • Dates", "bullets": ["• bullet1", ...] }
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid JSON"}, status_code=400)

    header  = (body.get("header") or "").strip()
    bullets = body.get("bullets") or []
    if not header:
        return JSONResponse({"error": "header required"}, status_code=400)

    bullets_text = "\n".join(bullets) if bullets else "(no bullets provided)"
    prompt = (
        f"You are a professional resume writer. Rewrite the following work experience role "
        f"to be much stronger and more impactful for a tech recruiter.\n\n"
        f"Role: {header}\n"
        f"Current bullets:\n{bullets_text}\n\n"
        f"Instructions:\n"
        f"- Write exactly 3-4 strong bullet points\n"
        f"- Each bullet MUST start with a powerful past-tense action verb "
        f"(e.g., administered, coordinated, authored, negotiated, investigated, engineered, "
        f"programmed, coached, trained, initiated, implemented, monitored, achieved, reduced, spearheaded)\n"
        f"- Each bullet MUST include a quantified result (%, $, time saved, team size, etc.) "
        f"— if the original has no numbers, invent plausible but conservative estimates\n"
        f"- Remove weak verbs (helped, assisted, worked on, was responsible for)\n"
        f"- Remove pronouns (I, my, we)\n"
        f"- Format: return ONLY the bullet points, one per line, each starting with •\n"
        f"- Do not include the role header, just the bullets"
    )

    loop = asyncio.get_event_loop()

    def _call_llm():
        try:
            # Reuse the same LLM routing (Gemini → Grok fallback) as the bullet rewriter
            text = ai_rewrite_bullet("", prompt, "")
            return text.strip() if text else None
        except Exception as exc:
            logger.warning(f"rewrite_role LLM call failed: {exc}")
            return None

    text = await loop.run_in_executor(None, _call_llm)
    if not text:
        return JSONResponse({"error": "LLM unavailable"}, status_code=503)

    # Parse out bullet lines
    rewritten = [l.strip() for l in text.splitlines() if l.strip() and re.match(r"^[•\-–*]", l.strip())]
    if not rewritten:
        rewritten = [l.strip() for l in text.splitlines() if l.strip()]

    return JSONResponse({"bullets": rewritten})


def _resume_coach_prompt(
    candidate_profile: str,
    job_description: str,
    digest: str,
    focus_gaps: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """Shared JD + résumé + optional web-digest prompt for suggest-changes (JSON coach)."""
    digest_block = ""
    ds = (digest or "").strip()
    if ds:
        digest_block = (
            "\n---\nJOB & MARKET CONTEXT (from live web search before this analysis — use for terminology, "
            "JD vocabulary, and honest keyword overlap only; do NOT add employers, degrees, dates, or metrics "
            "not already in the résumé or JD):\n"
            f"{ds[:4500]}\n\n"
        )

    gaps_block = ""
    if focus_gaps:
        lines = []
        for g in focus_gaps[:8]:
            name = str(g.get("name") or "").strip()
            score = g.get("score")
            if not name:
                continue
            score_str = f" ({score}/10)" if isinstance(score, (int, float)) else ""
            lines.append(f"  - {name}{score_str}")
        if lines:
            gaps_block = (
                "\n---\nPRIORITY GAPS FROM PREVIOUS SCORING (these criteria scored low — focus your "
                "suggestions on surfacing any related experience already in the résumé, reframing bullets "
                "to use the exact keywords, or advising the candidate to add missing context they may have "
                "omitted. Do NOT fabricate experience):\n"
                + "\n".join(lines)
                + "\n\n"
            )

    return (
        "You are an expert resume coach. Analyze this resume against the job description "
        "and return 5-8 specific, actionable improvements for individual bullets or sections.\n\n"
        f"RESUME:\n{candidate_profile[:6000]}\n\n"
        f"JOB DESCRIPTION:\n{job_description[:3000]}\n"
        f"{digest_block}"
        f"{gaps_block}"
        "Return a JSON object with this exact structure:\n"
        '{\n'
        '  "summary": "One sentence: the most important gap between this resume and the JD.",\n'
        '  "strategic_tips": [\n'
        '    "2-4 actionable coaching tips (1-2 sentences each) on how to strengthen the candidacy for THIS role — '
        'what to emphasize in interviews, how to reframe experience, or gaps to address in narrative. '
        'NOT resume bullet rewrites; those belong in suggestions."\n'
        '  ],\n'
        '  "interview_questions": [\n'
        '    "5-8 specific questions an interviewer is VERY LIKELY to ask for THIS exact role, based on the JD responsibilities and the candidate\'s gaps. '
        'Include a mix: technical skill questions, behavioural (STAR) questions, and gap-probe questions. '
        'Each question is a complete sentence ending in a question mark."\n'
        '  ],\n'
        '  "suggestions": [\n'
        '    {\n'
        '      "id": "s1",\n'
        '      "section": "Work Experience",\n'
        '      "original": "The exact bullet text from the resume (quote it verbatim)",\n'
        '      "suggested": "The improved version, tailored to the JD keywords",\n'
        '      "reason": "Why this change improves the match — 1 concise sentence.",\n'
        '      "priority": "high",\n'
        '      "category": "quantify_impact"\n'
        '    }\n'
        '  ]\n'
        '}\n\n'
        "Rules:\n"
        "- Only suggest changes to bullets that EXIST in the resume — quote them exactly.\n"
        "- Do NOT invent metrics, employers, or facts not in the resume.\n"
        "- When JOB & MARKET CONTEXT is present, use it only to sharpen **wording** and JD-aligned phrasing for facts already in the résumé.\n"
        "- When PRIORITY GAPS are listed, bias your suggestions toward addressing them — find the closest "
        "existing bullets and rewrite them to surface the relevant skill or experience.\n"
        "- If the résumé summary/objective exists, suggest rewriting it to open with the EXACT job title "
        "from the posting (verbatim, not a synonym) — e.g. if the JD title is 'Software Engineer, Full Stack', "
        "the summary should begin 'Software Engineer, Full Stack with…', not 'Full Stack Engineer with…'.\n"
        "- Priority: 'high' = missing critical JD keyword; 'medium' = wording improvement; 'low' = polish.\n"
        "- category: MUST be exactly one of: 'quantify_impact' (add metrics/numbers/%), "
        "'action_verbs' (stronger/more specific action verbs), "
        "'add_keywords' (surface JD-critical keywords missing from the bullet), "
        "'relevance' (reframe to align with the target role), "
        "'remove_filler' (cut vague/generic language), "
        "'strengthen_impact' (make outcomes and scope clearer). "
        "Pick the SINGLE best-fit category for each suggestion.\n"
        "- strategic_tips: honest, specific, second-person OK; do not repeat the summary sentence; "
        "do not invent employers or metrics.\n"
        "- interview_questions: 5-8 questions an interviewer would realistically ask; mix technical, behavioural, "
        "and gap-probe questions; tailor to THIS specific role and the candidate's gaps; full question sentences.\n"
        "- Return ONLY the JSON object, no markdown fences."
    )


SUGGESTION_CATEGORIES = {
    "quantify_impact", "action_verbs", "add_keywords",
    "relevance", "remove_filler", "strengthen_impact",
}

def _sanitize_suggestions(raw: object) -> List[dict]:
    """Validate and normalise suggestions list — ensures category is always a known value."""
    if not isinstance(raw, list):
        return []
    out: List[dict] = []
    for item in raw[:12]:
        if not isinstance(item, dict):
            continue
        original = str(item.get("original") or "").strip()
        suggested = str(item.get("suggested") or "").strip()
        if not original or not suggested or original == suggested:
            continue
        cat = str(item.get("category") or "").strip().lower()
        if cat not in SUGGESTION_CATEGORIES:
            # fall back: infer from reason text
            reason_lower = str(item.get("reason") or "").lower()
            if any(w in reason_lower for w in ["metric", "number", "quantif", "%", "percent"]):
                cat = "quantify_impact"
            elif any(w in reason_lower for w in ["verb", "action", "strong"]):
                cat = "action_verbs"
            elif any(w in reason_lower for w in ["keyword", "missing", "jd", "ats"]):
                cat = "add_keywords"
            elif any(w in reason_lower for w in ["filler", "vague", "generic", "weak"]):
                cat = "remove_filler"
            elif any(w in reason_lower for w in ["relevance", "align", "reframe", "target"]):
                cat = "relevance"
            else:
                cat = "strengthen_impact"
        priority = str(item.get("priority") or "medium").lower()
        if priority not in ("high", "medium", "low"):
            priority = "medium"
        out.append({
            "id": str(item.get("id") or f"s{len(out)+1}"),
            "section": str(item.get("section") or "Resume").strip(),
            "original": original,
            "suggested": suggested,
            "reason": str(item.get("reason") or "").strip(),
            "priority": priority,
            "category": cat,
        })
    return out


def _sanitize_strategic_tips(raw: object) -> List[str]:
    """Coaching tips shown before PDF generate (not résumé diff rows)."""
    if not isinstance(raw, list):
        return []
    out: List[str] = []
    for item in raw[:6]:
        t = str(item or "").strip()
        if len(t) < 24 or len(t) > 520:
            continue
        out.append(t)
    return out[:4]


def _sanitize_interview_questions(raw: object) -> List[str]:
    """Likely interview questions for this specific role."""
    if not isinstance(raw, list):
        return []
    out: List[str] = []
    for item in raw[:12]:
        q = str(item or "").strip()
        if len(q) < 15 or len(q) > 600:
            continue
        out.append(q)
    return out[:8]


def _sanitize_reuse_research_sources(raw: object) -> List[dict]:
    out: List[dict] = []
    if not isinstance(raw, list):
        return out
    for it in raw[:40]:
        if not isinstance(it, dict):
            continue
        url = str(it.get("url") or "").strip()[:2000]
        if not url:
            continue
        title = str(it.get("title") or "").strip()[:500] or None
        out.append({"title": title, "url": url})
    return out


def _try_suggest_reuse_research(body: dict) -> Optional[Tuple[str, List[str], List[dict]]]:
    """If the client sends a prior digest (same JD session), skip a second live web search."""
    d = str(body.get("reuse_research_digest") or "").strip()
    if len(d) < 40:
        return None
    rq_in = body.get("reuse_research_queries")
    rq = [
        str(q).strip()[:500]
        for q in (rq_in if isinstance(rq_in, list) else [])
        if str(q).strip()
    ][:40]
    rs = _sanitize_reuse_research_sources(body.get("reuse_research_sources"))
    return (d[:5000], rq, rs)


def _parse_focus_gaps(raw: object) -> List[Dict[str, Any]]:
    """Validate and sanitize the focus_gaps list from the request body."""
    if not isinstance(raw, list):
        return []
    out: List[Dict[str, Any]] = []
    for item in raw[:10]:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "").strip()[:120]
        if not name:
            continue
        score = item.get("score")
        try:
            score = float(score) if score is not None else None
        except (TypeError, ValueError):
            score = None
        out.append({"name": name, "score": score})
    return out


async def api_suggest_changes(request: Request):
    """POST /api/suggest-changes — analyze resume vs JD and return per-bullet suggestions.

    Body: { "candidate_profile": str, "job_description": str,
            optional reuse_research_digest, reuse_research_queries, reuse_research_sources }
    Returns: { "summary", "suggestions", optional "research_queries", "research_sources", "research_digest" }
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid json"}, status_code=400)

    candidate_profile = (body.get("candidate_profile") or "").strip()
    job_description   = (body.get("job_description") or "").strip()

    if not candidate_profile:
        return JSONResponse({"error": "candidate_profile required"}, status_code=400)
    if not job_description:
        return JSONResponse({"error": "job_description required"}, status_code=400)

    loop = asyncio.get_event_loop()
    digest = ""
    research_queries: List[str] = []
    research_sources: List[dict] = []
    reuse = _try_suggest_reuse_research(body)
    if reuse:
        digest, research_queries, research_sources = reuse
        logger.info("suggest-changes: reusing client-provided research digest (%s chars)", len(digest))
    else:
        try:
            digest, research_queries, research_sources = await loop.run_in_executor(
                None, run_tailor_research_job_context, job_description
            )
        except Exception as exc:
            logger.warning("pre-suggestion web research failed (suggestions will use resume+JD only): %s", exc)

    focus_gaps = _parse_focus_gaps(body.get("focus_gaps"))
    prompt = _resume_coach_prompt(candidate_profile, job_description, digest, focus_gaps=focus_gaps)

    def _call():
        return coach_suggestions_llm(prompt)

    try:
        text = await loop.run_in_executor(None, _call)
        if text.startswith("```"):
            text = re.sub(r"^```[a-z]*\n?", "", text)
            text = re.sub(r"\n?```$", "", text)
        data = json.loads(text)
        if isinstance(data, dict):
            data["strategic_tips"] = _sanitize_strategic_tips(data.get("strategic_tips"))
            data["interview_questions"] = _sanitize_interview_questions(data.get("interview_questions"))
            data["research_queries"] = research_queries
            data["research_sources"] = research_sources
            if digest.strip():
                data["research_digest"] = digest.strip()[:2500]
        return JSONResponse(data)
    except json.JSONDecodeError as exc:
        logger.error(f"suggest-changes JSON parse error: {exc}  raw={text[:200]}")
        return JSONResponse({"error": "AI response could not be parsed."}, status_code=500)
    except Exception as exc:
        logger.exception("suggest-changes failed")
        return JSONResponse({"error": str(exc)}, status_code=500)


async def api_suggest_gap_fix(request: Request):
    """POST /api/suggest-gap-fix — return 2-3 targeted bullet rewrites for a single gap criterion.

    Much faster than the full suggest-changes call: only looks at one gap and finds the
    best matching bullets to rewrite, without web research or strategic tips.

    Body: {
        "gap_name": str,          # e.g. "Zendesk Administration"
        "gap_notes": str,         # the gap explanation text shown in the UI
        "candidate_profile": str, # plain-text resume
        "job_description": str,
    }
    Returns: {
        "suggestions": [{"id", "original", "suggested", "reason", "section", "priority"}]
    }
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid json"}, status_code=400)

    gap_name          = (body.get("gap_name") or "").strip()
    gap_notes         = (body.get("gap_notes") or "").strip()
    candidate_profile = (body.get("candidate_profile") or "").strip()
    job_description   = (body.get("job_description") or "").strip()

    if not gap_name:
        return JSONResponse({"error": "gap_name required"}, status_code=400)
    if not candidate_profile:
        return JSONResponse({"error": "candidate_profile required"}, status_code=400)
    if not job_description:
        return JSONResponse({"error": "job_description required"}, status_code=400)

    notes_line = f"\nGap detail: {gap_notes}" if gap_notes else ""
    prompt = (
        "You are an expert resume coach. A candidate's résumé was scored and one specific criterion scored low.\n\n"
        f"LOW-SCORING CRITERION: {gap_name}{notes_line}\n\n"
        f"RÉSUMÉ:\n{candidate_profile[:5000]}\n\n"
        f"JOB DESCRIPTION:\n{job_description[:2000]}\n\n"
        "Task: Find the 2-3 bullets in the résumé that are most relevant to this gap and suggest improved "
        "rewrites that better surface the skill or experience. If no bullet directly addresses the gap, "
        "find the closest transferable experience and reframe it.\n\n"
        "Rules:\n"
        "- Only rewrite bullets that EXIST verbatim in the résumé — quote each original exactly.\n"
        "- Do NOT invent metrics, employers, dates, or facts not already in the résumé.\n"
        # A5 — length is flexible upward, never downward
        "- Length: rewrites should be the same length or longer than the original. They may grow\n"
        "  to fit new keywords, but they may NOT shrink by dropping content. One bullet → one\n"
        "  bullet (no splitting), but use as many words as you need to preserve everything.\n"
        "- Focus on vocabulary and framing that matches the job description keywords.\n"
        "- KEYWORD PRESERVATION (CRITICAL): the suggested rewrite MUST preserve every concrete\n"
        "  technical term from the original bullet. Specifically:\n"
        "    * AWS service names (Lambda, Cognito, API Gateway, Bedrock, S3, SQS, DynamoDB, etc.)\n"
        "    * Cloud / DevOps tools (Docker, Kubernetes, Terraform, Jenkins, GitHub Actions, etc.)\n"
        "    * Programming languages and frameworks (Python, Java, Spring Boot, FastAPI, React, etc.)\n"
        "    * Protocols and standards (gRPC, REST, GraphQL, OAuth, SAML, WCAG, etc.)\n"
        "    * Database / data engine names (PostgreSQL, Redis, Kafka, RabbitMQ, etc.)\n"
        "    * Specific product / project nouns (e.g. 'live audio and text', 'CI/CD pipelines',\n"
        "      'infrastructure as code') — do NOT generalize them to vague phrases like 'live data'.\n"
        "  You may ADD new keywords from the JD. You may REPHRASE. You may NOT delete any of the\n"
        "  named technologies, services, protocols, or domain-specific nouns above.\n"
        # A2 — anti-abstraction rule with concrete examples
        "- NO ABSTRACTING (CRITICAL): do not replace a specific term with a generic one. Examples\n"
        "  of changes you must NEVER make:\n"
        "    * 'live audio and text'  →  'live data'                       ❌ (lost the product)\n"
        "    * 'PostgreSQL'            →  'database' or 'a relational store' ❌ (lost the engine)\n"
        "    * 'AWS Lambda'            →  'serverless functions'           ❌ (lost the service)\n"
        "    * 'Kubernetes'            →  'container orchestration'         ❌ (lost the tool)\n"
        "    * 'gRPC streaming'        →  'real-time streaming'             ❌ (lost the protocol)\n"
        "    * 'CI/CD pipelines'       →  'automated deployment'            ❌ (lost the keyword)\n"
        "  The right move is to ADD JD vocabulary alongside the original specifics, not replace.\n"
        "- HONESTY: if the bullet's actual content has no genuine connection to the gap, do NOT\n"
        "  shoehorn the gap's vocabulary in. Skip that bullet and pick a different one, or return\n"
        "  fewer than 3 suggestions. A weak forced bridge is worse than a missing one.\n\n"
        # A3 — few-shot examples (good, bad, skip)
        "Examples of good / bad / skip — STUDY THESE before responding.\n\n"
        "ORIGINAL bullet: 'Built FastAPI service on AWS Lambda that ingests Kafka events into\n"
        "                  PostgreSQL with sub-second latency.'\n"
        "GAP: 'real-time data pipelines'\n"
        "✅ GOOD rewrite: 'Built FastAPI service on AWS Lambda that powers a real-time data\n"
        "                  pipeline, ingesting Kafka events into PostgreSQL with sub-second\n"
        "                  latency.'\n"
        "   Why: every original keyword (FastAPI, AWS Lambda, Kafka, PostgreSQL) is preserved.\n"
        "   'real-time data pipeline' is ADDED. Length grew by one phrase.\n"
        "❌ BAD rewrite:  'Built a real-time data pipeline with sub-second latency on AWS.'\n"
        "   Why: dropped FastAPI, Lambda (now just 'AWS'), Kafka, PostgreSQL — four keywords\n"
        "   gone in exchange for one. Bullet is shorter and weaker. NEVER do this.\n\n"
        "ORIGINAL bullet: 'Designed onboarding flow with progress bars and tooltips for new users.'\n"
        "GAP: 'experience with retail labor systems'\n"
        "⊘ SKIP this bullet. There is no honest bridge from onboarding UI work to retail labor\n"
        "   systems. Forcing a connection ('Designed retail labor onboarding flow…') would be a\n"
        "   lie. Pick a different bullet, or return fewer suggestions.\n\n"
        # A4 — self-critique / verification step
        "Self-check before responding: for EACH suggestion you produce, do these steps mentally:\n"
        "  1. List every capitalized noun, acronym, and CamelCase term in the ORIGINAL bullet.\n"
        "  2. Confirm each of those terms appears verbatim in your SUGGESTED rewrite.\n"
        "  3. If any term from step 1 is missing in step 2, DO NOT emit that suggestion — revise\n"
        "     it until every term is preserved, or drop the suggestion entirely.\n"
        "  4. Confirm your rewrite is at least as long as the original (in words).\n"
        "  5. Confirm at least one phrase from the JOB DESCRIPTION appears in the rewrite (added,\n"
        "     not substituted).\n"
        "Only return a suggestion that passes all five checks.\n\n"
        # Force the category to match what the rewrite actually does — the
        # server-side validator will reject 'quantification' suggestions that
        # add no numerals, and reject any rewrite that drops a numeral or
        # proper noun from the original.
        "CATEGORY TRUTH: pick the category that matches what your rewrite actually does.\n"
        "  - 'quantification' is only valid if your rewrite adds a numeral (digits, %, scale, or\n"
        "    a [X]/[$Y]/[~N] placeholder) that was NOT in the original.\n"
        "  - 'add_keywords' / 'relevance' are valid when you add JD vocabulary alongside the\n"
        "    original specifics.\n"
        "  - 'remove_filler' is the ONLY category that may shrink the bullet, and even then it\n"
        "    must preserve every numeral and named technology from the original.\n"
        "  - 'readability' / 'languageQuality' / 'action_verbs' must NOT delete numerals or\n"
        "    named technologies — they apply to wording and verb choice, not content removal.\n"
        "  The server will reject any suggestion whose rewrite drops a numeral or proper noun\n"
        "  from the original, and any 'quantification' suggestion that adds no new numerals.\n\n"
        "Return ONLY a JSON object:\n"
        '{\n'
        '  "suggestions": [\n'
        '    {\n'
        '      "id": "gf1",\n'
        '      "section": "Work Experience",\n'
        '      "original": "The exact bullet text from the résumé (verbatim)",\n'
        '      "suggested": "The improved bullet text",\n'
        '      "reason": "One sentence explaining how this rewrite addresses the gap.",\n'
        '      "category": "add_keywords",\n'
        '      "priority": "high"\n'
        '    }\n'
        '  ]\n'
        '}\n'
        "Return 2-3 suggestions maximum. Return ONLY the JSON, no markdown fences."
    )

    loop = asyncio.get_event_loop()

    def _call():
        return coach_suggestions_llm(prompt)

    try:
        logger.info("suggest-gap-fix  |  gap=%s  profile_chars=%s", gap_name, len(candidate_profile))
        text = await loop.run_in_executor(None, _call)
        if text.startswith("```"):
            text = re.sub(r"^```[a-z]*\n?", "", text)
            text = re.sub(r"\n?```$", "", text)
        data = json.loads(text)
        suggestions = data.get("suggestions") if isinstance(data, dict) else []
        if not isinstance(suggestions, list):
            suggestions = []
        # Drop any rewrite that silently deletes numerals or proper nouns
        # from the original bullet — those are the "lossy rewrite tagged
        # readability" failure mode we saw repeatedly. Better to return
        # fewer suggestions than a lying one.
        validated: List[dict] = []
        for s in suggestions:
            if not isinstance(s, dict):
                continue
            original = str(s.get("original") or "").strip()
            suggested = str(s.get("suggested") or "").strip()
            if not original or not suggested:
                continue
            cat = str(s.get("category") or "").strip().lower() or None
            ok, why = _validate_rewrite_against_original(original, suggested, category=cat)
            if not ok:
                logger.info(
                    "suggest-gap-fix dropped suggestion (%s): orig=%r  rewrite=%r",
                    why, original[:80], suggested[:80],
                )
                continue
            validated.append(s)
        return JSONResponse({"suggestions": validated})
    except json.JSONDecodeError as exc:
        logger.error("suggest-gap-fix JSON parse error: %s  raw=%s", exc, text[:200])
        return JSONResponse({"error": "AI response could not be parsed."}, status_code=500)
    except Exception as exc:
        logger.exception("suggest-gap-fix failed")
        return JSONResponse({"error": str(exc)}, status_code=500)

async def api_suggest_changes_stream(request: Request):
    """POST /api/suggest-changes-stream — same inputs as suggest-changes; SSE with live coach tokens.

    Body may include reuse_research_digest (+ optional reuse_research_queries / reuse_research_sources)
    to skip a second live JD web search when the job description is unchanged since the last pass.

    Events (JSON per line after ``data: ``):
      ``status`` {msg}
      ``research`` {research_queries, research_sources, research_digest}
      ``coach_delta`` {text} — fragment of the JSON response from the model
      ``coach_done`` {summary, suggestions, research_*} — final structured payload
      ``error`` {msg}
    """
    try:
        body = await request.json()
    except Exception:
        async def err_gen():
            yield {"data": json.dumps({"event": "error", "msg": "invalid json"})}
        return EventSourceResponse(err_gen())

    if not isinstance(body, dict):
        async def err_gen():
            yield {"data": json.dumps({"event": "error", "msg": "Expected a JSON object."})}
        return EventSourceResponse(err_gen())

    candidate_profile = (body.get("candidate_profile") or "").strip()
    job_description = (body.get("job_description") or "").strip()
    if not candidate_profile:
        async def err_gen():
            yield {"data": json.dumps({"event": "error", "msg": "candidate_profile required"})}
        return EventSourceResponse(err_gen())
    if not job_description:
        async def err_gen():
            yield {"data": json.dumps({"event": "error", "msg": "job_description required"})}
        return EventSourceResponse(err_gen())

    loop = asyncio.get_event_loop()
    queue: asyncio.Queue = asyncio.Queue()
    reuse_pack = _try_suggest_reuse_research(body)
    focus_gaps = _parse_focus_gaps(body.get("focus_gaps"))

    def producer():
        try:
            def qput(item: Dict[str, Any]) -> None:
                asyncio.run_coroutine_threadsafe(queue.put(item), loop).result()

            digest = ""
            research_queries: List[str] = []
            research_sources: List[dict] = []
            if reuse_pack:
                digest, research_queries, research_sources = reuse_pack
                logger.info("suggest-changes-stream: reusing client research digest (%s chars)", len(digest))
                qput({"event": "status", "msg": "Reusing web research from your last pass for this job…"})
            else:
                qput({"event": "status", "msg": "Gathering live context from the web…"})
                try:
                    digest, research_queries, research_sources = run_tailor_research_job_context(job_description)
                except Exception as exc:
                    logger.warning("pre-suggestion web research failed (suggestions stream): %s", exc)
            rd = digest.strip()[:2500] if digest.strip() else ""
            qput({
                "event": "research",
                "research_queries": research_queries,
                "research_sources": research_sources,
                "research_digest": rd,
            })
            qput({"event": "status", "msg": "Drafting tailored suggestions…"})
            prompt = _resume_coach_prompt(candidate_profile, job_description, digest, focus_gaps=focus_gaps)

            def on_delta(fragment: str) -> None:
                qput({"event": "coach_delta", "text": fragment})

            text = coach_suggestions_llm_stream(prompt, on_delta=on_delta)
            if text.startswith("```"):
                text = re.sub(r"^```[a-z]*\n?", "", text)
                text = re.sub(r"\n?```$", "", text)
            data = json.loads(text)
            if not isinstance(data, dict):
                raise RuntimeError("Coach returned non-object JSON")
            strategic_tips = _sanitize_strategic_tips(data.get("strategic_tips"))
            interview_questions = _sanitize_interview_questions(data.get("interview_questions"))
            suggestions = _sanitize_suggestions(data.get("suggestions"))
            data["strategic_tips"] = strategic_tips
            data["interview_questions"] = interview_questions
            data["suggestions"] = suggestions
            data["research_queries"] = research_queries
            data["research_sources"] = research_sources
            if digest.strip():
                data["research_digest"] = digest.strip()[:2500]
            qput({
                "event": "coach_done",
                "summary": data.get("summary"),
                "strategic_tips": strategic_tips,
                "interview_questions": interview_questions,
                "suggestions": suggestions,
                "research_queries": data.get("research_queries"),
                "research_sources": data.get("research_sources"),
                "research_digest": data.get("research_digest"),
            })
        except json.JSONDecodeError as exc:
            logger.error("suggest-changes-stream JSON parse error: %s", exc)
            asyncio.run_coroutine_threadsafe(
                queue.put({"event": "error", "msg": "AI response could not be parsed."}),
                loop,
            ).result()
        except Exception as exc:
            logger.exception("suggest-changes-stream failed")
            asyncio.run_coroutine_threadsafe(
                queue.put({"event": "error", "msg": _sse_friendly_error(exc)}),
                loop,
            ).result()
        finally:
            asyncio.run_coroutine_threadsafe(queue.put(None), loop).result()

    threading.Thread(target=producer, daemon=True).start()

    async def event_gen():
        while True:
            item = await queue.get()
            if item is None:
                break
            yield {"data": json.dumps(item)}

    return EventSourceResponse(event_gen())


async def serve_pdf(request: Request):
    folder   = request.path_params["folder"]
    filename = request.path_params["filename"]

    if not filename.endswith(".pdf") or ".." in folder or ".." in filename:
        return JSONResponse({"error": "not found"}, status_code=404)

    pdf_path = os.path.join(LIBRARY_ROOT, folder, filename)
    if not os.path.isfile(pdf_path):
        return JSONResponse({"error": "not found"}, status_code=404)

    logger.info(f"Serving PDF  |  {pdf_path}")
    return FileResponse(pdf_path, media_type="application/pdf")


async def api_rewrite_bullet(request: Request):
    """POST /api/rewrite-bullet — score and optionally rewrite a single bullet.

    Reusable endpoint: works standalone from Analyze inline editing, Resume Builder
    suggestions, or any other UI that needs per-bullet AI feedback.

    Request body (JSON):
      bullet      string  — the bullet text to evaluate (required)
      jd          string  — job description for keyword alignment (optional)
      role        string  — target role title for context (optional)
      company     string  — target company name (optional)
      rewrite     bool    — if true, return an improved version; default true

    Response (JSON):
      {
        "original":    string,
        "score":       int (0–100),
        "issues":      string[],
        "improved":    string | null,   // null when rewrite=false
        "explanation": string
      }
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid JSON"}, status_code=400)

    bullet  = (body.get("bullet") or "").strip()
    jd      = (body.get("jd") or "").strip()
    role    = (body.get("role") or "").strip()
    company = (body.get("company") or "").strip()
    want_rewrite = body.get("rewrite", True)

    if not bullet:
        return JSONResponse({"error": "bullet required"}, status_code=400)

    context_parts: list[str] = []
    if role or company:
        context_parts.append(f"Target role: {role} at {company}".strip(" at"))
    if jd:
        context_parts.append(f"Job description (first 1500 chars):\n{jd[:1500]}")
    context_block = ("\n\n" + "\n\n".join(context_parts)) if context_parts else ""

    rewrite_instruction = (
        '"improved": "A stronger rewrite using an action verb, quantified where possible, '
        'aligned with JD keywords. Do NOT invent metrics not in the original.",'
        if want_rewrite else
        '"improved": null,'
    )

    prompt = f"""You are an expert resume coach. Evaluate the following resume bullet and return ONLY a JSON object — no markdown, no prose.

BULLET:
{bullet}{context_block}

Return this exact JSON schema:
{{
  "score": <integer 0-100 — overall bullet quality>,
  "issues": ["short issue label", ...],
  "explanation": "One sentence: the single most important weakness.",
  {rewrite_instruction}
}}

Scoring rubric:
- 80-100: Strong action verb, specific outcome/metric, relevant to JD
- 60-79: Good verb, decent specificity, minor improvements possible
- 40-59: Weak verb or vague, duty-focused, no metric
- 0-39: Passive voice, responsibilities-only, no impact shown

Issues labels (use only these): "No metric", "Weak verb", "Passive voice", "Too vague",
"Duty-focused", "Too long", "Missing JD keyword", "Starts with date"

Return ONLY the JSON object."""

    raw = _llm_json_call(prompt)
    if not raw or not isinstance(raw, dict):
        return JSONResponse({"error": "LLM unavailable"}, status_code=503)

    score = max(0, min(100, int(raw.get("score") or 50)))
    issues = [str(i) for i in (raw.get("issues") or []) if str(i).strip()][:6]
    improved = str(raw.get("improved") or "").strip() or None if want_rewrite else None
    explanation = str(raw.get("explanation") or "").strip()

    return JSONResponse({
        "original":    bullet,
        "score":       score,
        "issues":      issues,
        "improved":    improved,
        "explanation": explanation,
    })


def _build_docx_bytes_from_structured(
    structured: dict,
    accepted_edits: Optional[dict] = None,
) -> bytes:
    """Render a Word document from a structured resume dict (Analyze / Builder export)."""
    accepted_edits = accepted_edits or {}
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError:
        raise RuntimeError("python-docx is not installed. Add 'python-docx' to requirements.txt.")

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin   = Pt(54)
        section.right_margin  = Pt(54)

    def _h(text: str, size: int = 11, bold: bool = False, color: tuple = (0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(*color)
        return p

    full_name = structured.get("full_name") or "Candidate"
    _h(full_name, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    contact_parts = [v for k in ("email", "phone", "linkedin", "github", "location")
                     if (v := structured.get(k))]
    if contact_parts:
        _h(" | ".join(contact_parts), size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(80, 80, 80))

    if structured.get("summary"):
        doc.add_paragraph()
        _h("SUMMARY", size=9, bold=True, color=(70, 70, 70))
        doc.add_paragraph().add_run(structured["summary"]).font.size = Pt(10)

    skills = structured.get("skills") or []
    if skills:
        doc.add_paragraph()
        _h("SKILLS", size=9, bold=True, color=(70, 70, 70))
        for sk in skills:
            cat   = sk.get("category", "")
            items = sk.get("items") or []
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{cat}: " if cat else "")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(", ".join(items)).font.size = Pt(10)

    experience = structured.get("experience") or []
    if experience:
        doc.add_paragraph()
        _h("EXPERIENCE", size=9, bold=True, color=(70, 70, 70))
        for ei, exp in enumerate(experience):
            header_parts = []
            if exp.get("role"):
                header_parts.append(exp["role"])
            if exp.get("company"):
                header_parts.append(exp["company"])
            role_line = " | ".join(header_parts)
            p = doc.add_paragraph()
            r = p.add_run(role_line)
            r.bold = True
            r.font.size = Pt(10.5)
            if exp.get("dates"):
                p.add_run(f"  {exp['dates']}").font.size = Pt(9)

            bullets = exp.get("bullets") or []
            ei_str  = str(ei)
            for bi, bullet in enumerate(bullets):
                bi_str = str(bi)
                text = accepted_edits.get(ei_str, {}).get(bi_str) or bullet
                p2 = doc.add_paragraph(style="List Bullet")
                p2.add_run(text).font.size = Pt(10)

    for sec in (structured.get("extra_sections") or []):
        title = sec.get("title", "")
        lines = sec.get("lines") or []
        if not lines:
            continue
        doc.add_paragraph()
        _h(title.upper(), size=9, bold=True, color=(70, 70, 70))
        for line in lines:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line).font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_attachment_filename(stem: str, fallback: str = "resume") -> str:
    safe = re.sub(r"[^\w.\-]+", "_", (stem or "").strip()).strip("._")[:80] or fallback
    return safe if safe.lower().endswith(".docx") else f"{safe}.docx"


async def api_export_docx(request: Request):
    """POST /api/export-docx — generate a DOCX from a ResumeDocModel JSON payload.

    Body (JSON):
      structuredResume — dict matching _resume_doc_to_dict() output
      acceptedEdits    — optional {experienceIdx: {bulletIdx: newText}} patch map
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid JSON"}, status_code=400)

    structured = body.get("structuredResume")
    if not structured or not isinstance(structured, dict):
        return JSONResponse({"error": "structuredResume required"}, status_code=400)

    accepted_edits: dict = body.get("acceptedEdits") or {}

    try:
        docx_bytes = await asyncio.get_event_loop().run_in_executor(
            None,
            partial(_build_docx_bytes_from_structured, structured, accepted_edits),
        )
    except RuntimeError as exc:
        return JSONResponse({"error": str(exc)}, status_code=501)
    except Exception as exc:
        logger.exception("export_docx failed")
        return JSONResponse({"error": str(exc)}, status_code=500)

    filename = f"resume-{(structured.get('full_name') or 'export').replace(' ', '_')}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def api_apply_suggestions(request: Request):
    """POST /api/apply-suggestions

    Apply accepted suggestion rewrites to a resume folder WITHOUT an LLM rewrite pass.

    Flow:
      1. Load ResumeDocModel from Supabase resumes table (by folder).
      2. Patch each accepted suggestion via _apply_accepted_edits_to_doc (string replacement).
      3. Re-render via JinjaLatexRenderer → deterministic .tex (no LLM).
      4. Compile with pdflatex.
      5. Upload new PDF to storage.
      6. UPDATE Supabase resumes row: resume_doc + pdf_url + applied_suggestions log.
      7. Return {pdf_url, patches_applied, patches_failed, folder}.
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid JSON"}, status_code=400)

    folder   = str(body.get("folder") or "").strip()
    accepted = body.get("accepted_suggestions") or []
    user_id  = str(body.get("user_id") or "local").strip()
    # Optional: caller can pass resume_doc directly (avoids a Supabase round-trip)
    resume_doc_override = body.get("resume_doc")

    if not folder:
        return JSONResponse({"error": "folder is required"}, status_code=400)
    if not accepted:
        return JSONResponse({"error": "accepted_suggestions is required"}, status_code=400)

    # ── 1. Load ResumeDocModel ────────────────────────────────────────────────
    doc: Optional["ResumeDocModel"] = None
    raw_doc_dict: Optional[dict] = None

    # Try caller-supplied doc first (fastest, no DB round-trip)
    if isinstance(resume_doc_override, dict) and resume_doc_override:
        try:
            raw_doc_dict = resume_doc_override
            doc = _resume_doc_from_parsed(raw_doc_dict)
        except Exception as exc:
            logger.warning(f"apply-suggestions: could not parse caller resume_doc: {exc}")
            doc = None

    # Fall back to Supabase
    if doc is None:
        try:
            supabase = _supabase_table("resumes")
            res = (
                supabase
                .select("resume_doc, reference_folder")
                .eq("folder", folder)
                .limit(1)
                .execute()
            )
            row = (res.data or [None])[0]
            if row and row.get("resume_doc"):
                raw_doc_dict = row["resume_doc"]
                doc = _resume_doc_from_parsed(raw_doc_dict)
        except Exception as exc:
            logger.warning(f"apply-suggestions: Supabase load failed: {exc}")

    if doc is None:
        return JSONResponse(
            {"error": "Could not load resume data for this folder. Re-generate the resume first."},
            status_code=404,
        )

    # ── 2. Patch the doc (pure string replacement, no LLM) ──────────────────
    _apply_accepted_edits_to_doc(doc, accepted)
    raw_doc_dict = _resume_doc_to_dict(doc)

    # Count how many bullets actually matched
    patches_applied = 0
    patches_failed  = 0
    for item in accepted:
        original  = str(item.get("original") or "").strip()
        suggested = str(item.get("suggested") or "").strip()
        if not original or not suggested:
            continue
        # Check if the suggested text appears anywhere in the serialised doc
        doc_text = str(raw_doc_dict)
        if suggested in doc_text or original not in doc_text:
            patches_applied += 1
        else:
            patches_failed += 1

    # ── 3. Re-render via Jinja (deterministic, no LLM) ───────────────────────
    try:
        renderer = JinjaLatexRenderer()
        ref_folder = (row or {}).get("reference_folder") or ""
        template_name = _template_name_for_reference(ref_folder)
        new_tex = renderer.render(doc, template_name=template_name)
    except Exception as exc:
        logger.exception("apply-suggestions: Jinja render failed")
        return JSONResponse({"error": f"LaTeX render failed: {exc}"}, status_code=500)

    # ── 4. Compile ───────────────────────────────────────────────────────────
    try:
        compiled = recompile_resume_from_tex(folder, new_tex)
    except Exception as exc:
        logger.exception("apply-suggestions: pdflatex failed")
        return JSONResponse({"error": f"PDF compilation failed: {exc}"}, status_code=500)

    if not compiled.get("compiled"):
        return JSONResponse(
            {"error": "PDF compilation failed", "details": compiled.get("compile_error", "")},
            status_code=500,
        )

    # ── 5. Upload PDF ────────────────────────────────────────────────────────
    pdf_url: Optional[str] = None
    try:
        pdf_url = upload_pdf(user_id, folder, compiled["pdf_path"])
    except Exception as exc:
        logger.warning(f"apply-suggestions: upload_pdf failed: {exc}")
        # Non-fatal — local PDF still available

    # Fall back to the locally-served PDF route if Supabase upload failed
    # so the frontend always gets a usable URL it can reload the viewer with.
    if not pdf_url and compiled.get("pdf_path") and os.path.isfile(compiled["pdf_path"]):
        pdf_url = f"/pdf/{folder}/{Path(compiled['pdf_path']).name}"

    # Upload updated .tex too
    try:
        upload_tex(user_id, folder, compiled["tex_path"])
    except Exception as exc:
        logger.warning(f"apply-suggestions: upload_tex failed: {exc}")

    # ── 6. Update Supabase ───────────────────────────────────────────────────
    try:
        update_payload: dict = {
            "resume_doc": raw_doc_dict,
            "applied_patch": {
                "suggestions": [
                    {"original": s.get("original"), "suggested": s.get("suggested"), "category": s.get("category")}
                    for s in accepted if s.get("original")
                ],
                "applied_at": __import__("datetime").datetime.utcnow().isoformat() + "Z",
            },
        }
        if pdf_url:
            update_payload["pdf_url"] = pdf_url
        _supabase_table("resumes").update(update_payload).eq("folder", folder).execute()
    except Exception as exc:
        logger.warning(f"apply-suggestions: Supabase update failed: {exc}")

    return JSONResponse({
        "pdf_url":        pdf_url,
        "patches_applied": patches_applied,
        "patches_failed":  patches_failed,
        "folder":         folder,
    })


async def api_builder_export_docx(request: Request):
    """POST /api/builder-export-docx — DOCX for a tailored résumé folder (Builder flow).

    Loads the generated .tex, applies accepted coach suggestions, and returns Word.
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid JSON"}, status_code=400)

    folder = (body.get("folder") or "").strip()
    if not folder or ".." in folder or "/" in folder:
        return JSONResponse({"error": "folder required"}, status_code=400)

    user_id = (body.get("user_id") or "").strip()
    if user_id == "local":
        user_id = ""
    accepted = body.get("accepted_suggestions")
    download_stem = (body.get("download_name") or "").strip()

    tex = get_resume_tex(folder)
    if tex is None and user_id:
        tex = download_tex(user_id, folder)
    if not tex:
        return JSONResponse({"error": "resume not found"}, status_code=404)

    def _build() -> bytes:
        parsed = parse_resume_tex(tex)
        doc = _resume_doc_from_parsed(parsed)
        _apply_accepted_edits_to_doc(doc, accepted if isinstance(accepted, list) else None)
        structured = _resume_doc_to_dict(doc)
        return _build_docx_bytes_from_structured(structured, {})

    try:
        docx_bytes = await asyncio.get_event_loop().run_in_executor(None, _build)
    except RuntimeError as exc:
        return JSONResponse({"error": str(exc)}, status_code=501)
    except Exception as exc:
        logger.exception("builder_export_docx failed")
        return JSONResponse({"error": str(exc)}, status_code=500)

    filename = _docx_attachment_filename(
        download_stem or folder,
        fallback=(folder.split("_")[0] if "_" in folder else folder) or "resume",
    )
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _llm_tailor_to_jd(doc: "ResumeDocModel", jd: str, role: str, company: str) -> "ResumeDocModel":
    """Re-tailor an already-extracted ResumeDocModel to a specific JD.

    Returns the tailored doc, or the original doc unchanged if LLM fails.
    """
    jd_snippet     = (jd or "")[:3000].strip()
    profile_snippet = (jd or "")[:100]  # just for context; real content is from doc

    # Serialize the existing model's content into a compact profile block
    lines: list[str] = []
    if doc.summary:
        lines.append(f"SUMMARY: {doc.summary}")
    for exp in (doc.experience or []):
        lines.append(f"\nROLE: {exp.role} at {exp.company} ({exp.dates})")
        for b in exp.bullets:
            lines.append(f"  - {b}")
    serialized = "\n".join(lines)[:5000]

    prompt = f"""You are an expert resume writer. Rewrite the professional summary and experience bullets
to target this specific role and JD. Keep all dates, company names, and core facts exactly as given.
Strengthen verbs, quantify where data is present, and naturally weave in 1-2 JD keywords per bullet.
Do NOT fabricate metrics.

TARGET ROLE: {role} at {company}

JOB DESCRIPTION:
{jd_snippet}

CURRENT RESUME CONTENT:
{serialized}

Output ONLY valid JSON (no markdown) with this schema:
{{
  "summary": "string — new 2-3 sentence summary tailored to JD",
  "experience": [
    {{
      "company": "string — exact original value",
      "role": "string — exact original value",
      "bullets": ["string — rewritten bullet"]
    }}
  ]
}}"""

    try:
        raw = _llm_json_call(prompt)
        if not raw or not isinstance(raw, dict):
            return doc

        new_summary = _clean_model_text(str(raw.get("summary") or ""))
        if new_summary:
            doc = _resume_doc_with_updates(doc, summary=new_summary)

        exp_list = raw.get("experience") or []
        new_experience = list(doc.experience or [])
        for ei, raw_exp in enumerate(exp_list):
            if ei >= len(new_experience):
                break
            raw_bullets = [_clean_model_text(str(b)) for b in (raw_exp.get("bullets") or []) if _clean_model_text(str(b))]
            if raw_bullets:
                old = new_experience[ei]
                new_experience[ei] = ExperienceItem(
                    company=old.company, role=old.role, dates=old.dates,
                    location=old.location, bullets=raw_bullets,
                )

        return _resume_doc_with_updates(doc, experience=new_experience)
    except Exception as exc:
        logger.warning(f"_llm_tailor_to_jd failed: {exc}")
        return doc


def _doc_from_structured_dict(
    structured: dict,
    accepted_edits: "dict[str, dict[str, str]]",
) -> "ResumeDocModel":
    """Rebuild a ResumeDocModel from a structuredResume dict (Phase 1 format),
    patching in any accepted bullet edits from the Analyze UI.

    accepted_edits format: { "<experienceIdx>": { "<bulletIdx>": "new text" } }
    """
    experience: list[ExperienceItem] = []
    for ei, exp in enumerate(structured.get("experience") or []):
        raw_bullets = list(exp.get("bullets") or [])
        ei_key = str(ei)
        patched = []
        for bi, bullet in enumerate(raw_bullets):
            override = (accepted_edits.get(ei_key) or {}).get(str(bi))
            patched.append(_clean_model_text(override) if override else _clean_model_text(bullet))
        experience.append(ExperienceItem(
            company=_clean_model_text(str(exp.get("company") or "")) or "Company",
            role=_clean_model_text(str(exp.get("role") or "")),
            dates=_clean_model_text(str(exp.get("dates") or "")),
            location=_clean_model_text(str(exp.get("location") or "")),
            bullets=[b for b in patched if b],
        ))

    skills: list[tuple[str, list[str]]] = []
    for sk in (structured.get("skills") or []):
        cat   = _clean_model_text(str(sk.get("category") or "Skills"))
        items = normalize_skill_items(str(i) for i in (sk.get("items") or []))
        if items:
            skills.append((cat or "Skills", items))

    education: list[EducationItem] = []
    for edu in structured.get("education") or []:
        if isinstance(edu, dict):
            row = _education_item_from_dict(edu)
            if row:
                education.append(row)

    projects = _project_items_from_llm_projects(structured.get("projects"))

    extra_sections: list[tuple[str, list[str]]] = []
    for sec in (structured.get("extra_sections") or []):
        title = str(sec.get("title") or "")
        lines = [_clean_model_text(str(l)) for l in (sec.get("lines") or []) if _clean_model_text(str(l))]
        if title and lines:
            extra_sections.append((title, lines))

    return ResumeDocModel(
        full_name=_clean_model_text(str(structured.get("full_name") or "Candidate")) or "Candidate",
        headline=_clean_model_text(str(structured.get("headline") or "")),
        location=_clean_model_text(str(structured.get("location") or "")),
        email=_clean_model_text(str(structured.get("email") or "")),
        phone=_clean_model_text(str(structured.get("phone") or "")),
        linkedin=_clean_model_text(str(structured.get("linkedin") or "")),
        github=_clean_model_text(str(structured.get("github") or "")),
        summary=_clean_model_text(str(structured.get("summary") or "")),
        skills=skills,
        experience=experience,
        education=education,
        projects=projects,
        extra_sections=extra_sections,
    )


async def api_analyze_export_pdf(request: Request):
    """POST /api/analyze-export-pdf — render accepted Analyze edits to a PDF.

    Body (JSON):
      structuredResume  — dict from /api/analyze-upload (required)
      acceptedEdits     — { "<ei>": { "<bi>": "new text" } } (optional)
      referenceFolder   — LaTeX template key, e.g. "Harshibar_Template1" (optional)
      jd                — job description text; if present, runs _llm_tailor_to_jd first (optional)
      role              — job title (optional, used for tailoring + folder naming)
      company           — company name (optional, used for tailoring + folder naming)
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "invalid JSON"}, status_code=400)

    structured = body.get("structuredResume")
    if not structured or not isinstance(structured, dict):
        return JSONResponse({"error": "structuredResume required"}, status_code=400)

    accepted_edits: dict  = body.get("acceptedEdits") or {}
    reference_folder: str = (body.get("referenceFolder") or "Harshibar_Template1").strip()
    jd: str               = (body.get("jd") or "").strip()
    role: str             = (body.get("role") or structured.get("headline") or "Candidate").strip()
    company: str          = (body.get("company") or "").strip()

    loop = asyncio.get_event_loop()

    def _build_pdf() -> tuple[bytes, str]:
        # 1. Rebuild ResumeDocModel from structured dict + accepted edits
        doc = _doc_from_structured_dict(structured, accepted_edits)

        # 2. Optionally tailor to JD
        if jd:
            doc = _llm_tailor_to_jd(doc, jd, role, company)

        # 3. Render LaTeX
        renderer = JinjaLatexRenderer()
        template_name = _template_name_for_reference(reference_folder)
        new_tex = renderer.render(doc, template_name=template_name)

        # 4. Compile to PDF
        out_folder, _ = _create_structured_output_folder(None, reference_folder, role, company or "export")
        compiled = recompile_resume_from_tex(out_folder, new_tex)

        pdf_path = compiled.get("pdf_path")
        if not pdf_path or not os.path.isfile(pdf_path):
            raise RuntimeError("PDF compilation failed — check pdflatex is installed.")

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        safe_name = (doc.full_name or "resume").replace(" ", "_")
        filename  = f"{safe_name}_tailored.pdf"
        return pdf_bytes, filename

    try:
        pdf_bytes, filename = await loop.run_in_executor(None, _build_pdf)
    except Exception as exc:
        logger.exception("analyze_export_pdf failed")
        return JSONResponse({"error": str(exc)}, status_code=500)

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def api_export_pdf_html(request: Request) -> Response:
    """Render arbitrary HTML to PDF via Playwright/Chromium.

    Body: { html: string, filename?: string }
    Returns: PDF bytes (application/pdf)
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "Invalid JSON body"}, status_code=400)

    html_content = body.get("html", "")
    filename = body.get("filename", "resume.pdf")

    if not html_content:
        return JSONResponse({"error": "html field is required"}, status_code=400)

    import asyncio
    from concurrent.futures import ThreadPoolExecutor
    from playwright.sync_api import sync_playwright

    def render_pdf():
        with sync_playwright() as p:
            browser = p.chromium.launch(
                args=["--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu"]
            )
            page = browser.new_page()
            page.set_content(html_content, wait_until="networkidle")
            pdf_bytes = page.pdf(
                format="Letter",
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
                print_background=True,
            )
            browser.close()
            return pdf_bytes

    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor(max_workers=1) as executor:
        pdf_bytes = await loop.run_in_executor(executor, render_pdf)

    safe_filename = filename.replace('"', '').replace('\\', '') or "resume.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_filename}"'},
    )


routes = [
    Route("/",                              homepage),
    Route("/api/health",                    api_health,          methods=["GET"]),
    Route("/api/resumes",                   api_resumes),
    Route("/api/generate-stream",           api_generate_stream, methods=["POST"]),
    Route("/api/upload-resume",             api_upload_resume,   methods=["POST"]),
    Route("/api/extract-jd",              api_extract_jd,     methods=["POST"]),
    Route("/api/resume/{folder}",          api_resume_parsed,  methods=["GET"]),
    Route("/api/resume/{folder}",          api_resume_save,    methods=["POST"]),
    Route("/api/ai-edit-bullet",           api_ai_edit_bullet,methods=["POST"]),
    Route("/api/generate-skills",          api_generate_skills, methods=["POST"]),
    Route("/api/ats-check/{folder}",     api_ats_check,     methods=["POST"]),
    Route("/api/doctor-check",             api_doctor_check,   methods=["POST"]),
    Route("/api/resume-analysis/{folder}", api_resume_analysis, methods=["POST"]),
    Route("/api/share/{folder}",           api_share_create,  methods=["POST"]),
    Route("/api/share/{shortid}",         api_share_resolve, methods=["GET"]),
    Route("/api/share/{shortid}",         api_share_revoke, methods=["DELETE"]),
    Route("/api/version/{folder}",        api_version_save, methods=["POST"]),
    Route("/api/version/{folder}",        api_version_list, methods=["GET"]),
    Route("/api/version/{folder}/{version}", api_version_load, methods=["GET"]),
    Route("/api/storage-status",            api_storage_status,methods=["GET"]),
    Route("/api/backfill-tex",              api_backfill_tex,  methods=["POST"]),
    Route("/api/analyze-upload",           api_analyze_upload,  methods=["POST"]),
    Route("/api/my-analyses",             api_my_analyses,     methods=["GET"]),
    Route("/api/cohort-stats",            api_cohort_stats,    methods=["GET"]),
    Route("/api/student-detail",          api_student_detail,  methods=["GET"]),
    Route("/api/export-docx",             api_export_docx,     methods=["POST"]),
    Route("/api/apply-suggestions",        api_apply_suggestions,   methods=["POST"]),
    Route("/api/builder-export-docx",     api_builder_export_docx, methods=["POST"]),
    Route("/api/analyze-export-pdf",      api_analyze_export_pdf, methods=["POST"]),
    Route("/api/analyze-folder/{folder}", api_analyze_folder,  methods=["POST"]),
    Route("/api/rewrite-role",            api_rewrite_role,    methods=["POST"]),
    Route("/api/rewrite-bullet",          api_rewrite_bullet,  methods=["POST"]),
    Route("/api/suggest-changes",         api_suggest_changes, methods=["POST"]),
    Route("/api/suggest-changes-stream",  api_suggest_changes_stream, methods=["POST"]),
    Route("/api/suggest-gap-fix",         api_suggest_gap_fix, methods=["POST"]),
    Route("/api/export-pdf-html",         api_export_pdf_html,     methods=["POST"]),
    Route("/pdf/{folder}/{filename}",      serve_pdf),
]

middleware = [
    Middleware(
        CORSMiddleware,
        allow_origins=ALLOWED_ORIGINS,
        # Allow any GitHub Pages domain + any resunova.io subdomain
        allow_origin_regex=r"https://(.*\.github\.io|(.*\.)?resunova\.io)",
        allow_methods=["GET", "POST", "OPTIONS", "DELETE"],
        allow_headers=["*"],
        allow_credentials=False,
    ),
]

app = Starlette(routes=routes, middleware=middleware)

if __name__ == "__main__":
    host = "0.0.0.0" if os.environ.get("RAILWAY_ENVIRONMENT") else "127.0.0.1"
    logger.info(f"Resume Generator starting on http://{host}:{PORT}")
    uvicorn.run(app, host=host, port=PORT, log_level="info")
