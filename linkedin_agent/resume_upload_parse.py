"""
Resume upload: document → markdown (MarkItDown), then LLM → structured JSON (Pydantic).

Mirrors the Resume-Matcher pattern (extraction vs understanding) using this repo's
Grok/Gemini stack from ``resume_library``.
"""

from __future__ import annotations

import io
import json
import logging
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

from pydantic import BaseModel, Field, field_validator

logger = logging.getLogger(__name__)

# Align with Resume Matcher user-mistake guardrails (docs/user-mistake-guards.md).
RESUME_UPLOAD_MAX_BYTES = 4 * 1024 * 1024

ALLOWED_UPLOAD_SUFFIXES = frozenset({".pdf", ".doc", ".docx"})
EXPECTED_MIME_BY_SUFFIX: Dict[str, frozenset[str]] = {
    ".pdf": frozenset({"application/pdf"}),
    ".doc": frozenset({"application/msword"}),
    ".docx": frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    ),
}


def validate_resume_upload_file(content_type: Optional[str], filename: str) -> None:
    """
    Validate extension + browser-reported MIME. Raises ValueError with a safe user message (400).
    ``application/octet-stream`` or missing type skips MIME check (many clients omit a precise type).
    """
    name = (filename or "").strip() or "resume"
    suf = Path(name).suffix.lower()
    if suf not in ALLOWED_UPLOAD_SUFFIXES:
        raise ValueError(
            f"Unsupported file type ({suf or 'no extension'}). "
            "Allowed: PDF, Microsoft Word .doc, or .docx."
        )

    ct = (content_type or "").split(";")[0].strip().lower()
    if not ct or ct == "application/octet-stream":
        return

    if ct.startswith(("image/", "video/", "audio/")):
        raise ValueError(
            "The upload looks like an image or media file, not a résumé document. "
            "Export or save as PDF or Word (.doc / .docx) from your editor — do not rename image files."
        )
    if ct.startswith("text/") or ct in ("application/json", "application/xml", "text/xml"):
        raise ValueError(
            "The upload looks like plain text or code, not a PDF or Word file. "
            "Please upload an actual .pdf or .doc / .docx export."
        )

    expected = EXPECTED_MIME_BY_SUFFIX.get(suf, frozenset())
    if ct in expected:
        return

    # Legacy / uncommon PDF MIME seen in the wild
    if suf == ".pdf" and ct in ("application/x-pdf", "binary/octet-stream"):
        return

    if suf == ".pdf":
        raise ValueError(
            "The file does not look like a real PDF (the browser reported a different format). "
            "If you renamed a non-PDF to .pdf, use Save as PDF or Print to PDF instead."
        )
    if suf == ".docx":
        if ct == "application/msword":
            raise ValueError(
                "The file is named .docx but was reported as an older .doc type. "
                "Re-save as .docx in Word, or rename to .doc if it is actually Word 97–2003 format."
            )
        raise ValueError(
            "The file does not look like a Word .docx (the browser reported a different format). "
            "Export as .docx from Word or Google Docs."
        )
    if suf == ".doc" and ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        raise ValueError(
            "The file is named .doc but was reported as .docx format. "
            "Rename to .docx or re-save with the correct extension."
        )
    raise ValueError(
        "The file’s reported type does not match the extension. "
        "Upload a genuine PDF or Word export (avoid renaming unrelated files)."
    )


@dataclass(frozen=True)
class UploadMarkdownOutcome:
    """Result of binary → text extraction for résumé upload."""

    markdown: str
    """Set when ``markdown`` is empty; drives HTTP 422 messaging."""
    empty_reason: Optional[str] = None


def message_for_empty_resume_extract(empty_reason: Optional[str]) -> str:
    if empty_reason == "corrupt_or_unreadable_pdf":
        return (
            "Could not read this PDF — it may be corrupted, encrypted, or not a valid PDF. "
            "Try opening it in a PDF viewer; if it fails, re-export from the original document."
        )
    if empty_reason == "corrupt_or_unreadable_word":
        return (
            "Could not read this Word file — it may be corrupted, password-protected, or not a valid .doc/.docx. "
            "Try opening it in Word; re-save a copy without encryption."
        )
    if empty_reason == "pdf_no_text":
        return (
            "No text could be extracted from this PDF. It is often a scanned (image-only) résumé. "
            "Upload a PDF with selectable text, or use a .docx export, or run OCR first."
        )
    if empty_reason == "word_no_text":
        return (
            "No text could be extracted from this Word document. "
            "The file may be empty, heavily restricted, or not a real Word export."
        )
    return (
        "Could not extract readable text from the file. "
        "Try another format (PDF with selectable text or .docx) or a different export."
    )


def extract_upload_markdown(
    content: bytes,
    filename: str,
    pdf_plain_fallback: Optional[str] = None,
) -> UploadMarkdownOutcome:
    """
    PDF or Word → markdown/plain text via MarkItDown; PDF falls back to pdfplumber
    when MarkItDown returns empty. Records ``empty_reason`` when output is blank.
    """
    suffix = Path(filename or "resume.pdf").suffix.lower() or ".pdf"
    if suffix not in ALLOWED_UPLOAD_SUFFIXES:
        suffix = ".pdf"

    markitdown_failed = False
    tmp_path: Optional[Path] = None
    try:
        from markitdown import MarkItDown  # type: ignore[import-untyped]

        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)

        md = MarkItDown()
        result = md.convert(str(tmp_path))
        text = (getattr(result, "text_content", None) or "").strip()
    except Exception as exc:
        logger.warning("MarkItDown convert failed (%s): %s", filename, exc)
        markitdown_failed = True
        text = ""
    finally:
        if tmp_path is not None:
            tmp_path.unlink(missing_ok=True)

    if text:
        return UploadMarkdownOutcome(text, None)

    if suffix == ".pdf":
        fb = (pdf_plain_fallback or "").strip() or _extract_pdf_text_pdfplumber(content)
        if fb.strip():
            logger.info("resume_upload_parse: MarkItDown empty; used pdfplumber fallback for %s", filename)
            return UploadMarkdownOutcome(fb.strip(), None)
        if markitdown_failed:
            return UploadMarkdownOutcome("", "corrupt_or_unreadable_pdf")
        return UploadMarkdownOutcome("", "pdf_no_text")

    if markitdown_failed:
        return UploadMarkdownOutcome("", "corrupt_or_unreadable_word")
    return UploadMarkdownOutcome("", "word_no_text")

# ── Date repair (year-only → month-inclusive from markdown) ─────────────────

_MD_DATE_RE = re.compile(
    r"(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?"
    r"|Dec(?:ember)?)"
    r"\.?\s+\d{4})"
    r"(?:\s*[-–—]\s*"
    r"(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?"
    r"|Dec(?:ember)?)"
    r"\.?\s+\d{4}"
    r"|Present|Current|Now|Ongoing))?",
    re.IGNORECASE,
)


def _extract_markdown_dates(markdown: str) -> List[str]:
    return _MD_DATE_RE.findall(markdown)


def restore_dates_from_markdown(
    parsed_data: Dict[str, Any],
    markdown: str,
) -> Dict[str, Any]:
    """Patch year-only ``years`` fields using month-inclusive spans found in markdown."""
    md_dates = _extract_markdown_dates(markdown)
    if not md_dates:
        return parsed_data

    year_only_re = re.compile(r"\d{4}")
    year_to_full: Dict[str, str] = {}
    for md_date in md_dates:
        years_in_date = year_only_re.findall(md_date)
        if years_in_date:
            year_key = " - ".join(years_in_date)
            if year_key not in year_to_full:
                normalized = re.sub(r"\s*[-–—]\s*", " - ", md_date.strip())
                year_to_full[year_key] = normalized

    if not year_to_full:
        return parsed_data

    month_re = re.compile(
        r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)",
        re.IGNORECASE,
    )

    def patch_years(years_val: str) -> Optional[str]:
        if not isinstance(years_val, str) or not years_val.strip():
            return None
        if month_re.search(years_val):
            return None
        key = years_val.strip()
        if key in year_to_full:
            return year_to_full[key]
        return None

    patched = 0
    for key in ("work_experience", "education", "projects"):
        for entry in parsed_data.get(key) or []:
            if not isinstance(entry, dict):
                continue
            new_y = patch_years(entry.get("years") or "")
            if new_y:
                entry["years"] = new_y
                patched += 1

    if patched:
        logger.info("resume_upload_parse: restored months in %d date fields from markdown", patched)

    return parsed_data


# ── Pydantic schema (strict enough to validate; extra keys ignored) ───────────


class WorkExperienceItem(BaseModel):
    title: str = ""
    organization: str = ""
    location: str = ""
    years: str = ""
    bullets: List[str] = Field(default_factory=list)

    @field_validator("bullets", mode="before")
    @classmethod
    def _bullets(cls, v: Any) -> List[str]:
        if v is None:
            return []
        if isinstance(v, str) and v.strip():
            return [v.strip()]
        if isinstance(v, list):
            return [str(x).strip() for x in v if str(x).strip()]
        return []


class EducationItem(BaseModel):
    school: str = ""
    degree: str = ""
    years: str = ""
    details: str = ""


class ProjectItem(BaseModel):
    name: str = ""
    description: str = ""
    years: str = ""
    bullets: List[str] = Field(default_factory=list)

    @field_validator("bullets", mode="before")
    @classmethod
    def _bullets(cls, v: Any) -> List[str]:
        if v is None:
            return []
        if isinstance(v, str) and v.strip():
            return [v.strip()]
        if isinstance(v, list):
            return [str(x).strip() for x in v if str(x).strip()]
        return []


class UploadedResumeStructured(BaseModel):
    model_config = {"extra": "ignore"}

    full_name: str = ""
    headline: str = ""
    email: str = ""
    phone: str = ""
    linkedin: str = ""
    github: str = ""
    website: str = ""
    summary: str = ""
    skills: List[str] = Field(default_factory=list)
    work_experience: List[WorkExperienceItem] = Field(default_factory=list)
    education: List[EducationItem] = Field(default_factory=list)
    projects: List[ProjectItem] = Field(default_factory=list)

    @field_validator("skills", mode="before")
    @classmethod
    def _skills(cls, v: Any) -> List[str]:
        if v is None:
            return []
        if isinstance(v, str):
            parts = re.split(r"[,;|]\s*|\n+", v)
            return [p.strip() for p in parts if p.strip()]
        if isinstance(v, list):
            return [str(x).strip() for x in v if str(x).strip()]
        return []


RESUME_JSON_SCHEMA_HINT = """{
  "full_name": "string",
  "headline": "string",
  "email": "string",
  "phone": "string",
  "linkedin": "string (URL or path)",
  "github": "string (URL or path)",
  "website": "string",
  "summary": "string — 2-5 sentences professional summary if present else empty",
  "skills": ["skill1", "skill2"],
  "work_experience": [
    {
      "title": "Job title",
      "organization": "Company",
      "location": "City, ST or Remote",
      "years": "Jan 2020 - Present",
      "bullets": ["Achievement with metric", "…"]
    }
  ],
  "education": [
    {"school": "…", "degree": "BS Computer Science", "years": "2016 - 2020", "details": "Honors, GPA if worth keeping"}
  ],
  "projects": [
    {"name": "…", "description": "one line", "years": "", "bullets": ["…"]}
  ]
}"""


def _build_parse_prompt(markdown: str) -> str:
    clipped = markdown.strip()
    if len(clipped) > 52000:
        clipped = clipped[:52000] + "\n\n[… document truncated for parsing …]\n"
    return (
        "You extract structured résumé data from the following document text (Markdown or plain).\n"
        "Rules:\n"
        "- Copy facts faithfully; do not invent employers, degrees, dates, or metrics.\n"
        "- If a field is unknown, use empty string or empty array.\n"
        "- Preserve employment date ranges with months when shown in the source.\n"
        "- Split skills into individual array entries (not one long comma string).\n"
        "- work_experience bullets: action-led lines; keep numbers/units from source.\n\n"
        "Return ONLY valid JSON matching this shape (keys required; use defaults for missing sections):\n"
        f"{RESUME_JSON_SCHEMA_HINT}\n\n"
        "DOCUMENT:\n"
        f"{clipped}"
    )


def structured_to_plain_resume_text(data: UploadedResumeStructured) -> str:
    """Render structured résumé as plain text for the builder textarea (pipe-friendly lines)."""
    lines: List[str] = []
    name = (data.full_name or "").strip()
    if name:
        lines.append(name)
    head = (data.headline or "").strip()
    if head:
        lines.append(head)
    lines.append("")

    meta_bits = []
    for val in (
        (data.email or "").strip(),
        (data.phone or "").strip(),
        (data.linkedin or "").strip(),
        (data.github or "").strip(),
        (data.website or "").strip(),
    ):
        if val:
            meta_bits.append(val)
    if meta_bits:
        lines.append("  |  ".join(meta_bits))
        lines.append("")

    summ = (data.summary or "").strip()
    if summ:
        lines.append("SUMMARY")
        lines.append(summ)
        lines.append("")

    skills = [s.strip() for s in data.skills if s and str(s).strip()]
    if skills:
        lines.append("SKILLS")
        lines.append(", ".join(skills))
        lines.append("")

    if data.work_experience:
        lines.append("EXPERIENCE")
        for w in data.work_experience:
            title = (w.title or "").strip()
            org = (w.organization or "").strip()
            loc = (w.location or "").strip()
            yrs = (w.years or "").strip()
            pipe = " | ".join(p for p in (title, org, loc, yrs) if p)
            if pipe:
                lines.append(pipe)
            for b in w.bullets:
                bt = (b or "").strip()
                if bt:
                    lines.append(f"• {bt}")
            lines.append("")

    if data.education:
        lines.append("EDUCATION")
        for e in data.education:
            school = (e.school or "").strip()
            deg = (e.degree or "").strip()
            yrs = (e.years or "").strip()
            det = (e.details or "").strip()
            pipe = " | ".join(p for p in (school, deg, yrs) if p)
            if pipe:
                lines.append(pipe)
            if det:
                lines.append(f"• {det}")
            lines.append("")

    if data.projects:
        lines.append("PROJECTS")
        for p in data.projects:
            nm = (p.name or "").strip()
            desc = (p.description or "").strip()
            yrs = (p.years or "").strip()
            head_line = " | ".join(x for x in (nm, desc, yrs) if x)
            if head_line:
                lines.append(head_line)
            for b in p.bullets:
                bt = (b or "").strip()
                if bt:
                    lines.append(f"• {bt}")
            lines.append("")

    return "\n".join(lines).strip()


def _post_clean_resume_text(text: str) -> str:
    """Normalize PDF placeholder glyphs; keep line breaks (same semantics as resume_gui.app)."""
    text = re.sub(r"\(\s*cid\s*:\s*\d+\)", " • ", text, flags=re.IGNORECASE)
    text = re.sub(r"\bUsed\s*:\s*to\b", "Used to", text, flags=re.IGNORECASE)
    lines = [" ".join(ln.split()) for ln in text.splitlines()]
    return "\n".join(lines).strip()


def _merge_split_word_tokens(tokens: List[str]) -> List[str]:
    """Fuse a lone uppercase letter with the following lowercase token (pdfplumber quirk)."""
    skip_single = frozenset({"I", "A"})
    out: List[str] = []
    i = 0
    while i < len(tokens):
        t = tokens[i]
        if (
            len(t) == 1
            and t.isalpha()
            and t.isupper()
            and t not in skip_single
            and i + 1 < len(tokens)
        ):
            nxt = tokens[i + 1]
            if nxt and len(nxt) >= 2 and nxt[0].islower():
                out.append(t + nxt)
                i += 2
                continue
        out.append(t)
        i += 1
    return out


def _extract_pdf_text_pdfplumber(content: bytes) -> str:
    """Word-spacing-aware PDF extract when MarkItDown returns empty."""
    import pdfplumber

    pages_text: List[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            words = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=False)
            if not words:
                pages_text.append(page.extract_text() or "")
                continue
            line_map: Dict[int, List[str]] = {}
            for w in words:
                y_key = round(float(w["top"]) / 4) * 4
                line_map.setdefault(y_key, []).append(w["text"])
            page_lines = [
                " ".join(_merge_split_word_tokens(tokens))
                for _, tokens in sorted(line_map.items())
            ]
            pages_text.append("\n".join(page_lines))
    return _post_clean_resume_text("\n".join(pages_text))


def markdown_from_upload_bytes(
    content: bytes,
    filename: str,
    pdf_plain_fallback: Optional[str] = None,
) -> str:
    """Backward-compatible: returns markdown/plain text only (see ``extract_upload_markdown``)."""
    return extract_upload_markdown(content, filename, pdf_plain_fallback).markdown


def llm_resume_markdown_to_structured(markdown: str) -> Dict[str, Any]:
    """
    Call configured LLM(s) to produce structured JSON, validate, apply date repair.

    Raises RuntimeError if no provider succeeds.
    """
    from google.genai import types

    from resume_library import (
        _GROK_FALLBACK_MODELS,
        _backoff_if_rate_limited,
        _is_grok,
        _json_grok,
        _model_chain,
        _optional_gemini_client,
        _transient_provider_error,
        grok_preferred_for_throughput,
        primary_gemini_flash_model,
        primary_llm_model_for_resume_workloads,
    )

    prompt = _build_parse_prompt(markdown)

    if grok_preferred_for_throughput():
        seen: set[str] = set()
        chain = []
        prim = primary_llm_model_for_resume_workloads()
        for m in (prim,) + _GROK_FALLBACK_MODELS + tuple(_model_chain(primary_gemini_flash_model())):
            if m and m not in seen:
                seen.add(m)
                chain.append(m)
    else:
        chain = _model_chain(primary_gemini_flash_model())

    api_key = (os.environ.get("GOOGLE_API_KEY") or os.environ.get("GEMINI_API_KEY") or "").strip()
    client: Any = _optional_gemini_client() if api_key else None

    last_exc: Optional[BaseException] = None

    for model in chain:
        try:
            if _is_grok(model):
                parsed = _json_grok(model, prompt, temperature=0.1)
                if not parsed or not isinstance(parsed, dict):
                    continue
            else:
                if client is None:
                    continue
                cfg = types.GenerateContentConfig(
                    temperature=0.1,
                    response_mime_type="application/json",
                )
                r = client.models.generate_content(
                    model=model,
                    contents=prompt,
                    config=cfg,
                )
                raw = (getattr(r, "text", None) or "").strip()
                if not raw:
                    continue
                parsed = json.loads(raw)
                if not isinstance(parsed, dict):
                    continue

            patched = restore_dates_from_markdown(parsed, markdown)
            validated = UploadedResumeStructured.model_validate(patched)
            logger.info("resume_upload_parse: LLM structured parse succeeded on model %s", model)
            return validated.model_dump()

        except Exception as exc:
            last_exc = exc
            logger.warning("resume_upload_parse: model %s failed: %s", model, exc)
            if _transient_provider_error(exc):
                _backoff_if_rate_limited(exc)
            continue

    if last_exc:
        raise RuntimeError(f"Structured resume parse failed: {last_exc}") from last_exc
    raise RuntimeError(
        "Structured resume parse failed: no LLM response (check GOOGLE_API_KEY / GEMINI_API_KEY or XAI_API_KEY)."
    )


# ── Deterministic section / date helpers ────────────────────────────────────

_SECTION_MAP: Dict[str, tuple[str, ...]] = {
    "experience": (
        "experience", "work history", "employment", "professional background",
        "career history", "work experience", "professional experience",
        "relevant experience", "employment history",
    ),
    "education": (
        "education", "academic", "qualification", "academics", "schooling",
        "educational background", "academic background",
    ),
    "skills": (
        "skills", "technical skills", "core competencies", "competencies",
        "technologies", "expertise", "proficiencies", "highlights",
        "qualifications", "technical qualifications", "core skills",
        "key skills", "areas of expertise",
    ),
    "summary": (
        "summary", "profile", "objective", "about", "professional summary",
        "career objective", "overview", "career overview", "career focus",
        "career summary", "professional profile", "executive summary",
        "personal statement", "personal profile",
    ),
    "projects": ("project", "portfolio", "key projects"),
    "certifications": (
        "certification", "certificate", "license", "credential",
        "training", "courses", "professional development",
    ),
}

_EXTRACT_DATE_RE = re.compile(
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s+\d{4}"
    r"(?:\s*[-–—]\s*"
    r"(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s+\d{4}|Present|Current|Now|Ongoing))?",
    re.IGNORECASE,
)
_YEAR_RANGE_RE = re.compile(
    r"\b((?:19|20)\d{2})\s*[-–—]\s*((?:19|20)\d{2}|Present|Current|Now)\b",
    re.IGNORECASE,
)


def _det_canonical_section(text: str) -> Optional[str]:
    low = text.lower().strip()
    for sec, kws in _SECTION_MAP.items():
        if any(k in low for k in kws):
            return sec
    return None


def _det_extract_date(text: str) -> str:
    m = _EXTRACT_DATE_RE.search(text)
    if m:
        return re.sub(r"\s*[-–—]\s*", " - ", m.group()).strip()
    m = _YEAR_RANGE_RE.search(text)
    if m:
        return f"{m.group(1)} - {m.group(2)}"
    return ""


def _det_extract_contact(lines: List[str]) -> Dict[str, str]:
    block = "\n".join(lines[:15])
    contact: Dict[str, str] = {}
    em = re.search(r"[\w.+%-]+@[\w.-]+\.\w{2,}", block)
    ph = re.search(r"\+?1?[\s.-]?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}", block)
    li = re.search(r"(?:linkedin\.com/in/|linkedin\.com/)[\w%-]+(?:/[\w%-]+)*", block, re.I)
    gh = re.search(r"github\.com/[\w-]+", block, re.I)
    url = re.search(r"https?://(?!(?:www\.)?(?:linkedin|github)\.com)\S+", block, re.I)
    if em:  contact["email"]    = em.group().lower()
    if ph:  contact["phone"]    = re.sub(r"\s+", " ", ph.group()).strip()
    if li:  contact["linkedin"] = li.group().rstrip(".,)")
    if gh:  contact["github"]   = gh.group().rstrip(".,)")
    if url: contact["website"]  = url.group().rstrip(".,)")
    return contact


def _det_quality_ok(result: "UploadedResumeStructured") -> bool:
    """Return True if deterministic extraction found enough to be useful."""
    has_identity = bool(result.email) or bool(result.full_name) or bool(result.phone)
    has_content = bool(
        result.work_experience or result.education or result.skills or result.summary
    )
    return has_identity or has_content


# ── PyMuPDF PDF extractor ───────────────────────────────────────────────────

def _extract_pdf_structured(content: bytes) -> Optional["UploadedResumeStructured"]:
    """Deterministic PDF → UploadedResumeStructured via PyMuPDF span metadata."""
    try:
        import fitz  # type: ignore[import-untyped]
    except ImportError:
        logger.warning("_extract_pdf_structured: PyMuPDF not installed")
        return None

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        logger.warning("_extract_pdf_structured: fitz.open failed: %s", exc)
        return None

    # ── Reconstruct lines from span data ────────────────────────────────────
    raw_lines: List[Dict[str, Any]] = []
    for page in doc:
        spans: List[Dict[str, Any]] = []
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    t = span["text"].strip()
                    if t:
                        spans.append({
                            "text": t,
                            "y": round(span["origin"][1] / 4) * 4,
                            "x": span["origin"][0],
                            "size": span["size"],
                            "bold": bool(span["flags"] & 16),
                        })
        line_map: Dict[int, List[Dict]] = {}
        for s in spans:
            line_map.setdefault(s["y"], []).append(s)
        for y in sorted(line_map):
            row = sorted(line_map[y], key=lambda s: s["x"])
            text = " ".join(s["text"] for s in row).strip()
            sizes = [s["size"] for s in row]
            raw_lines.append({
                "text": text,
                "size": sum(sizes) / len(sizes),
                "bold": all(s["bold"] for s in row),
                "all_caps": text == text.upper() and any(c.isalpha() for c in text),
            })
    doc.close()

    if not raw_lines:
        return None

    # Detect flat-format PDF (no typography variation)
    from collections import Counter
    size_counts = Counter(round(ln["size"]) for ln in raw_lines)
    body_size = size_counts.most_common(1)[0][0]
    bold_ratio = sum(1 for ln in raw_lines if ln["bold"]) / len(raw_lines)
    is_flat = len(size_counts) <= 2 and bold_ratio < 0.05

    def is_heading(ln: Dict[str, Any]) -> bool:
        text = ln["text"]
        if not text or len(text) > 70:
            return False
        if is_flat:
            return len(text.split()) <= 5 and _det_canonical_section(text) is not None
        if ln["all_caps"] and len(text) <= 60 and any(c.isalpha() for c in text):
            return True
        if ln["bold"] and len(text) <= 50:
            return True
        if ln["size"] > body_size + 1 and len(text) <= 60:
            return True
        return False

    plain_lines = [ln["text"] for ln in raw_lines]
    return _build_structured_from_lines(plain_lines, is_heading_fn=lambda i: is_heading(raw_lines[i]))


# ── python-docx DOCX extractor ──────────────────────────────────────────────

def _extract_docx_structured(content: bytes) -> Optional["UploadedResumeStructured"]:
    """Deterministic DOCX → UploadedResumeStructured via python-docx paragraph styles.

    Uses full element.body traversal so table-only resumes (where doc.paragraphs
    returns nothing) are handled correctly.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.oxml.ns import qn  # type: ignore[import-untyped]
        from docx.text.paragraph import Paragraph as DocxParagraph  # type: ignore[import-untyped]
        from docx.table import Table as DocxTable  # type: ignore[import-untyped]
    except ImportError:
        logger.warning("_extract_docx_structured: python-docx not installed")
        return None

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        logger.warning("_extract_docx_structured: Document() failed: %s", exc)
        return None

    def _para_meta(para: Any) -> Dict[str, Any]:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        runs = [r for r in para.runs if r.text.strip()]
        all_bold = bool(runs and all(r.bold for r in runs))
        all_caps = text == text.upper() and any(c.isalpha() for c in text)
        is_list = any(s in style for s in ("List", "Bullet", "Item")) or \
                  bool(text and text[0] in ("•", "●", "▪", "◦", "○"))
        return {"text": text, "style": style, "all_bold": all_bold,
                "all_caps": all_caps, "is_list": is_list}

    # Traverse document body in reading order, including table cells
    para_data: List[Dict[str, Any]] = []
    seen_cell_texts: set = set()

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            para = DocxParagraph(child, doc)
            meta = _para_meta(para)
            if meta["text"]:
                para_data.append(meta)
        elif tag == "tbl":
            table = DocxTable(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if not cell_text or cell_text in seen_cell_texts:
                        continue
                    seen_cell_texts.add(cell_text)
                    for para in cell.paragraphs:
                        meta = _para_meta(para)
                        if meta["text"]:
                            para_data.append(meta)

    if not para_data:
        return None

    # Detect uniform-bold (like flat PDFs): disable bold-as-heading when >50% lines are bold
    bold_ratio = sum(1 for p in para_data if p["all_bold"]) / len(para_data)
    uniform_bold = bold_ratio > 0.5

    def is_heading(i: int) -> bool:
        p = para_data[i]
        text = p["text"]
        if not text or len(text) > 70:
            return False
        style = p["style"]
        if any(style.startswith(s) for s in ("Heading 1", "Heading 2", "Title")):
            return True
        if p["all_caps"] and len(text) <= 60:
            return True
        if uniform_bold:
            # Fall back to keyword-only when bold is uniform across the document
            return len(text.split()) <= 5 and _det_canonical_section(text) is not None
        if p["all_bold"] and len(text) <= 50 and not p["is_list"]:
            return True
        return False

    plain_lines = [p["text"] for p in para_data]
    return _build_structured_from_lines(plain_lines, is_heading_fn=lambda i: is_heading(i))


# ── Shared line-based structured builder ────────────────────────────────────

def _build_structured_from_lines(
    lines: List[str],
    is_heading_fn: Any,
) -> Optional["UploadedResumeStructured"]:
    """
    Convert a flat list of text lines into UploadedResumeStructured.
    ``is_heading_fn(i)`` returns True when line i is a section heading.
    """
    if not lines:
        return None

    result = UploadedResumeStructured()

    # Contact from header block
    contact = _det_extract_contact(lines)
    result.email    = contact.get("email", "")
    result.phone    = contact.get("phone", "")
    result.linkedin = contact.get("linkedin", "")
    result.github   = contact.get("github", "")
    result.website  = contact.get("website", "")

    # Name: first short line with no special chars
    for text in lines[:5]:
        words = text.split()
        if (
            1 <= len(words) <= 5
            and not re.search(r"[@|/\\]", text)
            and not re.search(r"\d{4}", text)
            and not any(kw in text.lower() for kw in ("resume", "curriculum", "vitae", "cv"))
        ):
            result.full_name = text
            break

    current_section: Optional[str] = None
    current_job: Optional[WorkExperienceItem] = None
    current_edu: Optional[EducationItem] = None
    current_proj: Optional[ProjectItem] = None
    skills_raw: List[str] = []
    summary_lines: List[str] = []

    for i, text in enumerate(lines):
        if is_heading_fn(i):
            sec = _det_canonical_section(text)
            if sec:
                current_section = sec
                current_job = None
                current_edu = None
                current_proj = None
                continue

        date_str = _det_extract_date(text)
        is_bullet = text and text[0] in ("•", "●", "▪", "◦", "○", "–", "—") or \
                    (len(text) > 1 and text[:2] == "- ")
        bullet_text = re.sub(r"^[•●▪◦○–—\-]\s*", "", text).strip() if is_bullet else ""

        if current_section == "summary":
            if not is_heading_fn(i):
                summary_lines.append(text)

        elif current_section == "skills":
            if not is_heading_fn(i):
                skills_raw.append(text)

        elif current_section == "experience":
            if is_bullet and current_job:
                current_job.bullets.append(bullet_text or text)
            elif not is_bullet:
                text_no_date = _EXTRACT_DATE_RE.sub("", text).strip(" |–—-").strip()
                parts = [p.strip() for p in re.split(r"\s*[|,]\s*", text_no_date) if p.strip()]
                if date_str and (not current_job or current_job.title):
                    current_job = WorkExperienceItem(years=date_str)
                    result.work_experience.append(current_job)
                    if len(parts) >= 2:
                        current_job.title = parts[0]
                        current_job.organization = parts[1]
                    elif parts:
                        current_job.title = parts[0]
                elif not current_job or (current_job.title and current_job.organization):
                    current_job = WorkExperienceItem(title=text_no_date or text)
                    result.work_experience.append(current_job)
                elif current_job and not current_job.organization:
                    current_job.organization = text_no_date or text
                elif current_job and date_str and not current_job.years:
                    current_job.years = date_str

        elif current_section == "education":
            if not is_bullet:
                text_no_date = _EXTRACT_DATE_RE.sub("", text).strip(" |–—-").strip()
                parts = [p.strip() for p in re.split(r"\s*[|,]\s*", text_no_date) if p.strip()]
                if not current_edu or (current_edu.school and current_edu.degree):
                    current_edu = EducationItem()
                    result.education.append(current_edu)
                    if len(parts) >= 2:
                        current_edu.school = parts[0]
                        current_edu.degree = parts[1]
                    elif parts:
                        current_edu.school = parts[0]
                    if date_str:
                        current_edu.years = date_str
                elif current_edu and not current_edu.degree:
                    current_edu.degree = text_no_date or text
                elif current_edu and date_str and not current_edu.years:
                    current_edu.years = date_str
                elif current_edu:
                    current_edu.details = (current_edu.details + " " + text).strip()

        elif current_section == "projects":
            if is_bullet and current_proj:
                current_proj.bullets.append(bullet_text or text)
            elif not is_bullet:
                if not current_proj or current_proj.name:
                    current_proj = ProjectItem(name=text, years=date_str)
                    result.projects.append(current_proj)
                elif not current_proj.description:
                    current_proj.description = text

    # Process skills
    if skills_raw:
        raw = " | ".join(skills_raw)
        parts = re.split(r"[,;|\n•●▪]+", raw)
        result.skills = [p.strip() for p in parts if 1 < len(p.strip()) < 80]

    if summary_lines:
        result.summary = " ".join(summary_lines)

    if not _det_quality_ok(result):
        return None

    return result


# ── Smart upload router ─────────────────────────────────────────────────────

def parse_upload_bytes(
    content: bytes,
    filename: str,
    markdown: str,
) -> "tuple[Dict[str, Any], str, str, List[str]]":
    """
    Smart router: tries deterministic extraction first (PyMuPDF for PDF,
    python-docx for DOCX), falls back to LLM only when needed.

    Returns ``(structured_dict, plain_text, status, hints)``.
    ``status`` is one of: ``ready``, ``ready_deterministic``, ``llm_failed``.
    """
    suffix = Path(filename or "resume.pdf").suffix.lower()
    hints: List[str] = []
    det_result: Optional[UploadedResumeStructured] = None

    if suffix == ".pdf":
        try:
            det_result = _extract_pdf_structured(content)
        except Exception as exc:
            logger.warning("parse_upload_bytes: PDF deterministic extract failed: %s", exc)

    elif suffix in (".docx", ".doc"):
        try:
            det_result = _extract_docx_structured(content)
        except Exception as exc:
            logger.warning("parse_upload_bytes: DOCX deterministic extract failed: %s", exc)

    if det_result is not None:
        logger.info(
            "parse_upload_bytes: deterministic extract succeeded for %s "
            "(jobs=%d edu=%d skills=%d)",
            filename,
            len(det_result.work_experience),
            len(det_result.education),
            len(det_result.skills),
        )
        return (
            det_result.model_dump(),
            structured_to_plain_resume_text(det_result),
            "ready_deterministic",
            hints,
        )

    # Fallback: LLM pipeline on markdown text
    logger.info("parse_upload_bytes: deterministic extract insufficient, falling back to LLM for %s", filename)
    return parse_upload_resume_full_pipeline(markdown)


def parse_upload_resume_full_pipeline(markdown: str) -> tuple[Dict[str, Any], str, str, List[str]]:
    """
    Returns ``(structured_dict, plain_text, status, hints)``.

    ``status`` is ``ready`` or ``llm_failed``. ``hints`` is non-empty when structured
    parsing failed but extracted text is still usable.
    """
    hints: List[str] = []
    try:
        structured = llm_resume_markdown_to_structured(markdown)
        model = UploadedResumeStructured.model_validate(structured)
        return structured, structured_to_plain_resume_text(model), "ready", hints
    except Exception as exc:
        logger.warning("resume_upload_parse: LLM pipeline failed, falling back to markdown-only text: %s", exc)
        fallback = re.sub(r"^#+\s*", "", markdown, flags=re.MULTILINE).strip()
        hints.append(
            "Structured résumé parsing did not finish (model unavailable or output invalid). "
            "You can still edit the extracted text below — try again later for automatic section layout."
        )
        return {}, fallback, "llm_failed", hints
